# -*- coding: utf-8 -*-
import os
import re
import asyncio
import tempfile
import subprocess
import time
import json
import html
import logging
import random
import shutil
import textwrap
from datetime import datetime, date
from collections import defaultdict
from io import BytesIO
from pathlib import Path
import base64
from contextvars import ContextVar
import xml.etree.ElementTree as ET

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except Exception:
    plt = None

import yt_dlp
import httpx
from PIL import Image
import qrcode
import jdatetime
from docx import Document

# Optional PDF reader for the Smart File Reader. Install with: pip install pypdf
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None
from bs4 import BeautifulSoup

# Optional advanced math engine. Install with: pip install sympy
try:
    import sympy as sp
except ImportError:
    sp = None

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters
)
from telegram.constants import ChatAction


# =========================================================
# CONFIG
# =========================================================

# =========================================================
# API KEYS / TOKENS
# فعلاً کلیدها را مستقیم همین‌جا قرار بده.
# بعداً برای امنیت بیشتر می‌توانی آن‌ها را به Environment Variable منتقل کنی.
# =========================================================
TOKEN = os.getenv("BOT_TOKEN", "")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
AUDD_API_KEY = os.getenv("AUDD_API_KEY", "")
KURDISH_STT_KEY = os.getenv("KURDISH_STT_KEY", "")

GROQ_CHAT_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_CHAT_MODEL = "openai/gpt-oss-120b"
# Vision model for image question solving and OCR.
GROQ_VISION_MODEL = "qwen/qwen3.6-27b"

ADMIN_IDS = [8946824076]  # آیدی عددی ادمین — در صورت نیاز آیدی خودت را اضافه کن

MAX_FILE_SIZE = 49 * 1024 * 1024
MAX_DOCUMENT_SIZE = 2 * 1024 * 1024 * 1024

# Smart File Reader limits
SMART_FILE_MAX_SIZE = 25 * 1024 * 1024
SMART_FILE_MAX_CHARS = 60000
DAILY_LIMIT = 25
DAILY_CONFIG_LIMIT = 20
MAX_SUBTITLE_VIDEO_DURATION = 90

DOWNLOAD_SEMAPHORE = asyncio.Semaphore(1)
UPLOAD_SEMAPHORE = asyncio.Semaphore(2)

STATS_FILE = "bot_stats.json"
BLACKLIST_FILE = "blacklist.json"
COOKIES_FILE = "cookies.txt"
LOG_FILE = "bot.log"
CONFIG_USAGE_FILE = "config_usage.json"

CONFIG_SOURCES = {
    "vless": [
        "https://raw.githubusercontent.com/Epodonios/v2ray-configs/main/Splitted-By-Protocol/vless.txt",
        "https://raw.githubusercontent.com/barry-far/V2ray-config/main/Splitted-By-Protocol/vless.txt",
    ],
    "vmess": [
        "https://raw.githubusercontent.com/Epodonios/v2ray-configs/main/Splitted-By-Protocol/vmess.txt",
        "https://raw.githubusercontent.com/barry-far/V2ray-config/main/Splitted-By-Protocol/vmess.txt",
    ],
    "trojan": [
        "https://raw.githubusercontent.com/Epodonios/v2ray-configs/main/Splitted-By-Protocol/trojan.txt",
        "https://raw.githubusercontent.com/barry-far/V2ray-config/main/Splitted-By-Protocol/trojan.txt",
    ],
    "ss": [
        "https://raw.githubusercontent.com/Epodonios/v2ray-configs/main/Splitted-By-Protocol/ss.txt",
        "https://raw.githubusercontent.com/barry-far/V2ray-config/main/Splitted-By-Protocol/ss.txt",
    ],
}


# =========================================================
# LOGGING
# =========================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# =========================================================
# DATA
# =========================================================

user_downloads = defaultdict(lambda: {"date": str(datetime.now().date()), "count": 0})
user_config_usage = defaultdict(lambda: {"date": str(date.today()), "count": 0})
total_downloads = 0
blacklist = set()
cancel_flags = {}
price_alerts = defaultdict(list)
all_users = set()  # همه کاربرانی که /start زده‌اند
user_join_dates = {}  # user_id -> ISO date string
reminders = defaultdict(list)  # user_id -> list of {text, due_ts, created}
ALERTS_FILE = "price_alerts.json"
USERS_FILE = "users.json"
JOIN_DATES_FILE = "join_dates.json"
REMINDERS_FILE = "reminders.json"



def load_data():
    global total_downloads, user_downloads, blacklist, user_config_usage, all_users, user_join_dates, reminders
    try:
        if os.path.exists(STATS_FILE):
            with open(STATS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                total_downloads = data.get("total", 0)
                user_downloads = defaultdict(
                    lambda: {"date": str(datetime.now().date()), "count": 0},
                    {int(k): v for k, v in data.get("users", {}).items()}
                )
    except Exception as e:
        logger.error(f"Load stats error: {e}")

    try:
        if os.path.exists(BLACKLIST_FILE):
            with open(BLACKLIST_FILE, "r", encoding="utf-8") as f:
                blacklist = set(json.load(f))
    except Exception as e:
        logger.error(f"Load blacklist error: {e}")

    try:
        if os.path.exists(CONFIG_USAGE_FILE):
            with open(CONFIG_USAGE_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                user_config_usage = defaultdict(
                    lambda: {"date": str(date.today()), "count": 0},
                    {int(k): v for k, v in data.items()}
                )
    except Exception as e:
        logger.error(f"Load config usage error: {e}")

    try:
        if os.path.exists(ALERTS_FILE):
            with open(ALERTS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                price_alerts.clear()
                for uid, alerts in data.items():
                    price_alerts[int(uid)] = alerts
    except Exception as e:
        logger.error(f"Load alerts error: {e}")

    try:
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE, "r", encoding="utf-8") as f:
                all_users.clear()
                all_users.update(int(x) for x in json.load(f))
    except Exception as e:
        logger.error(f"Load users error: {e}")
    # مهاجرت: کاربران دانلود را هم به لیست کل اضافه کن
    all_users.update(user_downloads.keys())

    try:
        if os.path.exists(JOIN_DATES_FILE):
            with open(JOIN_DATES_FILE, "r", encoding="utf-8") as f:
                user_join_dates.clear()
                user_join_dates.update({int(k): v for k, v in json.load(f).items()})
    except Exception as e:
        logger.error(f"Load join dates error: {e}")

    try:
        if os.path.exists(REMINDERS_FILE):
            with open(REMINDERS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                reminders.clear()
                for uid, items in data.items():
                    reminders[int(uid)] = items
    except Exception as e:
        logger.error(f"Load reminders error: {e}")


def save_alerts():
    try:
        with open(ALERTS_FILE, "w", encoding="utf-8") as f:
            json.dump({str(k): v for k, v in price_alerts.items()}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Save alerts error: {e}")


def save_stats():
    try:
        with open(STATS_FILE, "w", encoding="utf-8") as f:
            json.dump({"total": total_downloads, "users": {str(k): v for k, v in user_downloads.items()}}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Save stats error: {e}")


def save_config_usage():
    try:
        with open(CONFIG_USAGE_FILE, "w", encoding="utf-8") as f:
            json.dump({str(k): v for k, v in user_config_usage.items()}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Save config usage error: {e}")


def save_blacklist():
    try:
        with open(BLACKLIST_FILE, "w", encoding="utf-8") as f:
            json.dump(list(blacklist), f)
    except Exception as e:
        logger.error(f"Save blacklist error: {e}")


def save_users():
    try:
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(sorted(all_users), f)
    except Exception as e:
        logger.error(f"Save users error: {e}")


def save_join_dates():
    try:
        with open(JOIN_DATES_FILE, "w", encoding="utf-8") as f:
            json.dump({str(k): v for k, v in user_join_dates.items()}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Save join dates error: {e}")


def save_reminders():
    try:
        with open(REMINDERS_FILE, "w", encoding="utf-8") as f:
            json.dump({str(k): v for k, v in reminders.items()}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Save reminders error: {e}")


def register_user(user_id: int):
    """ثبت کاربر برای آمار و پیام همگانی"""
    if user_id not in all_users:
        all_users.add(user_id)
        save_users()
    if user_id not in user_join_dates:
        user_join_dates[user_id] = str(date.today())
        save_join_dates()


def check_daily_limit(user_id: int) -> bool:
    today = str(datetime.now().date())
    info = user_downloads[user_id]
    if info["date"] != today:
        info["date"] = today
        info["count"] = 0
    return info["count"] < DAILY_LIMIT


def increase_download_count(user_id: int):
    global total_downloads
    today = str(datetime.now().date())
    info = user_downloads[user_id]
    if info["date"] != today:
        info["date"] = today
        info["count"] = 0
    info["count"] += 1
    total_downloads += 1
    save_stats()


def check_config_limit(user_id: int) -> bool:
    today = str(date.today())
    info = user_config_usage[user_id]
    if info["date"] != today:
        info["date"] = today
        info["count"] = 0
    return info["count"] < DAILY_CONFIG_LIMIT


def increase_config_count(user_id: int):
    today = str(date.today())
    info = user_config_usage[user_id]
    if info["date"] != today:
        info["date"] = today
        info["count"] = 0
    info["count"] += 1
    save_config_usage()


def is_admin(user_id: int) -> bool:
    try:
        return int(user_id) in ADMIN_IDS
    except Exception:
        return False


# =========================================================
# STATIC MULTILINGUAL UI (NO GROQ FOR UI)
# =========================================================
LANGUAGES_FILE = "user_languages.json"
SUPPORTED_LANGUAGES = ("fa", "ckb", "en")
LANGUAGE_NAMES = {"fa": "🇮🇷 فارسی", "ckb": "☀️ کوردی", "en": "🇬🇧 English"}
user_languages = {}
_CURRENT_UI_LANG = ContextVar("asoland_ui_lang", default="fa")

UI_TRANSLATIONS = {'🏠 منوی اصلی': ('🏠 Main Menu', '🏠 سەرەکی'), '🔙 تغییر پس\u200cزمینه': ('🔙 Change Background', '🔙 گۆڕینی پاشبنەما'), '❌ لغو': ('❌ Cancel', '❌ هەڵوەشاندنەوە'), '🔙 تغییر رنگ متن': ('🔙 Change Text Color', '🔙 گۆڕینی ڕەنگی دەق'), '🌤 آب\u200cوهوا': ('🌤 Weather', '🌤 کەش و هەوا'), '🔮 فال روزانه': ('🔮 Daily Fortune', '🔮 بەختی ڕۆژانە'), '🧮 محاسبه\u200cگر': ('🧮 Calculator', '🧮 حیسابکەر'), '💱 تبدیل ارز': ('💱 Currency Converter', '💱 گۆڕینی دراو'), '📅 تقویم امروز': ("📅 Today's Calendar", '📅 ڕۆژژمێری ئەمڕۆ'), '🔔 یادآور': ('🔔 Reminder', '🔔 بیرخەرەوە'), '👤 پنل من': ('👤 My Panel', '👤 پەنێڵی من'), '📰 اخبار': ('📰 News', '📰 هەواڵ'), '🎓 ابزار دانشجویی': ('🎓 Student Tools', '🎓 ئامرازە خوێندکارییەکان'), '📸 حل سؤال از عکس': ('📸 Solve from Image', '📸 چارەسەری پرسیار لە وێنە'), '🔤 OCR عکس': ('🔤 Image OCR', '🔤 OCRی وێنە'), '💰 قیمت دلار و سکه': ('💰 Dollar & Coin Prices', '💰 نرخی دۆلار و دراو'), '🖼 با عکس': ('🖼 With Photo', '🖼 لەگەڵ وێنە'), '🎨 بدون عکس': ('🎨 Without Photo', '🎨 بەبێ وێنە'), '📥 دانلود فیلم و ویدیو': ('📥 Download Video', '📥 داگرتنی ڤیدیۆ'), '📸 اینستاگرام': ('📸 Instagram', '📸 ئینستاگرام'), '▶️ یوتیوب': ('▶️ YouTube', '▶️ یوتیوب'), '🎵 تیک\u200cتاک': ('🎵 TikTok', '🎵 تیک\u200cتۆک'), '🐦 توییتر': ('🐦 Twitter', '🐦 تویتەر'), '🔴 ردیت': ('🔴 Reddit', '🔴 ڕێدیت'), '🔍 جستجوی یوتیوب': ('🔍 YouTube Search', '🔍 گەڕانی یوتیوب'), '🔗 تبدیل ویدیو به MP3': ('🔗 Video to MP3', '🔗 گۆڕینی ڤیدیۆ بۆ MP3'), '🎵 موسیقی و آهنگ': ('🎵 Music', '🎵 میوزیک و گۆرانی'), '🎶 پیدا کردن آهنگ': ('🎶 Find a Song', '🎶 دۆزینەوەی گۆرانی'), '📝 متن آهنگ': ('📝 Lyrics', '📝 دەقی گۆرانی'), '📜 زیرنویس یوتیوب': ('📜 YouTube Subtitles', '📜 ژێرنووسی یوتیوب'), '🚀 امکانات هوشمند': ('🚀 Smart Features', '🚀 تایبەتمەندییە زیرەکەکان'), '🛠 ابزارهای کاربردی': ('🛠 Useful Tools', '🛠 ئامرازە بەسوودەکان'), '🖼 ساخت استیکر': ('🖼 Create Sticker', '🖼 دروستکردنی ستیکەر'), '📱 QR Code': ('📱 QR Code', '📱 کۆدی QR'), '🗜 فشرده عکس': ('🗜 Compress Image', '🗜 پەستاندنی وێنە'), '🗜 فشرده ویدیو': ('🗜 Compress Video', '🗜 پەستاندنی ڤیدیۆ'), '📄 عکس به PDF': ('📄 Photos to PDF', '📄 وێنە بۆ PDF'), '📝 متن به ورد': ('📝 Text to Word', '📝 دەق بۆ Word'), '✍ زیباسازی متن': ('✍ Fancy Text', '✍ جوانکردنی دەق'), '🔗 کوتاه لینک': ('🔗 Shorten Link', '🔗 کورتکردنەوەی بەستەر'), '🕔 ساعت': ('🕔 Clock', '🕔 کات'), '📁 فایل\u200cخوان هوشمند': ('📁 Smart File Reader', '📁 خوێنەری زیرەکی فایل'), '🤖 هوش مصنوعی': ('🤖 AI', '🤖 زیرەکی دەستکرد'), '💬 چت AI': ('💬 AI Chat', '💬 گفتوگۆی AI'), '🌐 ترجمه': ('🌐 Translate', '🌐 وەرگێڕان'), '📝 خلاصه\u200cسازی': ('📝 Summarize', '📝 کورتەکردنەوە'), '🎤 ویس به متن': ('🎤 Voice to Text', '🎤 دەنگ بۆ دەق'), '🇬🇧 معلم انگلیسی': ('🇬🇧 English Teacher', '🇬🇧 مامۆستای ئینگلیزی'), '🎬 ساخت زیرنویس (کوردی / فارسی / انگلیسی)': ('🎬 Create Subtitles (Kurdish / Persian / English)', '🎬 دروستکردنی ژێرنووس (کوردی / فارسی / ئینگلیزی)'), '🩷 کانفیگ رایگان': ('🩷 Free Config', '🩷 کۆنفیگی بەخۆڕایی'), 'ℹ️ راهنما': ('ℹ️ Help', 'ℹ️ ڕێنمایی'), '📩 پشتیبانی': ('📩 Support', '📩 پشتگیری'), '☀️ کوردی': ('🟡 Sorani Kurdish', '☀️ کوردی'), '🇮🇷 فارسی': ('🇮🇷 Persian', '🇮🇷 فارسی'), '🇬🇧 انگلیسی': ('🇬🇧 English', '🇬🇧 ئینگلیزی'), '📄 دریافت فایل SRT': ('📄 Get SRT File', '📄 وەرگرتنی فایلی SRT'), '🎥 دریافت ویدیوی زیرنویس\u200cدار': ('🎥 Get Subtitled Video', '🎥 وەرگرتنی ڤیدیۆی ژێرنووسکراو'), '📦 دریافت هر دو': ('📦 Get Both', '📦 هەردووکی وەرگرە'), '🎲 تصادفی': ('🎲 Random', '🎲 هەڕەمەکی'), '🔙 برگشت': ('🔙 Back', '🔙 گەڕانەوە'), '💬 مکالمه آزاد': ('💬 Free Conversation', '💬 گفتوگۆی ئازاد'), '✍️ تصحیح جمله من': ('✍️ Correct My Sentence', '✍️ ڕاستکردنەوەی ڕستەکەم'), '📚 یادگیری لغت جدید': ('📚 Learn New Words', '📚 فێربوونی وشەی نوێ'), '📖 تمرین گرامر': ('📖 Grammar Practice', '📖 ڕاهێنانی ڕێزمان'), '🎯 تنظیم سطح من': ('🎯 Set My Level', '🎯 دیاریکردنی ئاستی من'), '❌ خروج از معلم زبان': ('❌ Exit Teacher', '❌ دەرچوون لە مامۆستا'), '🟢 مبتدی': ('🟢 Beginner', '🟢 سەرەتایی'), '🟡 متوسط': ('🟡 Intermediate', '🟡 مامناوەند'), '🔴 پیشرفته': ('🔴 Advanced', '🔴 پێشکەوتوو'), '📊 آمار ربات': ('📊 Bot Statistics', '📊 ئاماری بۆت'), '📢 پیام همگانی': ('📢 Broadcast', '📢 پەیامی گشتی'), '🚫 بلاک\u200cلیست': ('🚫 Blacklist', '🚫 لیستی ڕێگریکراو'), '🎵 فقط صدا (MP3)': ('🎵 Audio Only (MP3)', '🎵 تەنها دەنگ (MP3)'), '❌لغو': ('❌ Cancel', '❌ هەڵوەشاندنەوە'), '🔄 کیفیت دیگر': ('🔄 Other Quality', '🔄 کوالیتیی تر'), '🎵 صدا (MP3)': ('🎵 Audio (MP3)', '🎵 دەنگ (MP3)'), '❌خروج از چت': ('❌ Exit Chat', '❌ دەرچوون لە گفتوگۆ'), '📌 خلاصه': ('📌 Summary', '📌 کورتە'), '🔍 جستجو': ('🔍 Search', '🔍 گەڕان'), '❓ سؤال از فایل': ('❓ Ask the File', '❓ پرسیار لە فایل'), '📝 استخراج نکات': ('📝 Extract Notes', '📝 دەرهێنانی خاڵەکان'), '🎓 ساخت آزمون': ('🎓 Create Quiz', '🎓 دروستکردنی تاقیکردنەوە'), '🧠 فلش\u200cکارت': ('🧠 Flashcards', '🧠 کارتی فێربوون'), '🗑 حذف فایل': ('🗑 Delete File', '🗑 سڕینەوەی فایل'), '₿ بیت\u200cکوین ۷ روزه': ('₿ Bitcoin 7 Days', '₿ بیتکۆین ٧ ڕۆژ'), '🔷 اتریوم ۷ روزه': ('🔷 Ethereum 7 Days', '🔷 ئیتریۆم ٧ ڕۆژ'), '🌤 نمودار دمای شهر': ('🌤 City Temperature Chart', '🌤 هێڵکاری گەرمای شار'), '🌠عمومی': ('🌠 General', '🌠 گشتی'), '💻 فناوری': ('💻 Technology', '💻 تەکنەلۆجیا'), '💵 اقتصاد': ('💵 Economy', '💵 ئابووری'), '₿ کریپتو': ('₿ Crypto', '₿ کریپتۆ'), '🎓 آموزش': ('🎓 Education', '🎓 پەروەردە'), '🔄 بروزرسانی': ('🔄 Refresh', '🔄 نوێکردنەوە'), '📸 حل از عکس': ('📸 Solve from Image', '📸 چارەسەر لە وێنە'), '🔄 فال دوباره': ('🔄 Fortune Again', '🔄 بەخت دووبارە'), '🔤 تغییر فونت': ('🔤 Change Font', '🔤 گۆڕینی فۆنت'), '🔙 تغییر فونت': ('🔙 Change Font', '🔙 گۆڕینی فۆنت'), '🎨 تغییر رنگ\u200cها': ('🎨 Change Colors', '🎨 گۆڕینی ڕەنگەکان'), '🩷 کانفیگ جدید': ('🩷 New Config', '🩷 کۆنفیگی نوێ'), '🔙 منوی معلم': ('🔙 Teacher Menu', '🔙 پێڕستی مامۆستا'), '❌انصراف': ('❌ Cancel', '❌ هەڵوەشاندنەوە'), 'تشخیص خودکار': ('Auto Detect', 'دۆزینەوەی خۆکار'), '⬇️ دانلود MP3': ('⬇️ Download MP3', '⬇️ داگرتنی MP3'), '📠متن آهنگ': ('📠 Lyrics', '📠 دەقی گۆرانی'), '🎤 ویس جدید': ('🎤 New Voice', '🎤 دەنگی نوێ'), '💱 تبدیل جدید': ('💱 New Conversion', '💱 گۆڕینەوەی نوێ'), '🌤 شهر دیگر': ('🌤 Another City', '🌤 شارێکی تر'), '🧮 محاسبه جدید': ('🧮 New Calculation', '🧮 حیسابکردنی نوێ'), '🎓 سؤال جدید': ('🎓 New Question', '🎓 پرسیاری نوێ'), '📠خلاصه جدید': ('📠 New Summary', '📠 کورتەی نوێ'), '⬇️ دانلود ویدیو': ('⬇️ Download Video', '⬇️ داگرتنی ڤیدیۆ'), '🎵 دانلود MP3': ('🎵 Download MP3', '🎵 داگرتنی MP3'), '📸 سؤال جدید': ('📸 New Question', '📸 پرسیاری نوێ'), '🎤 دوباره': ('🎤 Again', '🎤 دووبارە'), '🎨 تغییر رنگ متن': ('🎨 Change Text Color', '🎨 گۆڕینی ڕەنگی دەق'), '📸 عکس جدید': ('📸 New Photo', '📸 وێنەی نوێ'), '🎨 ساخت بدون عکس': ('🎨 Create Without Photo', '🎨 دروستکردن بەبێ وێنە'), '🎨 تغییر پس\u200cزمینه': ('🎨 Change Background', '🎨 گۆڕینی پاشبنەما'), '🖼 ساخت با عکس': ('🖼 Create With Photo', '🖼 دروستکردن لەگەڵ وێنە'), '✅ متن جدید': ('✅ New Text', '✅ دەقی نوێ'), '🔤 عکس جدید': ('🔤 New Photo', '🔤 وێنەی نوێ'), '🔄 تلاش دوباره': ('🔄 Try Again', '🔄 هەوڵدانەوە')}
UI_PHRASES = {'فارسی': ('Persian', 'فارسی'), 'کوردی': ('Sorani Kurdish', 'کوردی'), 'انگلیسی': ('English', 'ئینگلیزی'), 'منوی اصلی:': ('Main menu:', 'پێڕستی سەرەکی:'), 'از منوی زیر استفاده کن.': ('Use the menu below.', 'لە پێڕستی خوارەوە بەکاربهێنە.'), 'ربات همه\u200cکاره و هوشمند': ('All-in-one smart bot', 'بۆتی زیرەکی هەمووکارە'), 'راهنمای AsoLand': ('AsoLand Help', 'ڕێنمایی AsoLand'), 'پنل مدیریت': ('Admin Panel', 'پەنێڵی بەڕێوەبەرایەتی'), 'آمار، پیام همگانی و بلاک\u200cلیست از دکمه\u200cهای زیر:': ('Statistics, broadcast and blacklist are below:', 'ئامار، پەیامی گشتی و لیستی ڕێگریکراو لە خوارەوەن:'), 'در حال دریافت': ('Fetching', 'لە وەرگرتندایە'), 'در حال پردازش': ('Processing', 'لە پرۆسەکردندایە'), 'در حال ارسال': ('Sending', 'لە ناردندایە'), 'دوباره تلاش کن': ('Try again', 'دووبارە هەوڵ بدە'), 'لینک معتبر': ('valid link', 'بەستەری دروست'), 'پیدا نشد': ('Not found', 'نەدۆزرایەوە'), 'آماده شد': ('Ready', 'ئامادەیە'), 'ارسال شد': ('Sent', 'نێردرا'), 'لغو شد': ('Cancelled', 'هەڵوەشێندرایەوە'), 'دسترسی ندارید': ('Access denied', 'دەستگەیشتنت نییە'), 'عکسی نیست': ('No photo', 'وێنە نییە'), 'متن بفرست': ('Send the text', 'دەق بنێرە'), 'جستجو': ('Search', 'گەڕان'), 'زبان رو انتخاب کن': ('Choose the language', 'زمان هەڵبژێرە'), 'حالا ویس رو بفرست': ('Now send the voice message', 'ئێستا دەنگەکە بنێرە'), 'لطفاً زبان را انتخاب کنید:': ('Please choose a language:', 'تکایە زمانێک هەڵبژێرە:'), 'زبان با موفقیت تغییر کرد.': ('Language changed successfully.', 'زمان بە سەرکەوتوویی گۆڕدرا.'), 'انتخاب زبان': ('Choose Language', 'هەڵبژاردنی زمان'), 'زبان فعلی': ('Current language', 'زمانی ئێستا'), '🎯 سطح خودت رو انتخاب کن:': ('🎯 Choose your level:', '🎯 ئاستەکەت هەڵبژێرە:'), 'متن مورد نظرت رو بفرست.': ('Send your text.', 'دەقەکەت بنێرە.'), 'متن رو بفرست': ('Send the text', 'دەقەکە بنێرە'), 'عکس موردنظرت را بفرست.': ('Send your photo.', 'وێنەکەت بنێرە.'), 'عکس سؤال را بفرست.': ('Send the question image.', 'وێنەی پرسیارەکە بنێرە.'), 'نام شهر را بفرست': ('Send the city name', 'ناوی شارەکە بنێرە'), 'سؤال درسی را بفرست': ('Send the study question', 'پرسیاری خوێندکاری بنێرە'), 'دسته خبر را از منو انتخاب کن': ('Choose a news category from the menu', 'جۆری هەواڵ لە پێڕستەکە هەڵبژێرە'), 'از منوی امکانات هوشمند نمودار را انتخاب کن': ('Choose a chart from Smart Features', 'هێڵکارییەک لە تایبەتمەندییە زیرەکەکان هەڵبژێرە'), 'مقدار را بفرست': ('Send the amount', 'بڕەکە بنێرە'), 'کیفیت رو انتخاب کن': ('Choose the quality', 'کوالیتی هەڵبژێرە'), 'فقط لینک یوتیوب': ('YouTube link only', 'تەنها بەستەری یوتیوب'), 'در حال تبدیل ویس به متن': ('Converting voice to text', 'دەنگ دەگۆڕدرێت بۆ دەق'), 'در حال شناسایی آهنگ': ('Recognizing the song', 'گۆرانی دەناسرێتەوە'), 'در حال فشرده\u200cسازی': ('Compressing', 'لە پەستاندندایە'), 'ساخت PDF': ('Creating PDF', 'دروستکردنی PDF'), 'ساخت زیرنویس': ('Creating subtitles', 'دروستکردنی ژێرنووس'), 'ساخت استیکر': ('Creating sticker', 'دروستکردنی ستیکەر'), 'متن آهنگ': ('Lyrics', 'دەقی گۆرانی'), 'متن پیدا نشد': ('Lyrics not found', 'دەق نەدۆزرایەوە'), 'اطلاعاتی نیست': ('No information', 'زانیاری نییە'), 'لینک منقضی شده': ('The link has expired', 'بەستەرەکە بەسەرچووە'), 'حجم بیش از حد مجاز': ('File is too large', 'قەبارە لە سنوور زیاترە'), 'محدودیت روزانه تموم شده': ('Daily limit reached', 'سنووری ڕۆژانە تەواوبووە'), 'کانفیگ سالم پیدا نشد': ('No valid config found', 'کۆنفیگی دروست نەدۆزرایەوە'), 'خطا در ارتباط با هوش مصنوعی': ('AI connection error', 'هەڵە لە پەیوەندی بە زیرەکی دەستکرد'), 'خطا در خلاصه\u200cسازی متن': ('Summarization error', 'هەڵە لە کورتەکردنەوەی دەق'), 'خطا در ترجمه': ('Translation error', 'هەڵە لە وەرگێڕان'), 'خطا': ('Error', 'هەڵە'), 'نامشخص': ('Unknown', 'نادیار'), 'وضعیت نامشخص': ('Unknown status', 'دۆخی نادیار'), 'آسمان صاف': ('Clear sky', 'ئاسمانی ڕوون'), 'نیمه\u200cابری': ('Partly cloudy', 'تاڕادەیەک هەوراو'), 'ابری': ('Cloudy', 'هەوراو'), 'باران': ('Rain', 'باران'), 'برف': ('Snow', 'بەفر'), 'مه': ('Fog', 'تەم'), 'بالا': ('above', 'سەرەوە'), 'پایین': ('below', 'خوارەوە')}

UI_TRANSLATIONS.update({
    "🌐 زبان / Language": ("🌐 Language", "🌐 زمان"),
    "سلام": ("Hello", "سڵاو"),
    "کاربر": ("User", "بەکارهێنەر"),
    "روز": ("day", "ڕۆژ"),
    "روزانه": ("daily", "ڕۆژانە"),
    "قبل": ("Before", "پێش"),
    "بعد": ("After", "دوای"),
    "صفحه": ("page", "لاپەڕە"),
    "آبی": ("Blue", "شین"),
    "سبز": ("Green", "سەوز"),
    "بنفش": ("Purple", "مۆر"),
    "سفید": ("White", "سپی"),
    "طلایی": ("Gold", "زێڕین"),
})
UI_PHRASES.update({
    "ساعت و تاریخ شمسی": ("Solar Hijri Date & Time", "بەروار و کاتی هەتاوی"),
    "تاریخ شمسی": ("Solar Hijri Date", "بەرواری هەتاوی"),
    "ساعت": ("Clock", "کات"),
    "راهنمای AsoLand": ("AsoLand Help", "ڕێنمایی AsoLand"),
    "دانلود فیلم و ویدیو": ("Download films and videos", "داگرتنی فیلم و ڤیدیۆ"),
    "موسیقی، پیدا کردن آهنگ و متن ترانه": ("Music, song finder and lyrics", "میوزیک، دۆزینەوەی گۆرانی و دەقی گۆرانی"),
    "قیمت دلار و سکه + تبدیل ارز": ("Dollar and coin prices + currency conversion", "نرخی دۆلار و دراو + گۆڕینی دراو"),
    "امکانات هوشمند": ("Smart features", "تایبەتمەندییە زیرەکەکان"),
    "ابزارهای کاربردی": ("Useful tools", "ئامرازە بەسوودەکان"),
    "هوش مصنوعی": ("Artificial intelligence", "زیرەکی دەستکرد"),
    "ساخت زیرنویس": ("Create subtitles", "دروستکردنی ژێرنووس"),
    "کانفیگ رایگان": ("Free config", "کۆنفیگی بەخۆڕایی"),
    "پشتیبانی": ("Support", "پشتگیری"),
    "محدودیت روزانه": ("Daily limits", "سنوورە ڕۆژانەکان"),
    "دانلود": ("Downloads", "داگرتن"),
    "کانفیگ": ("Config", "کۆنفیگ"),
    "آمار": ("Statistics", "ئامار"),
    "پیام همگانی": ("Broadcast", "پەیامی گشتی"),
    "بلاک‌لیست": ("Blacklist", "لیستی ڕێگریکراو"),
    "راهنما": ("Help", "ڕێنمایی"),
    "منوی اصلی": ("Main menu", "پێڕستی سەرەکی"),
    "از منوی زیر استفاده کن.": ("Use the menu below.", "لە پێڕستی خوارەوە بەکاربهێنە."),
    "بفرست": ("Send", "بنێرە"),
    "انتخاب کن": ("Choose", "هەڵبژێرە"),
    "انتخاب کنید": ("Choose", "هەڵبژێرە"),
    "در حال": ("Processing", "لە پرۆسەکردندایە"),
    "آماده شد": ("Ready", "ئامادەیە"),
    "موفقیت": ("success", "سەرکەوتن"),
    "نامشخص": ("Unknown", "نادیار"),
    "خطا": ("Error", "هەڵە"),
    "یافت نشد": ("Not found", "نەدۆزرایەوە"),
    "پیدا نشد": ("Not found", "نەدۆزرایەوە"),
    "لغو شد": ("Cancelled", "هەڵوەشێندرایەوە"),
    "لطفاً": ("Please", "تکایە"),
    "حالا": ("Now", "ئێستا"),
    "متن": ("Text", "دەق"),
    "عکس": ("Photo", "وێنە"),
    "فایل": ("File", "فایل"),
    "ویدیو": ("Video", "ڤیدیۆ"),
    "آهنگ": ("Song", "گۆرانی"),
    "خواننده": ("Artist", "گۆرانیبێژ"),
    "شهر": ("City", "شار"),
    "امروز": ("Today", "ئەمڕۆ"),
    "فردا": ("Tomorrow", "بەیانی"),
    "ساعت": ("Time", "کات"),
    "تقویم": ("Calendar", "ڕۆژژمێر"),
    "دلار": ("Dollar", "دۆلار"),
    "سکه": ("Coin", "دراو"),
    "قیمت": ("Price", "نرخ"),
    "نتیجه": ("Result", "ئەنجام"),
    "درصد": ("Percent", "سەد"),
    "کیفیت": ("Quality", "کوالیتی"),
    "لغت": ("Word", "وشە"),
    "گرامر": ("Grammar", "ڕێزمان"),
    "سؤال": ("Question", "پرسیار"),
    "پاسخ": ("Answer", "وەڵام"),
    "صفحه": ("Page", "لاپەڕە"),
    "روز": ("day", "ڕۆژ"),
    "قبل: ": ("Before: ", "پێش: "),
    "بعد: ": ("After: ", "دوای: "),
    "صفحه": ("page", "لاپەڕە"),
    "روز": ("day", "ڕۆژ"),
    "فایل زیرنویس": ("Subtitle file", "فایلی ژێرنووس"),
    "ویدیو با زیرنویس": ("Subtitled video", "ڤیدیۆی ژێرنووسکراو"),
    "PDF با ": ("PDF with ", "PDF لەگەڵ "),
    "فایل ورد آماده شد.": ("Word file is ready.", "فایلی Word ئامادەیە."),
    "یادآور": ("Reminder", "بیرخەرەوە"),
    "هدف": ("Target", "ئامانج"),
})

# Complete translations for previously hard-coded UI messages.
UI_PHRASES.update({
    "اول مشخص کن استیکر را با چه پس‌زمینه‌ای می‌خواهی:": ("First choose the sticker background:", "سەرەتا دیاری بکە ستیکەرەکە بە چ جۆرە پاشبنەمایەک دەته‌وێت:"),
    "با عکس: عکس خودت را می‌فرستی و متن روی عکس قرار می‌گیرد.": ("With photo: send your photo and the text will be placed on it.", "لەگەڵ وێنە: وێنەکەت دەنێریت و دەقەکە لەسەر وێنەکە دادەنرێت."),
    "بدون عکس: یک رنگ برای پس‌زمینه انتخاب می‌کنی و متن روی آن قرار می‌گیرد.": ("Without photo: choose a background color and the text will be placed on it.", "بەبێ وێنە: ڕەنگێک بۆ پاشبنەما هەڵدەبژێریت و دەقەکە لەسەری دادەنرێت."),
    "اول رنگ پس‌زمینه را انتخاب کن:": ("First choose the background color:", "سەرەتا ڕەنگی پاشبنەما هەڵبژێرە:"),
    "حالا رنگ متن را انتخاب کن:": ("Now choose the text color:", "ئێستا ڕەنگی دەق هەڵبژێرە:"),
    "حالا فونت را انتخاب کن": ("Now choose the font", "ئێستا فۆنت هەڵبژێرە"),
    "فونت موردنظرت را انتخاب کن:": ("Choose your desired font:", "فۆنتی دڵخوازت هەڵبژێرە:"),
    "حالا متن فارسی یا کوردی را بفرست.": ("Now send the text.", "ئێستا دەقەکە بنێرە."),
    "متن فارسی یا کوردی": ("text", "دەق"),
    "اول مشخص کن استیکر": ("First choose the sticker", "سەرەتا ستیکەرەکە هەڵبژێرە"),
    "استیکر متنی": ("Text Sticker", "ستیکەری دەقی"),
    "رنگ پس‌زمینه": ("Background color", "ڕەنگی پاشبنەما"),
    "رنگ متن": ("Text color", "ڕەنگی دەق"),
})


def load_languages():
    global user_languages
    try:
        if os.path.exists(LANGUAGES_FILE):
            with open(LANGUAGES_FILE, "r", encoding="utf-8") as f:
                raw = json.load(f)
                user_languages = {int(k): v for k, v in raw.items() if v in SUPPORTED_LANGUAGES}
    except Exception as e:
        logger.warning(f"Load languages error: {e}")

def save_languages():
    try:
        with open(LANGUAGES_FILE, "w", encoding="utf-8") as f:
            json.dump({str(k): v for k, v in user_languages.items()}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"Save languages error: {e}")

def get_user_lang(user_id: int, context=None) -> str:
    if context is not None:
        lang = context.user_data.get("language")
        if lang in SUPPORTED_LANGUAGES:
            return lang
    return user_languages.get(int(user_id), "fa")

def set_user_lang(user_id: int, lang: str, context=None):
    if lang not in SUPPORTED_LANGUAGES:
        lang = "fa"
    user_languages[int(user_id)] = lang
    if context is not None:
        context.user_data["language"] = lang
    save_languages()


# -------------------------------------------------------------------------
# COMPLETE SORANI UI FALLBACK
# -------------------------------------------------------------------------
# The bot has many dynamic UI messages assembled with f-strings.  A simple
# exact-phrase dictionary cannot catch those.  The fallback below translates
# common Persian UI vocabulary after the exact phrase table has run, while
# leaving URLs, HTML tags, numbers and Latin product names intact.
CKB_UI_EXTRA = {
    "📄 <b>فایل دریافت شد</b>": "📄 <b>فایلەکە وەرگیرا</b>",
    "📝 متن استخراج‌شده:": "📝 دەقی دەرهێنراو:",
    "چه کاری انجام بدم؟": "چ کارێک ئەنجام بدەم؟",
    "کاراکتر": "نووسە",
    "در حال دریافت قیمت دلار و سکه...": "لە وەرگرتنی نرخی دۆلار و دراودایە...",
    "قیمت دلار و سکه": "نرخی دۆلار و دراو",
    "دلار تهران": "دۆلاری تاران",
    "طلای ۱۸ عیار (هر گرم)": "زێڕی ١٨ عیار (هەر گرام)",
    "مثقال طلا": "موسقاڵی زێڕ",
    "سکه جدید (امامی)": "دراوی نوێ (ئیمامی)",
    "سکه قدیم (بهار)": "دراوی کۆن (بەهار)",
    "نیم سکه": "نیوە دراو",
    "ربع سکه": "چارەکی دراو",
    "سکه گرمی": "دراوی گرامی",
    "نرخ دلار": "نرخی دۆلار",
    "منبع: کانال تلگرام": "سەرچاوە: کەناڵی تەلەگرام",
    "❌ فعلاً قیمت در دسترس نیست.": "❌ ئێستا نرخ بەردەست نییە.",
    "اتصال اینترنت را چک کن و دوباره امتحان کن.": "پەیوەندی ئینتەرنێت بپشکنە و دووبارە هەوڵ بدە.",
    # Core UI / actions
    "اول مشخص کن استیکر را با چه پس\u200cزمینه\u200cای می\u200cخواهی:": "سەرەتا دیاری بکە ستیکەرەکە بە چ جۆرە پاشبنەمایەک دەته\u200cوێت:",
    "با عکس: عکس خودت را می\u200cفرستی و متن روی عکس قرار می\u200cگیرد.": "لەگەڵ وێنە: وێنەکەت دەنێریت و دەقەکە لەسەر وێنەکە دادەنرێت.",
    "بدون عکس: یک رنگ برای پس\u200cزمینه انتخاب می\u200cکنی و متن روی آن قرار می\u200cگیرد.": "بەبێ وێنە: ڕەنگێک بۆ پاشبنەما هەڵدەبژێریت و دەقەکە لەسەری دادەنرێت.",
    "اول رنگ پس\u200cزمینه را انتخاب کن:": "سەرەتا ڕەنگی پاشبنەما هەڵبژێرە:",
    "حالا رنگ متن را انتخاب کن:": "ئێستا ڕەنگی دەق هەڵبژێرە:",
    "حالا فونت را انتخاب کن": "ئێستا فۆنت هەڵبژێرە",
    "فونت موردنظرت را انتخاب کن:": "فۆنتی دڵخوازت هەڵبژێرە:",
    "حالا متن فارسی یا کوردی را بفرست.": "ئێستا دەقەکە بنێرە.",
    "متن فارسی یا کوردی": "دەقەکە بنێرە",
    "اول مشخص کن استیکر": "سەرەتا ستیکەرەکە هەڵبژێرە",
    "استیکر متنی": "ستیکەری دەقی",
    "رنگ پس\u200cزمینه": "ڕەنگی پاشبنەما",
    "رنگ متن": "ڕەنگی دەق",
    "❌نام شهر معتبر نیست.": "❌ ناوی شار دروست نییە.",
    "❌ فعلاً خبری دریافت نشد.\nممکن است موقتاً دسترسی به خبرگزار قطع باشد. چند دقیقه بعد دوباره تلاش کن.": "❌ ئێستا هیچ هەواڵێک بەردەست نییە.\nلەوانەیە دەستگەیشتن بە سەرچاوەی هەواڵ کاتێکی کەم پچڕابێت. چەند خولەکێک دواتر دووبارە هەوڵ بدە.",
    "❌ فعلاً نرخ دلار در دسترس نیست. چند دقیقه بعد دوباره امتحان کن.": "❌ ئێستا نرخی دۆلار بەردەست نییە. چەند خولەکێک دواتر دووبارە هەوڵ بدە.",
    "❌ دریافت آب\u200cوهوا ناموفق بود. چند لحظه بعد دوباره امتحان کن.": "❌ وەرگرتنی کەش و هەوا سەرکەوتوو نەبوو. چەند ساتێک دواتر دووبارە هەوڵ بدە.",
    "❌ خطا در دریافت قیمت.\nچند لحظه بعد دوباره امتحان کن.": "❌ هەڵە لە وەرگرتنی نرخ.\nچەند ساتێک دواتر دووبارە هەوڵ بدە.",
    "📡 منبع: کانال تلگرام @DO_L4": "📡 سەرچاوە: کەناڵی تەلەگرام @DO_L4",
    "❌ کلید هوش مصنوعی تنظیم نشده است.": "❌ کلیلەی زیرەکی دەستکرد ڕێک نەخراوە.",
    "❌ پاسخی دریافت نشد.": "❌ هیچ وەڵامێک وەرنەگیرا.",
    "❌ کلید Groq تنظیم نشده است.": "❌ کلیلەی Groq ڕێک نەخراوە.",
    "❌ خطا در خلاصه\u200cسازی متن.": "❌ هەڵە لە کورتەکردنەوەی دەق.",
    "❌ خطا در ترجمه.": "❌ هەڵە لە وەرگێڕان.",
    "❌ استخراج متن از تصویر ناموفق بود.": "❌ دەرهێنانی دەق لە وێنە سەرکەوتوو نەبوو.",
    "❌ حل سؤال از روی تصویر فعلاً در دسترس نیست. مطمئن شو مدل Vision در حساب API فعال است.": "❌ چارەسەرکردنی پرسیار لە وێنە ئێستا بەردەست نییە. دڵنیابە مۆدێلی Vision لە هەژماری API چالاکە.",
    "❌ فرمت فایل پشتیبانی نمی\u200cشود. فقط PDF، DOCX و TXT مجاز است.": "❌ فۆرماتی فایل پشتگیری ناکرێت. تەنها PDF، DOCX و TXT ڕێگەپێدراون.",
    "❌ حجم فایل بیشتر از ۲۵ مگابایت است.": "❌ قەبارەی فایل لە ٢٥ مێگابایت زیاترە.",
    "⏳ در حال پردازش فایل...": "⏳ فایلەکە لە ژێر پرۆسەکردندایە...",
    "📥 فایل دریافت شد.\n⏳ در حال استخراج متن...": "📥 فایلەکە وەرگیرا.\n⏳ دەقەکە دەردەهێنرێت...",
    "❌ از این فایل متن قابل خواندن پیدا نشد.": "❌ هیچ دەقێکی خوێندنەوە لەم فایلەدا نەدۆزرایەوە.",
    "❌ فایل فعالی وجود ندارد.": "❌ هیچ فایلێکی چالاک نییە.",
    "🗑 فایل از حافظه این جلسه حذف شد.": "🗑 فایلەکە لە بیرگەی ئەم دانیشتنە سڕایەوە.",
    "🔍 عبارت موردنظر برای جستجو در فایل را بفرست.": "🔍 وشە یان دەستەواژەی گەڕان لە فایلەکە بنێرە.",
    "❓ سؤالت را درباره محتوای فایل بنویس.": "❓ پرسیارەکەت دەربارەی ناوەڕۆکی فایلەکە بنووسە.",
    "🌐 زبان مقصد را بنویس. مثال: انگلیسی، عربی، فارسی، کردی": "🌐 زمانی ئامانج بنووسە. نموونە: ئینگلیزی، عەرەبی، فارسی، کوردی",
    "📁 عملیات دیگری روی همین فایل؟": "📁 کارێکی تری لەسەر هەمان فایل دەکەیت؟",
    "🔍 چیزی با این عبارت در متن استخراج\u200cشده پیدا نشد.": "🔍 هیچ شتێک بەو دەستەواژەیە لە دەقی دەرهێنراو نەدۆزرایەوە.",
    "🎓 سؤال درسی را بفرست.": "🎓 پرسیاری خوێندکاری بنێرە.",
    "📰 دسته خبر را از منو انتخاب کن.": "📰 جۆری هەواڵ لە پێڕستەکە هەڵبژێرە.",
    "📊 از منوی امکانات هوشمند نمودار را انتخاب کن.": "📊 هێڵکارییەک لە پێڕستی تایبەتمەندییە زیرەکەکان هەڵبژێرە.",
    "🌤 نام شهر را بفرست.": "🌤 ناوی شار بنێرە.",
    "💱 مقدار را بفرست؛ مثال: ۱۰۰ دلار": "💱 بڕەکە بنێرە؛ نموونە: ١٠٠ دۆلار",
    "🔔 مثال: یادآوری ساعت ۱۰ جلسه": "🔔 نموونە: بیرخەرەوە کاتژمێر ١٠ کۆبوونەوە",
    "🎨 رنگ جدید متن را انتخاب کن:": "🎨 ڕەنگی نوێی دەق هەڵبژێرە:",
    "🔤 فونت موردنظرت را انتخاب کن:": "🔤 فۆنتی دڵخوازت هەڵبژێرە:",
    "📸 عکس سؤال را بفرست.": "📸 وێنەی پرسیارەکە بنێرە.",
    "📸 عکس پیدا نشد. دوباره عکس را بفرست.": "📸 وێنە نەدۆزرایەوە. دووبارە وێنەکە بنێرە.",
    "📸 ابتدا عکس را بفرست.": "📸 سەرەتا وێنەکە بنێرە.",
    "🎨 ابتدا رنگ متن را انتخاب کن:": "🎨 سەرەتا ڕەنگی دەق هەڵبژێرە:",
    "ابتدا رنگ پس\u200cزمینه را انتخاب کن:": "سەرەتا ڕەنگی پاشبنەما هەڵبژێرە:",
    "ابتدا رنگ\u200cها را انتخاب کن:": "سەرەتا ڕەنگەکان هەڵبژێرە:",
    "فونت نامعتبر": "فۆنت دروست نییە",
    "رنگ نامعتبر": "ڕەنگ دروست نییە",
    "📱 متن یا لینک بفرست تا QR بسازم.": "📱 دەق یان بەستەر بنێرە تا QR دروست بکەم.",
    "🗜 عکس بفرست تا فشرده کنم.": "🗜 وێنە بنێرە تا قەبارەکەی کەم بکەم.",
    "🗜 ویدیو بفرست تا فشرده کنم.": "🗜 ڤیدیۆ بنێرە تا قەبارەکەی کەم بکەم.",
    "🔗 لینک بلند رو بفرست.": "🔗 بەستەری درێژ بنێرە.",
    "🎶 اسم آهنگ، ویس یا لینک بفرست.": "🎶 ناوی گۆرانی، دەنگ یان بەستەر بنێرە.",
    "🔗 لینک ویدیو رو بفرست تا به MP3 تبدیل کنم.": "🔗 بەستەری ڤیدیۆ بنێرە تا بیکەم بە MP3.",
    "📜 لینک یوتیوب رو بفرست.": "📜 بەستەری یوتیوب بنێرە.",
    "🔠اسم ویدیو یا آهنگ رو بنویس.": "🔠 ناوی ڤیدیۆ یان گۆرانی بنووسە.",
    "🎞 کیفیت رو انتخاب کن:": "🎞 کوالیتی هەڵبژێرە:",
    "❌لینک معتبر پیدا نشد.\nاز دکمه\u200cها استفاده کن.": "❌ بەستەری دروست نەدۆزرایەوە.\nلە دوگمەکان بەکاربهێنە.",
    "▶️ یوتیوب شناسایی شد\nکیفیت رو انتخاب کن:": "▶️ یوتیوب ناسێنرا\nکوالیتی هەڵبژێرە:",
    "⏳پردازش...": "⏳ لە پرۆسەکردندایە...",
    "🗜 در حال فشرده\u200cسازی...": "🗜 لە پەستاندندایە...",
    "🎧 در حال شناسایی آهنگ...": "🎧 گۆرانی دەناسرێتەوە...",
    "🎧 در حال تبدیل ویس به متن...": "🎧 دەنگ دەگۆڕدرێت بۆ دەق...",
    "🌤 در حال دریافت آب\u200cوهوا...": "🌤 کەش و هەوا وەردەگیرێت...",
    "💱 در حال تبدیل...": "💱 لە گۆڕینەوەدایە...",
    "🎨 در حال ساخت استیکر...": "🎨 ستیکەر دروست دەکرێت...",
    "📊 در حال ساخت نمودار دما...": "📊 هێڵکاریی گەرما دروست دەکرێت...",
    "📠در حال خلاصه\u200cسازی...": "📠 لە کورتەکردنەوەدایە...",
    "📠در حال ساخت فایل ورد...": "📠 فایلی Word دروست دەکرێت...",
    "📄 ساخت PDF...": "📄 PDF دروست دەکرێت...",
    "📠جستجو...": "📠 گەڕان...",
    "🔗 در حال کوتاه کردن...": "🔗 بەستەرەکە کورت دەکرێتەوە...",
    "📜 دریافت زیرنویس...": "📜 ژێرنووس وەردەگیرێت...",
    "❌ لینک معتبر بفرست.": "❌ بەستەرێکی دروست بنێرە.",
    "❌ لینک قبلی نیست.": "❌ هیچ بەستەرێکی پێشوو نییە.",
    "فایل پیدا نشد": "فایل نەدۆزرایەوە",
    "✅ با موفقیت ارسال شد.": "✅ بە سەرکەوتوویی نێردرا.",
    "ناموفق": "سەرکەوتوو نەبوو",
    "✅ فشرده شد.": "✅ قەبارەکەی کەم کرایەوە.",
    "❌ساخت نمودار ناموفق بود.": "❌ دروستکردنی هێڵکاری سەرکەوتوو نەبوو.",
    "❌متن بفرست.": "❌ دەق بنێرە.",
    "❌متن خیلی کوتاهه.": "❌ دەقەکە زۆر کورتە.",
    "📠اسم آهنگ و خواننده رو بفرست\nمثال: Shape of You - Ed Sheeran": "📠 ناوی گۆرانی و گۆرانیبێژ بنێرە\nنموونە: Shape of You - Ed Sheeran",
    "😕 متن پیدا نشد.": "😕 دەق نەدۆزرایەوە.",
    "😕 شناسایی نشد.": "😕 نەناسێنرا.",
    "❌خطا در شناسایی آهنگ.": "❌ هەڵە لە ناسینەوەی گۆرانی.",
    "❌تبدیل نشد.": "❌ نەگۆڕدرا.",
    "❌عکسی نیست.": "❌ هیچ وێنەیەک نییە.",
    "❌ساخت استیکر ناموفق بود.": "❌ دروستکردنی ستیکەر سەرکەوتوو نەبوو.",
    "😕 زیرنویس پیدا نشد.": "😕 ژێرنووس نەدۆزرایەوە.",
    "❌ خطا در ساخت فایل ورد.": "❌ هەڵە لە دروستکردنی فایلی Word.",
    "❌ساخت نمودار ناموفق بود.": "❌ دروستکردنی هێڵکاری سەرکەوتوو نەبوو.",
    "❌ فرمت نادرست. مثال: BTC above 100000": "❌ فۆرمات هەڵەیە. نموونە: BTC above 100000",
    "❌این نماد پشتیبانی نمی\u200cشود. دلار، طلا یا سکه را امتحان کن.": "❌ ئەم هێمایە پشتگیری ناکرێت. دۆلار، زێڕ یان دراو تاقی بکەرەوە.",
    "❌نامعتبر": "❌ دروست نییە",
    "❌خطا": "❌ هەڵە",
    "❌ لینک منقضی شده.": "❌ بەستەرەکە بەسەرچووە.",
    "❌اطلاعات منقضی شده. دوباره ویدیو بفرست.": "❌ زانیارییەکان بەسەرچوون. دووبارە ڤیدیۆکە بنێرە.",
    "⛔️ بلاک هستید.": "⛔️ بلۆک کراویت.",
    "⛔️ دسترسی ندارید.": "⛔️ دەستگەیشتنت نییە.",
    "👋 خداحافظ!": "👋 خواحافیز!",
    "👋 سلام! من <b>آسو</b> هستم 😎\n\nدستیار باحال AsoLand. هر چی بخوای بپرس؛ از قیمت دلار تا حل مسئله ریاضی، از شوخی تا ترجمه.\n\nبزن بریم 💬": "👋 سڵاو! من <b>ئاسو</b>م 😎\n\nیاریدەدەری خۆشی AsoLand. هەرچییەکت دەوێت بپرسە؛ لە نرخی دۆلارەوە تا چارەسەرکردنی پرسیاری بیرکاری و وەرگێڕان.\n\nبا دەست پێ بکەین 💬",
    "✅ خارج شدی.": "✅ دەرچوویت.",
    "❌ فایل فعالی وجود ندارد.": "❌ هیچ فایلێکی چالاک نییە.",
    "از دکمه\u200cها استفاده کن:": "لە دوگمەکان بەکاربهێنە:",
    "❌ هیچ کاربری برای ارسال وجود ندارد.": "❌ هیچ بەکارهێنەرێک بۆ ناردن نییە.",
    "❌ محدودیت روزانه تموم شده": "❌ سنووری ڕۆژانە تەواو بووە",
    "❌کانفیگ سالم پیدا نشد.": "❌ کۆنفیگی دروست نەدۆزرایەوە.",
    "✅ آماده شد.": "✅ ئامادە کرا.",
    "😕 پیدا نشد.": "😕 نەدۆزرایەوە.",
    "❌ لینک معتبر بفرست.": "❌ بەستەرێکی دروست بنێرە.",
    "🌠<b>ترجمه:</b>\n\n": "🌠<b>وەرگێڕان:</b>\n\n",
    "😕 چیزی پیدا نشد.": "😕 هیچ شتێک نەدۆزرایەوە.",
    "❌فقط لینک یوتیوب.": "❌ تەنها بەستەری یوتیوب.",
    "🎤 <b>تبدیل ویس به متن</b>\n\nزبان رو انتخاب کن:": "🎤 <b>گۆڕینی دەنگ بۆ دەق</b>\n\nزمان هەڵبژێرە:",
    "🌠متن رو بفرست تا ترجمه کنم.": "🌠 دەق بنێرە تا وەریبگێڕم.",
    "📸 لینک اینستاگرام رو بفرست": "📸 بەستەری ئینستاگرام بنێرە",
    "▶️ لینک یوتیوب رو بفرست": "▶️ بەستەری یوتیوب بنێرە",
    "🎵 لینک تیک\u200cتاک رو بفرست": "🎵 بەستەری تیکتۆک بنێرە",
    "🦠لینک توییتر رو بفرست": "🦠 بەستەری تویتەر بنێرە",
    "🔴 لینک ردیت رو بفرست": "🔴 بەستەری ڕێدیت بنێرە",
    "🔄 ری\u200cاستارت": "🔄 دووبارە دەستپێکردن",
    "🌐 زمان / Language": "🌐 زمان",
    "🔍 جستجو": "🔍 گەڕان",
    "متن آهنگ": "دەقی گۆرانی",
    "زیرنویس": "ژێرنووس",
    "تمام": "تەواو",
    "تموم": "تەواو",
    "آب\u200cوهوای": "کەش و هەوای",
    "آب و هوای": "کەش و هەوای",
    "قیمت دلار و سکه": "نرخی دۆلار و دراو",
    "قیمت دلار": "نرخی دۆلار",
    "قیمت ارز": "نرخی دراو",
    "قیمت طلا": "نرخی زێڕ",
    "قیمت سکه": "نرخی دراو",
    "در حال دریافت": "لە وەرگرتندایە",
    "در حال پردازش": "لە پرۆسەکردندایە",
    "در حال ارسال": "لە ناردندایە",
    "چند لحظه بعد دوباره امتحان کن": "چەند ساتێک دواتر دووبارە هەوڵ بدە",
    "چند لحظه بعد دوباره تلاش کن": "چەند ساتێک دواتر دووبارە هەوڵ بدە",
    "چند دقیقه بعد دوباره امتحان کن": "چەند خولەکێک دواتر دووبارە هەوڵ بدە",
    "پیدا نشد": "نەدۆزرایەوە",
    "ارسال شد": "نێردرا",
    "آماده شد": "ئامادە کرا",
    "محدودیت روزانه": "سنووری ڕۆژانە",
    "باقی مانده": "ماوە",
    "باقی\u200cمانده": "ماوە",
    "کل کاربران": "کۆی بەکارهێنەران",
    "کل دانلودها": "کۆی داگرتنەکان",
    "دانلود امروز": "داگرتنی ئەمڕۆ",
    "کانفیگ امروز": "کۆنفیگی ئەمڕۆ",
    "⚠️ ویدیوی زیرنویس‌دار در دسترس نیست.\nاز فایل SRT استفاده کنید.": "⚠️ ڤیدیۆی ژێرنووسکراو بەردەست نییە.\nفایلی SRT بەکاربهێنە.",
    "❌ اول یک فایل PDF، DOCX یا TXT ارسال کن.": "❌ سەرەتا فایلێکی PDF، DOCX یان TXT بنێرە.",
    "❌ فعلاً قیمت در دسترس نیست.\nاتصال اینترنت را چک کن و دوباره امتحان کن.": "❌ ئێستا نرخ بەردەست نییە.\nپەیوەندی ئینتەرنێت بپشکنە و دووبارە هەوڵ بدە.",
    "❌ پردازش هوشمند فایل با خطا مواجه شد. دوباره تلاش کن.": "❌ پرۆسەکردنی زیرەکی فایل بە هەڵە ڕووبەڕوو بوو. دووبارە هەوڵ بدە.",
    "❌خطا در ساخت فایل ورد.": "❌ هەڵە لە دروستکردنی فایلی Word.",
    "❌ این فایل پشتیبانی نمی‌شود.\n\n📄 فرمت‌های مجاز: PDF، DOCX، TXT": "❌ ئەم فایلە پشتگیری ناکرێت.\n\n📄 فۆرماتە ڕێگەپێدراوەکان: PDF، DOCX، TXT",
    "🖼 <b>ساخت استیکر با عکس</b>\n\nعکس موردنظرت را بفرست.\nبعد فونت را انتخاب می‌کنی و متن دلخواهت روی عکس نوشته می‌شود.": "🖼 <b>دروستکردنی ستیکەر لەگەڵ وێنە</b>\n\nوێنەکەت بنێرە.\nدواتر فۆنت هەڵبژێرە و دەقی دڵخوازت لەسەر وێنەکە دەنووسرێت.",
    "📸 عکس سؤال را بفرست. متن سؤال را تشخیص می‌دهم و مرحله‌به‌مرحله حل می‌کنم.": "📸 وێنەی پرسیارەکە بنێرە. دەقی پرسیارەکە دەناسینمەوە و هەنگاو بە هەنگاو چارەسەری دەکەم.",
    "📁 <b>فایل‌خوان هوشمند AsoLand</b>\n\nفایل PDF، Word (DOCX) یا TXT را همین‌جا ارسال کن.\n\nبعد از دریافت فایل می‌توانی خلاصه، جستجو، سؤال، ترجمه، نکات، آزمون یا فلش‌کارت بگیری.": "📁 <b>خوێنەری زیرەکی فایل AsoLand</b>\n\nفایلی PDF، Word (DOCX) یان TXT لێرە بنێرە.\n\nدوای وەرگرتنی فایل دەتوانیت کورتە، گەڕان، پرسیار، وەرگێڕان، خاڵەکان، تاقیکردنەوە یان فلەش‌کارت وەربگریت.",
    "❌نتونستم این عبارت رو حل کنم. مثال: (25+7)*3، sqrt(144)، x^2-5*x+6=0": "❌ نەمتوانی ئەم دەربڕینە چارەسەر بکەم. نموونە: (25+7)*3، sqrt(144)، x^2-5*x+6=0",
    "❌موتور محاسباتی نصب نیست. این دستور را اجرا کن: pip install sympy": "❌ بزوێنەری حیسابکردن دامەزراو نییە. ئەم فەرمانە جێبەجێ بکە: pip install sympy",
    "❌ این فایل پشتیبانی نمی‌شود.": "❌ ئەم فایلە پشتگیری ناکرێت.",
    "فرمت فایل پشتیبانی نمی‌شود. فقط PDF، DOCX و TXT مجاز است.": "فۆرماتی فایل پشتگیری ناکرێت. تەنها PDF، DOCX و TXT ڕێگەپێدراون.",
    "فرمت فایل پشتیبانی نمی‌شود. فقط PDF، DOCX و TXT مجاز است.": "فۆرماتی فایل پشتگیری ناکرێت. تەنها PDF، DOCX و TXT ڕێگەپێدراون.",
    "فایل زیر را منظم و دقیق خلاصه کن. عنوان‌های اصلی، ایده‌های مهم و نتیجه‌گیری را حفظ کن.": "فایلی خوارەوە بە ڕێک و پێک و ورد کورتە بکە. سەردێڕە سەرەکییەکان، بیرۆکە گرنگەکان و ئەنجامەکە بپارێزە.",
    "این سؤال را از روی تصویر بخوان و حل کن.": "ئەم پرسیارە لەسەر وێنەکە بخوێنەوە و چارەسەری بکە.",
    "متن داخل این تصویر را کامل استخراج کن.": "هەموو دەقی ناو ئەم وێنەیە بە تەواوی دەربهێنە.",
    "❌ از این فایل متن قابل خواندن پیدا نشد.": "❌ هیچ دەقێکی خوێندنەوە لەم فایلەدا نەدۆزرایەوە.",
    "❌ ساخت نمودار ناموفق بود.": "❌ دروستکردنی هێڵکاری سەرکەوتوو نەبوو.",
    "❌ نمودار آب‌وهوا ساخته نشد.": "❌ هێڵکاریی کەش و هەوا دروست نەکرا.",
    "❌ فرمت نادرست. مثال: BTC above 100000": "❌ فۆرمات هەڵەیە. نموونە: BTC above 100000",
    "❌این نماد پشتیبانی نمی‌شود. دلار، طلا یا سکه را امتحان کن.": "❌ ئەم هێمایە پشتگیری ناکرێت. دۆلار، زێڕ یان دراو تاقی بکەرەوە.",
    "✅ استیکر آماده شد.\nمی‌تونی متن جدید بفرستی یا تنظیمات را عوض کنی.": "✅ ستیکەر ئامادەیە.\nدەتوانیت دەقی نوێ بنێریت یان ڕێکخستنەکان بگۆڕیت.",
    "✅ <b>استیکر آماده شد.</b>\n\nمتن جدید بفرست تا با همین عکس و فونت، استیکر دیگری ساخته شود.": "✅ <b>ستیکەر ئامادەیە.</b>\n\nدەقی نوێ بنێرە تا بە هەمان وێنە و فۆنت ستیکەرێکی تر دروست بکرێت.",
    "لغو": "هەڵوەشاندنەوە",
    "سلام <b>": "سڵاو <b>",
    "منوی اصلی:": "پێڕستی سەرەکی:",
    "ربات همه\u200cکاره و هوشمند": "بۆتی زیرەکی هەمووکارە",
    "از منوی زیر استفاده کن.": "لە پێڕستی خوارەوە هەڵبژێرە.",
    "🌤 نام شهر را بفرست؛ نمودار دمای ۷ روزه می\u200cسازم.": "🌤 ناوی شار بنێرە؛ هێڵکاریی گەرمای ٧ ڕۆژ دروست دەکەم.",
    "🌤 <b>آب\u200cوهوا</b>\n\nنام شهر را بفرست؛ مثال: تهران، آمستردام، اربیل": "🌤 <b>کەش و هەوا</b>\n\nناوی شار بنێرە؛ نموونە: تاران، ئەمستەردام، هەولێر",
    "🧮 <b>محاسبه\u200cگر پیشرفته</b>\n\nعبارت ریاضی، تابع یا معادله را بفرست.": "🧮 <b>حیسابکەری پێشکەوتوو</b>\n\nدەربڕینی بیرکاری، فەنکشن یان هاوکێشە بنێرە.",
    "💱 <b>تبدیل ارز</b>\n\nمقدار را بفرست؛ مثال:": "💱 <b>گۆڕینی دراو</b>\n\nبڕەکە بنێرە؛ نموونە:",
    "🔔 <b>یادآور</b>\n\nمثال:": "🔔 <b>بیرخەرەوە</b>\n\nنموونە:",
    "🩷 <b>کانفیگ رایگان</b>\n\nنوع کانفیگ رو انتخاب کن:": "🩷 <b>کۆنفیگی بەخۆڕایی</b>\n\nجۆری کۆنفیگ هەڵبژێرە:",
    "📠<b>خلاصه\u200cسازی متن</b>\n\nمتن یا مقاله رو بفرست.": "📠<b>کورتەکردنەوەی دەق</b>\n\nدەق یان وتارەکە بنێرە.",
    "📠متن رو بفرست تا به فایل ورد تبدیل کنم.": "📠 دەقەکە بنێرە تا بیکەم بە فایلی Word.",
    "✅ <b>زیباسازی متن</b>\n\nمتن مورد نظرت رو بفرست.": "✅ <b>جوانکردنی دەق</b>\n\nدەقە دڵخوازەکەت بنێرە.",
    "🎬 <b>ساخت زیرنویس</b>\n\nزبان زیرنویس مورد نظرت رو انتخاب کن:": "🎬 <b>دروستکردنی ژێرنووس</b>\n\nزمانی ژێرنووس هەڵبژێرە:",
    "🎓 <b>ابزارهای دانشجویی</b>\n\nسؤال ریاضی، فیزیک، شیمی یا مسئله درسی را بفرست.": "🎓 <b>ئامرازە خوێندکارییەکان</b>\n\nپرسیاری بیرکاری، فیزیا، کیمیا یان پرسیارێکی خوێندکاری بنێرە.",
    "📁 <b>فایل\u200cخوان هوشمند AsoLand</b>\n\nفایل PDF، Word (DOCX) یا TXT را همین\u200cجا ارسال کن.": "📁 <b>خوێنەری زیرەکی فایل AsoLand</b>\n\nفایلی PDF، Word (DOCX) یان TXT لێرە بنێرە.",
    "⏳در حال آماده\u200cسازی...": "⏳ لە ئامادەکردندایە...",
    "⏳در حال آماده\u200cسازی MP3...": "⏳ لە ئامادەکردنی MP3دایە...",
    "📠در حال دریافت متن...": "📠 دەق وەردەگیرێت...",
    "🎬 شروع ساخت زیرنویس ": "🎬 دەستپێکردنی دروستکردنی ژێرنووس ",
    "❌لینک معتبر پیدا نشد.\nاز دکمه\u200cها استفاده کن.": "❌ بەستەری دروست نەدۆزرایەوە.\nلە دوگمەکان بەکاربهێنە.",
    "خطا در باز کردن استیکر متنی": "هەڵە لە کردنەوەی ستیکەری دەقی",
    "رنگ نامعتبر": "ڕەنگ دروست نییە",
    "فونت نامعتبر": "فۆنت دروست نییە",
}

# Common single-word UI vocabulary used in dynamic f-strings.
CKB_WORD_MAP = {
    "اول":"سەرەتا","حالا":"ئێستا","الان":"ئێستا","امروز":"ئەمڕۆ","فردا":"بەیانی",
    "لطفاً":"تکایە","لطفا":"تکایە","بفرست":"بنێرە","بفرستید":"بنێرن","انتخاب":"هەڵبژاردن",
    "انتخابی":"هەڵبژێردراو","انتخاب کن":"هەڵبژێرە","انتخاب کنید":"هەڵبژێرە","زبان":"زمان",
    "فعلی":"ئێستا","فعلاً":"ئێستا","فعلا":"ئێستا","تغییر":"گۆڕین","جدید":"نوێ",
    "دوباره":"دووبارە","دوبارە":"دووبارە","برگشت":"گەڕانەوە","خروج":"دەرچوون","ورود":"چوونەژوورەوە",
    "منو":"پێڕست","منوی":"پێڕستی","اصلی":"سەرەکی","راهنما":"ڕێنمایی","پشتیبانی":"پشتگیری",
    "کاربر":"بەکارهێنەر","کاربران":"بەکارهێنەران","نام":"ناو","آیدی":"ناسنامە","تاریخ":"بەروار",
    "امکانات":"تایبەتمەندییەکان","هوشمند":"زیرەک","هوش":"زیرەکی","مصنوعی":"دەستکرد","ابزار":"ئامراز",
    "ابزارها":"ئامرازەکان","کاربردی":"بەسوود","دانلود":"داگرتن","دریافت":"وەرگرتن","ارسال":"ناردن",
    "ارسال شد":"نێردرا","آماده":"ئامادە","آماده شد":"ئامادە کرا","پردازش":"پرۆسەکردن","پردازش...":"پرۆسەکردن...",
    "در حال":"لە پرۆسەکردندایە","خطا":"هەڵە","ناموفق":"سەرکەوتوو نەبوو","موفق":"سەرکەوتوو",
    "موفقیت":"سەرکەوتن","نتیجه":"ئەنجام","پاسخ":"وەڵام","سؤال":"پرسیار","سوال":"پرسیار",
    "پرسش":"پرسیار","متن":"دەق","عکس":"وێنە","تصویر":"وێنە","فایل":"فایل","ویدیو":"ڤیدیۆ",
    "آهنگ":"گۆرانی","خواننده":"گۆرانیبێژ","شهر":"شار","کشور":"وڵات","قیمت":"نرخ","نرخ":"نرخ",
    "دلار":"دۆلار","یورو":"یۆرۆ","پوند":"پاوەند","درهم":"دیرھەم","لیر":"لیرە","یوان":"یوان",
    "طلا":"زێڕ","طلای":"زێڕی","سکه":"دراو","جدید":"نوێ","قدیم":"کۆن","نیم":"نیوە","ربع":"چارەک",
    "گرمی":"گرام","گرم":"گرام","مثقال":"موسقاڵ","کیفیت":"کوالیتی","رنگ":"ڕەنگ","پس‌زمینه":"پاشبنەما",
    "پس زمینه":"پاشبنەما","فونت":"فۆنت","زیبا":"جوان","زیباسازی":"جوانکردن","ساخت":"دروستکردن",
    "ساختن":"دروستکردن","استیکر":"ستیکەر","عکس جدید":"وێنەی نوێ","رنگ متن":"ڕەنگی دەق","رنگ پس‌زمینه":"ڕەنگی پاشبنەما",
    "جستجو":"گەڕان","یافت":"دۆزرایەوە","یافت نشد":"نەدۆزرایەوە","پیدا نشد":"نەدۆزرایەوە","اطلاعات":"زانیاری",
    "محدودیت":"سنوور","روزانه":"ڕۆژانە","روز":"ڕۆژ","باقی‌مانده":"ماوە","باقی مانده":"ماوە",
    "کل":"کۆی","تعداد":"ژمارە","آمار":"ئامار","پیام":"پەیام","همگانی":"گشتی","بلاک‌لیست":"لیستی ڕێگریکراو",
    "یادآور":"بیرخەرەوە","یادآوری":"بیرخستنەوە","هدف":"ئامانج","زمان":"زمان","فارسی":"فارسی","کوردی":"کوردی","انگلیسی":"ئینگلیزی",
    "ترجمه":"وەرگێڕان","خلاصه":"کورتە","خلاصه‌سازی":"کورتەکردنەوە","خلاصه سازی":"کورتەکردنەوە","معلم":"مامۆستا",
    "انگلیسی":"ئینگلیزی","لغت":"وشە","گرامر":"ڕێزمان","یادگیری":"فێربوون","تمرین":"ڕاهێنان","سطح":"ئاست",
    "مبتدی":"سەرەتایی","متوسط":"مامناوەند","پیشرفته":"پێشکەوتوو","آموزش":"پەروەردە","دانشجویی":"خوێندکاری",
    "دانشجو":"خوێندکار","حل":"چارەسەرکردن","معادله":"هاوکێشە","محاسبه":"حیسابکردن","ماشین":"ئامێر",
    "هوا":"هەوا","آب‌وهوا":"کەش و هەوا","آب و هوا":"کەش و هەوا","باران":"باران","برف":"بەفر","مه":"تەم",
    "آسمان صاف":"ئاسمانی ڕوون","ابری":"هەوراو","نیمه‌ابری":"تاڕادەیەک هەوراو","رعدوبرق":"تۆفان و برووسکە",
    "دما":"گەرما","رطوبت":"شێ","باد":"با","احتمال":"ئەگەری","بارش":"باران","امروز":"ئەمڕۆ","کمینه":"کەمترین","بیشینه":"زۆرترین",
    "اخبار":"هەواڵ","خبر":"هەواڵ","فناوری":"تەکنەلۆجیا","اقتصاد":"ئابووری","کریپتو":"کریپتۆ","عمومی":"گشتی",
    "نمودار":"هێڵکاری","چارت":"هێڵکاری","هشدار":"ئاگادارکردنەوە","هشدار قیمت":"ئاگادارکردنەوەی نرخ",
    "فال":"بەخت","روزانه":"ڕۆژانە","آرامش":"ئارامی","فرصت":"دەرفەت","تازه":"نوێ","خبر خوب":"هەواڵی باش",
    "تصمیم":"بڕیار","مهم":"گرنگ","پیشرفت":"پێشکەوتن","دیدار":"سەردان","خوشایند":"خۆش","رنگ":"ڕەنگ",
    "آبی":"شین","قرمز":"سور","سبز":"سەوز","بنفش":"مۆر","مشکی":"ڕەش","نارنجی":"پرتەقاڵی","صورتی":"پەمەیی","سفید":"سپی","زرد":"زەرد",
    "پررنگ":"قەڵەو","ایتالیک":"لار","اصلی":"سەرەکی","با فاصله":"بە بۆشایی","نقطه‌دار":"خاڵدار",
    "تمام‌عرض":"پانی تەواو","فشرده":"پەستاندراو","فشرده‌سازی":"پەستاندن","کوتاه":"کورت","کوتاه کردن":"کورتکردنەوە",
    "لینک":"بەستەر","معتبر":"دروست","منقضی":"بەسەرچوو","حجم":"قەبارە","بیش از حد":"لە سنوور زیاتر",
    "مثال":"نموونە","نمونه":"نموونە","فقط":"تەنها","دکمه":"دوگمە","دکمه‌ها":"دوگمەکان","استفاده":"بەکاربردن",
    "منبع":"سەرچاوە","کانال":"کەناڵ","تلگرام":"تەلەگرام","ساعت":"کات","تقویم":"ڕۆژژمێر","مناسبت":"بۆنە",
    "امروز چندمه":"ئەمڕۆ چەندەمە","صبح":"بەیانی","شب":"شەو","دقیقه":"خولەک","ثانیه":"چرکە","روز":"ڕۆژ",
    "فعال":"چالاک","فعالیت":"چالاکی","حذف":"سڕینەوە","پاک":"سڕینەوە","تلاش":"هەوڵ","تلاش دوباره":"دووبارە هەوڵدان",
    "اول":"سەرەتا","حالا":"ئێستا","الان":"ئێستا","بعد":"دواتر","قبل":"پێش","وقتی":"کاتێک","اگر":"ئەگەر",
    "یا":"یان","برای":"بۆ","با":"لەگەڵ","بدون":"بەبێ","روی":"لەسەر","داخل":"ناو","از":"لە","به":"بۆ","را":"ەکە","را انتخاب":"هەڵبژێرە",
}

def _ckb_replace_words(text: str) -> str:
    # Protect HTML tags, URLs, code blocks and latin-only tokens.
    protected=[]
    def protect(m):
        protected.append(m.group(0)); return f"\x00P{len(protected)-1}\x00"
    text = re.sub(r"https?://[^\s]+|<[^>]+>|`[^`]*`|@[A-Za-z0-9_]+", protect, text)
    # Long phrases first; then word-level vocabulary.
    for src, dst in sorted(CKB_WORD_MAP.items(), key=lambda kv: len(kv[0]), reverse=True):
        if not re.search(r'[\u0600-\u06ff]', src):
            continue
        if re.search(r'\s', src):
            pattern = r'(?<![\u0600-\u06ff])' + re.escape(src) + r'(?![\u0600-\u06ff])'
        else:
            pattern = r'(?<![\u0600-\u06ff])' + re.escape(src) + r'(?![\u0600-\u06ff])'
        text = re.sub(pattern, dst, text)
    for i, value in enumerate(protected):
        text = text.replace(f"\x00P{i}\x00", value)
    return text

def translate_ckb_ui(text: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    # Exact full-message translations first.
    if text in CKB_UI_EXTRA:
        return CKB_UI_EXTRA[text]
    result = text
    # Apply exact/embedded UI phrases first. If a phrase was translated, do not
    # run word-level replacement over the Sorani result (that would corrupt
    # already-correct Kurdish words).
    for src, dst in sorted(CKB_UI_EXTRA.items(), key=lambda kv: len(kv[0]), reverse=True):
        result = result.replace(src, dst)
    return _ckb_replace_words(result)

def tr(text, lang="fa"):
    if not isinstance(text, str) or lang == "fa":
        return text
    item = UI_TRANSLATIONS.get(text)
    if item:
        return item[0] if lang == "en" else item[1]
    # Sorani must be handled before the generic Persian phrase table. The generic
    # table contains word-level replacements that can destroy an exact sentence.
    if lang == "ckb":
        return translate_ckb_ui(text)
    result = text
    for src, pair in UI_PHRASES.items():
        dst = pair[0] if lang == "en" else pair[1]
        result = result.replace(src, dst)
    return result

def LButton(text, *args, **kwargs):
    # Button callbacks do not carry user context, so the label is localized by a
    # per-message language marker when available. get_*_keyboard sets it via context.
    lang = kwargs.pop("_lang", _CURRENT_UI_LANG.get())
    return InlineKeyboardButton(tr(text, lang), *args, **kwargs)

def LKeyboardButton(text, *args, **kwargs):
    lang = kwargs.pop("_lang", _CURRENT_UI_LANG.get())
    return KeyboardButton(tr(text, lang), *args, **kwargs)

def _lang_from_context(context):
    try:
        return get_user_lang(context._chat_id, context)
    except Exception:
        return "fa"

def language_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🇮🇷 فارسی", callback_data="set_lang_fa"), InlineKeyboardButton("☀️ کوردی", callback_data="set_lang_ckb")],
        [InlineKeyboardButton("🇬🇧 English", callback_data="set_lang_en")],
        [InlineKeyboardButton("🏠 منوی اصلی", callback_data="main_menu")],
    ])

async def reply_text_localized(message, context, text, *args, **kwargs):
    return await message.reply_text(tr(text, get_user_lang(message.chat_id, context)), *args, **kwargs)

async def edit_text_localized(message, context, text, *args, **kwargs):
    return await message.edit_text(tr(text, get_user_lang(message.chat_id, context)), *args, **kwargs)

async def edit_query_localized(query, context, text, *args, **kwargs):
    return await query.edit_message_text(tr(text, get_user_lang(query.message.chat_id, context)), *args, **kwargs)

async def bot_send_message_localized(bot, user_id, text, *args, **kwargs):
    return await bot.send_message(user_id, tr(text, get_user_lang(user_id)), *args, **kwargs)

# =========================================================
# HELPERS
# =========================================================

def get_jalali_datetime(lang=None):
    """Localized Jalali date/time. Uses the current user's UI language by default."""
    lang = lang or _CURRENT_UI_LANG.get()
    now = jdatetime.datetime.now()
    weekdays = {
        "fa": {0: "شنبه", 1: "یکشنبه", 2: "دوشنبه", 3: "سه‌شنبه", 4: "چهارشنبه", 5: "پنج‌شنبه", 6: "جمعه"},
        "ckb": {0: "شەممە", 1: "یەکشەممە", 2: "دووشەممە", 3: "سێشەممە", 4: "چوارشەممە", 5: "پێنجشەممە", 6: "هەینی"},
        "en": {0: "Saturday", 1: "Sunday", 2: "Monday", 3: "Tuesday", 4: "Wednesday", 5: "Thursday", 6: "Friday"},
    }
    months = {
        "fa": {1: "فروردین", 2: "اردیبهشت", 3: "خرداد", 4: "تیر", 5: "مرداد", 6: "شهریور", 7: "مهر", 8: "آبان", 9: "آذر", 10: "دی", 11: "بهمن", 12: "اسفند"},
        "ckb": {1: "فەروردین", 2: "ئاردیبهشت", 3: "خورداد", 4: "تیر", 5: "گەلاوێژ", 6: "خەرمانان", 7: "ڕەزبەر", 8: "گەڵاڕێزان", 9: "سەرماوەرز", 10: "بەفرانبار", 11: "ڕێبەندان", 12: "ڕەشەمە"},
        "en": {1: "Farvardin", 2: "Ordibehesht", 3: "Khordad", 4: "Tir", 5: "Mordad", 6: "Shahrivar", 7: "Mehr", 8: "Aban", 9: "Azar", 10: "Dey", 11: "Bahman", 12: "Esfand"},
    }
    if lang == "en":
        return f"📅 {weekdays[lang][now.weekday()]} {now.day} {months[lang][now.month]} {now.year}\n🕠 {now.hour:02d}:{now.minute:02d}:{now.second:02d}"
    return f"📅 {weekdays.get(lang, weekdays['fa'])[now.weekday()]} {now.day} {months.get(lang, months['fa'])[now.month]} {now.year}\n🕠 {now.hour:02d}:{now.minute:02d}:{now.second:02d}"


def extract_url(text):
    pattern = r"https?://(?:www\.)?(?:instagram\.com|youtube\.com|youtu\.be|tiktok\.com|vm\.tiktok\.com|twitter\.com|x\.com|facebook\.com|fb\.watch|reddit\.com|redd\.it|aparat\.com)/[^\s]+"
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0).rstrip(").,!?]>}") if match else None


def normalize_media_url(url: str) -> str:
    """Strip tracking query params that break extractors (esp. Instagram share links)."""
    if not url:
        return url
    from urllib.parse import urlparse, urlunparse
    parsed = urlparse(url.strip())
    host = (parsed.netloc or "").lower()
    path = parsed.path or ""
    # Instagram: keep only /reel/ID or /p/ID or /tv/ID — drop igsh, igsi, etc.
    if "instagram.com" in host:
        m = re.search(r"/(reel|p|tv)/([A-Za-z0-9_-]+)", path)
        if m:
            clean_path = f"/{m.group(1)}/{m.group(2)}/"
            return urlunparse(("https", "www.instagram.com", clean_path, "", "", ""))
        # fallback: drop query/fragment
        return urlunparse(("https", "www.instagram.com", path.rstrip("/") + "/", "", "", ""))
    # TikTok share params often unnecessary
    if "tiktok.com" in host or "vm.tiktok.com" in host:
        return urlunparse((parsed.scheme or "https", parsed.netloc, path, "", "", ""))
    return url


def get_platform(url: str) -> str:
    url = url.lower()
    if "instagram.com" in url: return "Instagram"
    if "youtube.com" in url or "youtu.be" in url: return "YouTube"
    if "tiktok.com" in url or "vm.tiktok.com" in url: return "TikTok"
    if "twitter.com" in url or "x.com" in url: return "Twitter / X"
    if "facebook.com" in url or "fb.watch" in url: return "Facebook"
    if "reddit.com" in url or "redd.it" in url: return "Reddit"
    if "aparat.com" in url: return "Aparat"
    return "Unknown"


def format_size(size):
    if not size: return "نامشخص"
    size = float(size)
    if size < 1024: return f"{size:.0f} B"
    if size < 1024 ** 2: return f"{size / 1024:.2f} KB"
    if size < 1024 ** 3: return f"{size / 1024 ** 2:.2f} MB"
    return f"{size / 1024 ** 3:.2f} GB"


def format_duration(seconds):
    if not seconds: return "نامشخص"
    seconds = int(seconds)
    h, m, s = seconds // 3600, (seconds % 3600) // 60, seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def clean_title(title):
    if not title: return "Video"
    return re.sub(r'[\\/:*?"<>|]', "_", str(title)).strip()[:80]


def get_media_info(filename):
    try:
        cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration",
               "-show_entries", "stream=width,height", "-of", "json", filename]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        data = json.loads(result.stdout)
        duration = data.get("format", {}).get("duration")
        width = height = None
        for stream in data.get("streams", []):
            if stream.get("codec_type") == "video":
                width, height = stream.get("width"), stream.get("height")
                break
        return {"duration": float(duration) if duration else None, "width": width, "height": height}
    except Exception:
        return {"duration": None, "width": None, "height": None}


def seconds_to_srt_time(seconds: float) -> str:
    millis = int((seconds % 1) * 1000)
    total = int(seconds)
    h, m, s = total // 3600, (total % 3600) // 60, total % 60
    return f"{h:02d}:{m:02d}:{s:02d},{millis:03d}"


# =========================================================
# PRICE SYSTEM (از کانال تلگرام DO_L4)
# =========================================================

DO_L4_URL = "https://t.me/s/DO_L4"

TG_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
    "Accept-Language": "fa-IR,fa;q=0.9,en;q=0.8",
}

async def fetch_channel(url: str) -> str:
    try:
        async with httpx.AsyncClient(
            timeout=20.0,
            follow_redirects=True,
            headers=TG_HEADERS,
            verify=True,
        ) as client:
            resp = await client.get(url)
            if resp.status_code == 200:
                return resp.text
            logger.warning(f"Channel status {resp.status_code}: {url}")
    except Exception as e:
        logger.warning(f"Channel error ({url}): {type(e).__name__}: {e}")
    return ""

def extract_messages(html: str) -> list:
    """Extract message texts from t.me/s page (newest first)."""
    if not html:
        return []
    try:
        soup = BeautifulSoup(html, "html.parser")
        msgs = soup.select(".tgme_widget_message_text")
        return [m.get_text(" ", strip=True) for m in msgs]
    except Exception:
        return []

def clean_num(text) -> int | None:
    if not text:
        return None
    text = str(text).translate(str.maketrans("۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩", "01234567890123456789"))
    text = re.sub(r"[^\d]", "", text)
    if not text.isdigit():
        return None
    val = int(text)
    return val if val > 0 else None

def fmt_price(n):
    return f"{n:,}".replace(",", "٬") if n else "—"

def normalize_text(text: str) -> str:
    if not text:
        return ""
    text = str(text).translate(str.maketrans("۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩", "01234567890123456789"))
    text = text.replace("٬", ",").replace("，", ",").replace("ـ", "")
    text = re.sub(r"[ \t]+", " ", text)
    return text

def parse_do_l4_summary(text: str) -> dict:
    """
    Parse the 'نرخ فروش #دلار، #ارز، #سکه و #طلا' summary message.
    Currency values are already in تومان.
    Coin / gold values marked with هـ.تومان are in هزار تومان → multiply by 1000.
    """
    if not text:
        return {}
    text = normalize_text(text)
    out = {}

    # Currency patterns (already تومان)
    currency_map = {
        "dollar": [r"دلار[:\s]*([\d,]+)\s*تومان", r"💵\s*دلار[:\s]*([\d,]+)"],
        "euro": [r"یورو[:\s]*([\d,]+)\s*تومان", r"💶\s*یورو[:\s]*([\d,]+)"],
        "gbp": [r"پوند[:\s]*([\d,]+)\s*تومان", r"💷\s*پوند[:\s]*([\d,]+)"],
        "aed": [r"درهم[:\s]*([\d,]+)\s*تومان", r"🇦🇪\s*درهم[:\s]*([\d,]+)"],
        "try": [r"لیر[:\s]*([\d,]+)\s*تومان", r"🇹🇷\s*لیر[:\s]*([\d,]+)"],
        "cny": [r"یوان[:\s]*([\d,]+)\s*تومان", r"💴\s*یوان[:\s]*([\d,]+)"],
    }
    for key, pats in currency_map.items():
        for pat in pats:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                val = clean_num(m.group(1))
                if val and val > 100:
                    out[key] = val
                    break

    # Gold & coin — numbers are in هزار تومان
    coin_map = {
        "coin_new": [r"سکه\s*جدید[:\s]*([\d,]+)", r"🌕\s*سکه\s*جدید[:\s]*([\d,]+)"],
        "coin_old": [r"سکه\s*قدیم[:\s]*([\d,]+)", r"🌕\s*سکه\s*قدیم[:\s]*([\d,]+)"],
        "coin_half": [r"سکه\s*نیم[:\s]*([\d,]+)", r"🌕\s*سکه\s*نیم[:\s]*([\d,]+)"],
        "coin_quarter": [r"سکه\s*ربع[:\s]*([\d,]+)", r"🌕\s*سکه\s*ربع[:\s]*([\d,]+)"],
        "coin_gram": [r"سکه\s*گرمی[:\s]*([\d,]+)", r"🌕\s*سکه\s*گرمی[:\s]*([\d,]+)"],
        "mesghal": [r"مثقال\s*طلا[^\d]{0,30}([\d,]+)", r"🌟\s*مثقال[^\d]{0,30}([\d,]+)"],
        "gold18": [r"گرم\s*18\s*عیار[^\d]{0,40}([\d,]+)", r"💫\s*گرم\s*18[^\d]{0,40}([\d,]+)", r"گرم\s*۱۸\s*عیار[^\d]{0,40}([\d,]+)"],
    }
    for key, pats in coin_map.items():
        for pat in pats:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                val = clean_num(m.group(1))
                if val and val > 10:
                    # هـ.تومان → هزار تومان
                    out[key] = val * 1000
                    break

    return out

def parse_do_l4_live_dollar(text: str) -> int | None:
    """Parse live 'دلار فردایی تهران' posts for latest deal price."""
    if not text:
        return None
    text = normalize_text(text)
    # Prefer معامله
    m = re.search(r"دلار\s*فردایی\s*تهران[^\d]{0,40}([\d,]+)[^\d]{0,15}معامله", text, re.IGNORECASE)
    if m:
        val = clean_num(m.group(1))
        if val and 10000 < val < 1000000:
            return val
    m = re.search(r"دلار\s*فردایی\s*تهران[^\d]{0,40}\*{0,2}([\d,]+)\*{0,2}", text, re.IGNORECASE)
    if m:
        val = clean_num(m.group(1))
        if val and 10000 < val < 1000000:
            return val
    return None

async def get_do_l4_prices(lang=None) -> str:
    """Fetch and format prices in the user's selected UI language."""
    lang = lang or _CURRENT_UI_LANG.get() or "fa"
    labels = {
        "fa": dict(title="قیمت دلار و سکه", dollar="دلار تهران", euro="یورو", gbp="پوند", aed="درهم", try_="لیر", cny="یوان",
                   gold18="طلای ۱۸ عیار (هر گرم)", mesghal="مثقال طلا", coin_new="سکه جدید (امامی)", coin_old="سکه قدیم (بهار)",
                   coin_half="نیم سکه", coin_quarter="ربع سکه", coin_gram="سکه گرمی", toman="تومان", source="منبع: کانال تلگرام @DO_L4",
                   unavailable="❌ فعلاً قیمت در دسترس نیست.\nاتصال اینترنت را چک کن و دوباره امتحان کن.",
                   error="❌ خطا در دریافت قیمت.\nچند لحظه بعد دوباره امتحان کن."),
        "ckb": dict(title="نرخی دۆلار و دراو", dollar="دۆلاری تاران", euro="یۆرۆ", gbp="پاوەند", aed="دیرھەم", try_="لیرە", cny="یوان",
                    gold18="زێڕی ١٨ عیار (هەر گرام)", mesghal="موسقاڵی زێڕ", coin_new="دراوی نوێ (ئیمامی)", coin_old="دراوی کۆن (بەهار)",
                    coin_half="نیوە دراو", coin_quarter="چارەکی دراو", coin_gram="دراوی گرامی", toman="تومان", source="سەرچاوە: کەناڵی تەلەگرام @DO_L4",
                    unavailable="❌ ئێستا نرخ بەردەست نییە.\nپەیوەندی ئینتەرنێت بپشکنە و دووبارە هەوڵ بدە.",
                    error="❌ هەڵە لە وەرگرتنی نرخ.\nچەند ساتێک دواتر دووبارە هەوڵ بدە."),
        "en": dict(title="Dollar & Coin Prices", dollar="Tehran Dollar", euro="Euro", gbp="Pound", aed="Dirham", try_="Lira", cny="Yuan",
                  gold18="18K Gold (per gram)", mesghal="Gold Mesghal", coin_new="New Coin (Emami)", coin_old="Old Coin (Bahar)",
                  coin_half="Half Coin", coin_quarter="Quarter Coin", coin_gram="Gram Coin", toman="Toman", source="Source: Telegram @DO_L4",
                  unavailable="❌ Prices are currently unavailable.\nCheck your internet connection and try again.",
                  error="❌ Error while getting prices.\nPlease try again in a moment.")
    }
    L = labels.get(lang, labels["fa"])
    try:
        html_data = await fetch_channel(DO_L4_URL)
        messages = extract_messages(html_data)
        if not messages:
            return L["unavailable"]
        summary, live_dollar = {}, None
        for msg in messages:
            if "نرخ فروش" in msg and ("دلار" in msg or "سکه" in msg):
                parsed = parse_do_l4_summary(msg)
                if parsed:
                    summary = parsed
                    break
        for msg in messages:
            if "دلار فردایی" in msg or "دلار‌فردایی" in msg:
                live_dollar = parse_do_l4_live_dollar(msg)
                if live_dollar:
                    break
        if live_dollar:
            summary["dollar"] = live_dollar
        if not summary:
            return L["unavailable"]
        lines = [f"💰 <b>{L['title']}</b>\n"]
        for key, emoji in [("dollar","🇺🇸"),("euro","🇪🇺"),("gbp","🇬🇧"),("aed","🇦🇪"),("try","🇹🇷"),("cny","🇨🇳")]:
            if summary.get(key):
                lines.append(f"{emoji} {L[key if key != 'try' else 'try_']}: <b>{fmt_price(summary[key])}</b> {L['toman']}")
        lines.append("")
        for key, emoji in [("gold18","🥇"),("mesghal","⚖️"),("coin_new","🪙"),("coin_old","🪙"),("coin_half","🪙"),("coin_quarter","🪙"),("coin_gram","🪙")]:
            if summary.get(key):
                lines.append(f"{emoji} {L[key]}: <b>{fmt_price(summary[key])}</b> {L['toman']}")
        lines.append(f"\n🕔 {get_jalali_datetime(lang)}")
        lines.append(f"📡 {L['source']}")
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"DO_L4 price error: {e}")
        return L["error"]

WEATHER_CODES_FA = {
    0: "آسمان صاف",
    1: "عمدتاً صاف",
    2: "نیمه‌ابری",
    3: "ابری",
    45: "مه",
    48: "مه یخ‌زن",
    51: "نم‌نم باران سبک",
    53: "نم‌نم باران",
    55: "نم‌نم باران شدید",
    61: "باران سبک",
    63: "باران",
    65: "باران شدید",
    71: "برف سبک",
    73: "برف",
    75: "برف شدید",
    80: "رگبار سبک",
    81: "رگبار",
    82: "رگبار شدید",
    95: "رعدوبرق",
    96: "رعدوبرق با تگرگ",
    99: "رعدوبرق و تگرگ شدید",
}

async def get_weather(city: str) -> str:
    city = city.strip()
    if len(city) < 2:
        return "❌نام شهر معتبر نیست."

    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            geo = await client.get(
                "https://geocoding-api.open-meteo.com/v1/search",
                params={
                    "name": city,
                    "count": 1,
                    "language": "fa",
                    "format": "json",
                },
            )
            geo.raise_for_status()
            results = geo.json().get("results") or []
            if not results:
                return f"❌شهر «{html.escape(city)}» پیدا نشد."

            place = results[0]
            lat, lon = place["latitude"], place["longitude"]
            display_name = place.get("name") or city
            country = place.get("country") or ""

            weather = await client.get(
                "https://api.open-meteo.com/v1/forecast",
                params={
                    "latitude": lat,
                    "longitude": lon,
                    "current": ",".join([
                        "temperature_2m",
                        "relative_humidity_2m",
                        "apparent_temperature",
                        "precipitation",
                        "weather_code",
                        "wind_speed_10m",
                        "wind_direction_10m",
                        "cloud_cover",
                    ]),
                    "daily": ",".join([
                        "temperature_2m_max",
                        "temperature_2m_min",
                        "precipitation_probability_max",
                        "sunrise",
                        "sunset",
                    ]),
                    "forecast_days": 3,
                    "timezone": "auto",
                    "wind_speed_unit": "kmh",
                },
            )
            weather.raise_for_status()
            data = weather.json()
            cur = data.get("current", {})
            daily = data.get("daily", {})

            code = cur.get("weather_code")
            condition = WEATHER_CODES_FA.get(code, "وضعیت نامشخص")
            temp = cur.get("temperature_2m")
            feels = cur.get("apparent_temperature")
            humidity = cur.get("relative_humidity_2m")
            wind = cur.get("wind_speed_10m")
            cloud = cur.get("cloud_cover")
            rain = cur.get("precipitation")
            today_max = (daily.get("temperature_2m_max") or [None])[0]
            today_min = (daily.get("temperature_2m_min") or [None])[0]
            rain_prob = (daily.get("precipitation_probability_max") or [None])[0]

            return (
                f"🌤 <b>آب‌وهوای {html.escape(display_name)}</b>"
                f"{f'، {html.escape(country)}' if country else ''}\n\n"
                f"🌡 دما: <b>{temp if temp is not None else '—'}°C</b>\n"
                f"🤗 احساس واقعی: <b>{feels if feels is not None else '—'}°C</b>\n"
                f"☁️ وضعیت: <b>{condition}</b>\n"
                f"💧 رطوبت: <b>{humidity if humidity is not None else '—'}%</b>\n"
                f"💨 باد: <b>{wind if wind is not None else '—'} km/h</b>\n"
                f"☁️ پوشش ابر: <b>{cloud if cloud is not None else '—'}%</b>\n"
                f"🌧 بارش فعلی: <b>{rain if rain is not None else '—'} mm</b>\n"
                f"☔ احتمال بارش امروز: <b>{rain_prob if rain_prob is not None else '—'}%</b>\n"
                f"📈 کمینه/بیشینه امروز: <b>{today_min if today_min is not None else '—'}° / "
                f"{today_max if today_max is not None else '—'}°</b>\n\n"
                "📡 منبع: Open-Meteo"
            )
    except Exception as e:
        logger.error(f"Weather error: {e}")
        return "❌ دریافت آب‌وهوا ناموفق بود. چند لحظه بعد دوباره امتحان کن."


async def get_daily_fortune(user_name: str = "دوست من") -> str:
    """Entertainment-only daily fortune generated deterministically per Persian date/user."""
    seed = f"{user_name}:{date.today().isoformat()}"
    rng = random.Random(seed)
    signs = ["آرامش", "فرصت تازه", "خبر خوب", "تصمیم مهم", "پیشرفت", "دیدار خوشایند"]
    advice = [
        "امروز قبل از تصمیم عجولانه یک بار دیگر همه جوانب را بررسی کن.",
        "یک گفت‌وگویی صادقانه می‌تواند مسیر یک موضوع را بهتر کند.",
        "روی یک کار مهم تمرکز کن و کارهای کم‌اهمیت را عقب بینداز.",
        "اگر فرصتی کوچک دیدی، آن را دست‌کم نگیر؛ شروع‌های کوچک مهم‌اند.",
        "امروز نظم و صبر از سرعت بیشتر به نفع توست.",
        "به حس خوبت توجه کن، اما تصمیم‌های مالی را فقط بر پایه شانس نگیر.",
    ]
    lucky = rng.choice(signs)
    tip = rng.choice(advice)
    return (
        f"🔮 <b>فال روزانه {html.escape(user_name)}</b>\n\n"
        f"✨ حال‌وهوای امروز: <b>{lucky}</b>\n"
        f"💫 پیام امروز: {tip}\n"
        f"🎯 عدد نمادین: <b>{rng.randint(1, 99)}</b>\n"
        f"🎨رنگ نمادین: <b>{rng.choice(['آبی', 'سبز', 'بنفش', 'سفید', 'طلایی'])}</b>\n\n"
        "ℹ️ این بخش صرفاً برای سرگرمی است و پیش‌بینی قطعی آینده نیست."
    )


CALC_ALLOWED_NAMES = {
    "pi": sp.pi if sp else None,
    "e": sp.E if sp else None,
    "sqrt": sp.sqrt if sp else None,
    "sin": sp.sin if sp else None,
    "cos": sp.cos if sp else None,
    "tan": sp.tan if sp else None,
    "asin": sp.asin if sp else None,
    "acos": sp.acos if sp else None,
    "atan": sp.atan if sp else None,
    "log": sp.log if sp else None,
    "ln": sp.log if sp else None,
    "exp": sp.exp if sp else None,
    "abs": sp.Abs if sp else None,
    "factorial": sp.factorial if sp else None,
}

def normalize_math_expression(expr: str) -> str:
    trans = str.maketrans({
        "۰":"0","۱":"1","۲":"2","۳":"3","۴":"4","۵":"5","۶":"6","۷":"7","۸":"8","۹":"9",
        "٠":"0","١":"1","٢":"2","٣":"3","٤":"4","٥":"5","٦":"6","٧":"7","٨":"8","٩":"9",
        "×":"*","÷":"/","−":"-","–":"-","^":"**","π":"pi",
    })
    expr = expr.translate(trans).replace("٫", ".").replace("،", ",")
    expr = expr.replace("√", "sqrt")
    return expr.strip()

def calculate_advanced(expression: str) -> str:
    if not sp:
        return "❌موتور محاسباتی نصب نیست. این دستور را اجرا کن: pip install sympy"

    expression = normalize_math_expression(expression)
    if not expression or len(expression) > 500:
        return "❌عبارت خالی یا بیش از حد طولانی است."

    # Restrict the input language before passing it to SymPy.
    if not re.fullmatch(r"[0-9a-zA-Z_+\-*/%().,=\s]+", expression):
        return "❌عبارت شامل کاراکتر غیرمجاز است."
    names = set(re.findall(r"[A-Za-z_]\w*", expression))
    if any(name not in CALC_ALLOWED_NAMES and name not in {"x", "y", "z"} for name in names):
        return "❌نام تابع/متغیر غیرمجاز است."

    local_dict = dict(CALC_ALLOWED_NAMES)
    local_dict.update({"x": sp.Symbol("x"), "y": sp.Symbol("y"), "z": sp.Symbol("z")})

    try:
        if "=" in expression:
            if expression.count("=") != 1:
                return "❌فقط یک علامت مساوی مجاز است."
            left, right = [part.strip() for part in expression.split("=", 1)]
            lhs = sp.sympify(left, locals=local_dict)
            rhs = sp.sympify(right, locals=local_dict)
            symbols = sorted((lhs - rhs).free_symbols, key=lambda s: s.name)
            if not symbols:
                return "✅ نتیجه: <b>درست</b>" if sp.simplify(lhs-rhs) == 0 else "❌این تساوی نادرست است."
            solutions = sp.solve(sp.Eq(lhs, rhs), symbols)
            if not solutions:
                return "ℹ️ راه‌حلی پیدا نشد."
            return "🧮 <b>حل معادله</b>\n\n" + "\n".join(
                f"<b>{s}</b> = <code>{html.escape(str(sp.N(v, 12)))}</code>"
                for s, v in zip(symbols, solutions[0] if isinstance(solutions[0], (list, tuple)) else solutions)
            )

        result = sp.sympify(expression, locals=local_dict)
        simplified = sp.simplify(result)
        numeric = sp.N(simplified, 14)
        return (
            "🧮 <b>محاسبه</b>\n\n"
            f"عبارت: <code>{html.escape(expression)}</code>\n"
            f"نتیجه دقیق: <b>{html.escape(str(simplified))}</b>\n"
            f"مقدار عددی: <b>{html.escape(str(numeric))}</b>"
        )
    except Exception as e:
        logger.warning(f"Calculator error: {e}")
        return "❌نتونستم این عبارت رو حل کنم. مثال: (25+7)*3، sqrt(144)، x^2-5*x+6=0"




# =========================================================
# EXTENDED TOOLS: SMART ROUTER / CHARTS / ALERTS / NEWS / STUDY
# =========================================================

async def smart_route(text: str, context: ContextTypes.DEFAULT_TYPE):
    """Lightweight Persian intent router for common requests."""
    t = text.lower().strip()
    if any(k in t for k in ["آب و هوا", "آب‌وهوا", "هواشناسی", "weather"]):
        city = re.sub(r".*?(?:آب ?و ?هوا|هواشناسی|weather)\s*(?:ی|در|برای)?\s*", "", t).strip()
        if city and city not in ["امروز", "فردا"]:
            return "weather_direct", city
        return "weather_menu", None
    if any(k in t for k in ["قیمت دلار", "قیمت ارز", "قیمت طلا", "قیمت سکه", "دلار", "سکه"]):
        return "price_do_l4", None
    if any(k in t for k in ["فال امروز", "فال روزانه", "فال حافظ", "فال"]):
        return "daily_fortune", None
    if any(k in t for k in ["خبر", "اخبار", "news"]):
        return "news_menu", None
    if any(k in t for k in ["نمودار", "چارت", "chart"]):
        return "chart_menu", None
    if any(k in t for k in ["هشدار قیمت", "قیمت رسید", "وقتی قیمت"]):
        return "alert_menu", None
    if any(k in t for k in ["حل سوال", "حل سؤال", "سوال ریاضی", "سؤال ریاضی"]):
        return "student_menu", None
    if any(k in t for k in ["تبدیل ارز", "تبدیل دلار", "چند تومان", "به تومان", "به دلار"]):
        return "currency_convert", None
    if any(k in t for k in ["تقویم", "امروز چندمه", "مناسبت"]):
        return "today_calendar", None
    if any(k in t for k in ["یادآوری", "یادآور", "reminder"]):
        return "reminder_menu", None
    if any(k in t for k in ["پنل من", "وضعیت من", "محدودیت من"]):
        return "user_panel", None
    return None, None


async def get_crypto_history_chart(coin_id: str = "bitcoin", days: int = 7):
    if plt is None:
        return None
    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.get(
                f"https://api.coingecko.com/api/v3/coins/{coin_id}/market_chart",
                params={"vs_currency": "usd", "days": days, "interval": "daily" if days > 2 else "hourly"},
                headers={"accept": "application/json", "User-Agent": "AsoLandBot/1.0"}
            )
            r.raise_for_status()
            prices = r.json().get("prices", [])
        if len(prices) < 2:
            return None
        xs=[datetime.fromtimestamp(x/1000) for x,_ in prices]
        ys=[y for _,y in prices]
        fig, ax = plt.subplots(figsize=(9, 4.8), dpi=140)
        ax.plot(xs, ys, linewidth=2)
        ax.set_title(f"{coin_id.title()} - {days} روز")
        ax.set_ylabel("USD")
        ax.grid(alpha=.25)
        fig.autofmt_xdate()
        path=os.path.join(tempfile.gettempdir(), f"aso_{coin_id}_{days}_{int(time.time())}.png")
        fig.tight_layout(); fig.savefig(path, bbox_inches="tight"); plt.close(fig)
        return path
    except Exception as e:
        logger.warning(f"Chart error: {e}")
        return None


async def get_weather_chart(city: str):
    if plt is None:
        return None
    try:
        async with httpx.AsyncClient(timeout=15) as client:
            geo=await client.get("https://geocoding-api.open-meteo.com/v1/search", params={"name":city,"count":1,"language":"fa","format":"json"})
            geo.raise_for_status(); results=geo.json().get("results") or []
            if not results: return None
            p=results[0]
            w=await client.get("https://api.open-meteo.com/v1/forecast", params={"latitude":p["latitude"],"longitude":p["longitude"],"daily":"temperature_2m_max,temperature_2m_min","forecast_days":7,"timezone":"auto"})
            w.raise_for_status(); d=w.json().get("daily",{})
        fig,ax=plt.subplots(figsize=(9,4.8),dpi=140)
        days=d.get("time",[]); hi=d.get("temperature_2m_max",[]); lo=d.get("temperature_2m_min",[])
        ax.plot(days,hi,marker="o",label="بیشینه"); ax.plot(days,lo,marker="o",label="کمینه")
        ax.set_title(f"پیش‌بینی دما - {p.get('name',city)}"); ax.set_ylabel("°C"); ax.grid(alpha=.25); ax.legend(); fig.autofmt_xdate()
        path=os.path.join(tempfile.gettempdir(),f"weather_{int(time.time())}.png")
        fig.tight_layout(); fig.savefig(path,bbox_inches="tight"); plt.close(fig); return path
    except Exception as e:
        logger.warning(f"Weather chart error: {e}"); return None


async def fetch_news(category="general"):
    queries = {
        "general": "ایران OR جهان",
        "tech": "فناوری OR تکنولوژی OR هوش مصنوعی",
        "economy": "اقتصاد OR ارز OR طلا OR بورس",
        "crypto": "بیت کوین OR ارز دیجیتال OR کریپتو",
        "student": "دانشگاه OR آموزش OR کنکور",
    }
    q = queries.get(category, queries["general"])
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
        "Accept": "application/rss+xml, application/xml, text/xml, */*",
        "Accept-Language": "fa-IR,fa;q=0.9,en;q=0.8",
    }
    urls = [
        ("https://news.google.com/rss/search", {"q": q, "hl": "fa", "gl": "IR", "ceid": "IR:fa"}),
        ("https://news.google.com/rss", {"hl": "fa", "gl": "IR", "ceid": "IR:fa"}),
    ]
    items = []
    for url, params in urls:
        try:
            async with httpx.AsyncClient(timeout=20.0, follow_redirects=True, headers=headers) as client:
                r = await client.get(url, params=params)
                r.raise_for_status()
                content = r.text
            # حذف namespace برای پارس ساده‌تر
            content = re.sub(r'\sxmlns[^"]*"[^"]*"', "", content)
            root = ET.fromstring(content)
            for item in root.findall(".//item")[:10]:
                title = (item.findtext("title") or "").strip()
                link = (item.findtext("link") or "").strip()
                pub = (item.findtext("pubDate") or "").strip()
                if title and title not in [t for t, _, _ in items]:
                    items.append((title, link, pub))
            if items:
                break
        except Exception as e:
            logger.warning(f"News fetch error ({url}): {e}")
            continue
    return items[:8]


async def news_text(category="general"):
    items = await fetch_news(category)
    if not items:
        return (
            "❌ فعلاً خبری دریافت نشد.\n"
            "ممکن است موقتاً دسترسی به خبرگزار قطع باشد. چند دقیقه بعد دوباره تلاش کن."
        )
    names={"general":"عمومی","tech":"فناوری","economy":"اقتصاد","crypto":"کریپتو","student":"آموزش"}
    out=[f"📰 <b>آخرین اخبار {names.get(category,'عمومی')}</b>\n"]
    for i,(title,link,pub) in enumerate(items,1):
        safe=html.escape(title[:180]); out.append(f"{i}. <a href=\"{html.escape(link,quote=True)}\">{safe}</a>")
    return "\n".join(out)


async def solve_student_problem(text: str) -> str:
    prompt=("تو یک معلم حرفه‌ای ریاضی و علوم هستی. سؤال دانش‌آموز را مرحله‌به‌مرحله و دقیق حل کن. "
            "اگر ریاضی است فرمول و محاسبات را واضح بنویس و در پایان جواب نهایی را جدا کن. "
            "اگر سؤال مبهم است، فرض منطقی را شما بگو. پاسخ فارسی باشد.")
    return await ask_ai(text, context=_DummyContext(), system_prompt=prompt) if False else await ask_ai(text, _CURRENT_CONTEXT, prompt)


_CURRENT_CONTEXT = None

async def solve_student_problem_with_context(text, context):
    prompt=("تو معلم حرفه‌ای ریاضی، فیزیک و شیمی هستی. سؤال را مرحله‌به‌مرحله حل کن، "
            "فرمول‌ها را واضح بنویس، محاسبات را بررسی کن و جواب نهایی را با عنوان «جواب نهایی» بده. فارسی پاسخ بده.")
    return await ask_ai(text, context, prompt)


async def solve_image_question(file_path: str, context) -> str:
    try:
        with open(file_path,"rb") as f: b64=base64.b64encode(f.read()).decode()
        payload={"model":GROQ_VISION_MODEL,"messages":[
            {"role":"system","content":"تو یک معلم حرفه‌ای هستی. متن و سؤال داخل تصویر را دقیق بخوان و مرحله‌به‌مرحله حل کن. فارسی پاسخ بده و جواب نهایی را مشخص کن."},
            {"role":"user","content":[{"type":"text","text":"این سؤال را از روی تصویر بخوان و حل کن."},{"type":"image_url","image_url":{"url":f"data:image/jpeg;base64,{b64}"}}]}
        ],"temperature":0.1,"max_tokens":1800}
        async with httpx.AsyncClient(timeout=90) as client:
            r=await client.post("https://api.groq.com/openai/v1/chat/completions",headers={"Authorization":f"Bearer {GROQ_API_KEY}","Content-Type":"application/json"},json=payload)
            r.raise_for_status(); return r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"Image solve error: {e}")
        return "❌ حل سؤال از روی تصویر فعلاً در دسترس نیست. مطمئن شو مدل Vision در حساب API فعال است."


async def set_price_alert(user_id, symbol, target, direction):
    price_alerts[user_id].append({"symbol":symbol,"target":float(target),"direction":direction,"created":time.time()})
    save_alerts()


async def get_alert_price(symbol):
    """Get current price for alert check from DO_L4 (dollar / gold18 / coin)."""
    symbol = (symbol or "").lower().strip()
    try:
        html = await fetch_channel(DO_L4_URL)
        messages = extract_messages(html)
        summary = {}
        for msg in messages:
            if "نرخ فروش" in msg and ("دلار" in msg or "سکه" in msg):
                summary = parse_do_l4_summary(msg)
                if summary:
                    break
        for msg in messages:
            if "دلار فردایی" in msg:
                live = parse_do_l4_live_dollar(msg)
                if live:
                    summary["dollar"] = live
                    break
        if symbol in ("dollar", "usd", "دلار"):
            return summary.get("dollar")
        if symbol in ("gold", "gold18", "طلا", "طلای18", "طلای 18"):
            return summary.get("gold18")
        if symbol in ("coin", "سکه", "امامی", "سکه جدید"):
            return summary.get("coin_new")
    except Exception as e:
        logger.warning(f"get_alert_price error: {e}")
    return None


async def price_alert_loop(app):
    while True:
        try:
            for uid, alerts in list(price_alerts.items()):
                remaining=[]
                for a in alerts:
                    price=await get_alert_price(a.get("symbol",""))
                    hit=False
                    if price is not None:
                        hit = price >= a["target"] if a["direction"]=="above" else price <= a["target"]
                    if hit:
                        try:
                            await bot_send_message_localized(app.bot, uid, f"🔔 <b>هشدار قیمت</b>\n\n{html.escape(a['symbol'])}: <b>{price:,.2f}</b>\nهدف: <b>{a['target']:,.2f}</b>", parse_mode="HTML")
                        except Exception: pass
                    else: remaining.append(a)
                price_alerts[uid]=remaining
            save_alerts()
        except Exception as e: logger.warning(f"Alert loop: {e}")
        await asyncio.sleep(60)


def extended_tools_keyboard():
    return InlineKeyboardMarkup([
        [LButton("🌤 آب‌وهوا", callback_data="weather_menu"), LButton("🔮 فال روزانه", callback_data="daily_fortune")],
        [LButton("🧮 محاسبه‌گر", callback_data="advanced_calculator"), LButton("💱 تبدیل ارز", callback_data="currency_convert")],
        [LButton("📅 تقویم امروز", callback_data="today_calendar"), LButton("🔔 یادآور", callback_data="reminder_menu")],
        [LButton("📰 اخبار", callback_data="news_menu")],
        [LButton("🎓 ابزار دانشجویی", callback_data="student_menu"), LButton("📸 حل سؤال از عکس", callback_data="solve_image")],
        [LButton("🔤 OCR عکس", callback_data="ocr_image"), LButton("💰 قیمت دلار و سکه", callback_data="price_do_l4")],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")]
    ])


# =========================================================
# NEW FEATURES: converter / panel / calendar / reminder / sticker-text / OCR
# =========================================================

IRAN_OCCASIONS = {
    # (month, day): title  — ماه شمسی
    (1, 1): "🎉 نوروز – آغاز سال نو شمسی",
    (1, 2): "🎉 تعطیلات نوروز",
    (1, 3): "🎉 تعطیلات نوروز",
    (1, 4): "🎉 تعطیلات نوروز",
    (1, 12): "🇮🇷 روز جمهوری اسلامی",
    (1, 13): "🌿 سیزده‌به‌در",
    (3, 14): "🕊 رحلت امام خمینی",
    (3, 15): "🇮🇷 قیام ۱۵ خرداد",
    (11, 22): "🇮🇷 پیروزی انقلاب اسلامی",
    (12, 29): "🔥 روز ملی شدن صنعت نفت",
}

async def get_dollar_rate() -> int | None:
    """نرخ دلار تهران از DO_L4 (تومان)."""
    try:
        html = await fetch_channel(DO_L4_URL)
        messages = extract_messages(html)
        for msg in messages:
            if "دلار فردایی" in msg:
                val = parse_do_l4_live_dollar(msg)
                if val:
                    return val
        for msg in messages:
            if "نرخ فروش" in msg and "دلار" in msg:
                parsed = parse_do_l4_summary(msg)
                if parsed.get("dollar"):
                    return parsed["dollar"]
    except Exception as e:
        logger.warning(f"get_dollar_rate: {e}")
    return None


async def convert_currency(text: str) -> str:
    """تبدیل واحد پول با نرخ دلار لحظه‌ای."""
    t = normalize_text(text).replace(",", "").replace("٬", "")
    t = t.translate(str.maketrans("۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩", "01234567890123456789"))
    rate = await get_dollar_rate()
    if not rate:
        return "❌ فعلاً نرخ دلار در دسترس نیست. چند دقیقه بعد دوباره امتحان کن."

    # patterns: 100 دلار / 100 دلار به تومان / 5000000 تومان به دلار / 100 usd
    m = re.search(r"([\d.]+)\s*(دلار|dollar|usd|\$)\s*(?:به\s*)?(تومان|toman)?", t, re.I)
    if m:
        amount = float(m.group(1))
        toman = int(round(amount * rate))
        return (
            f"💱 <b>تبدیل ارز</b>\n\n"
            f"🇺🇸 {fmt_price(int(amount) if amount == int(amount) else amount)} دلار\n"
            f"🇮🇷 <b>{fmt_price(toman)}</b> تومان\n\n"
            f"نرخ دلار: <b>{fmt_price(rate)}</b> تومان\n"
            f"🕔 {get_jalali_datetime()}\n📡 منبع: @DO_L4"
        )

    m = re.search(r"([\d.]+)\s*(تومان|toman|تومن)\s*(?:به\s*)?(دلار|dollar|usd)?", t, re.I)
    if m:
        amount = float(m.group(1))
        dollars = amount / rate
        if dollars >= 1:
            d_txt = f"{dollars:,.2f}"
        else:
            d_txt = f"{dollars:,.4f}"
        return (
            f"💱 <b>تبدیل ارز</b>\n\n"
            f"🇮🇷 {fmt_price(int(amount) if amount == int(amount) else int(amount))} تومان\n"
            f"🇺🇸 <b>{d_txt}</b> دلار\n\n"
            f"نرخ دلار: <b>{fmt_price(rate)}</b> تومان\n"
            f"🕔 {get_jalali_datetime()}\n📡 منبع: @DO_L4"
        )

    return (
        "❌ فرمت را متوجه نشدم.\n\n"
        "مثال‌ها:\n"
        "• ۱۰۰ دلار\n"
        "• ۱۰۰ دلار به تومان\n"
        "• ۵۰۰۰۰۰۰ تومان به دلار"
    )


def get_user_panel(user_id: int, name: str) -> str:
    """پنل کاربری شخصی."""
    today = str(date.today())
    dl_info = user_downloads.get(user_id, {"date": today, "count": 0})
    cfg_info = user_config_usage.get(user_id, {"date": today, "count": 0})
    dl_count = dl_info["count"] if dl_info.get("date") == today else 0
    cfg_count = cfg_info["count"] if cfg_info.get("date") == today else 0
    join = user_join_dates.get(user_id, "—")
    rem_dl = max(0, DAILY_LIMIT - dl_count)
    rem_cfg = max(0, DAILY_CONFIG_LIMIT - cfg_count)
    user_rems = reminders.get(user_id, [])
    return (
        f"👤 <b>پنل کاربری</b>\n\n"
        f"نام: <b>{html.escape(name or 'کاربر')}</b>\n"
        f"آیدی: <code>{user_id}</code>\n"
        f"تاریخ عضویت: <b>{join}</b>\n\n"
        f"📥 دانلود امروز: <b>{dl_count}</b> / {DAILY_LIMIT}\n"
        f"⏳ باقی‌مانده دانلود: <b>{rem_dl}</b>\n\n"
        f"🩷 کانفیگ امروز: <b>{cfg_count}</b> / {DAILY_CONFIG_LIMIT}\n"
        f"⏳ باقی‌مانده کانفیگ: <b>{rem_cfg}</b>\n\n"
        f"🔔 یادآورهای فعال: <b>{len(user_rems)}</b>\n\n"
        f"🕔 {get_jalali_datetime()}"
    )


def get_today_calendar() -> str:
    """تقویم شمسی + مناسبت امروز."""
    now = jdatetime.datetime.now()
    weekdays = {0: "شنبه", 1: "یکشنبه", 2: "دوشنبه", 3: "سه‌شنبه", 4: "چهارشنبه", 5: "پنج‌شنبه", 6: "جمعه"}
    months = {1: "فروردین", 2: "اردیبهشت", 3: "خرداد", 4: "تیر", 5: "مرداد", 6: "شهریور",
              7: "مهر", 8: "آبان", 9: "آذر", 10: "دی", 11: "بهمن", 12: "اسفند"}
    occasion = IRAN_OCCASIONS.get((now.month, now.day))
    lines = [
        "📅 <b>تقویم امروز</b>\n",
        f"📌 {weekdays[now.weekday()]} {now.day} {months[now.month]} {now.year}",
        f"🕔 ساعت: <b>{now.hour:02d}:{now.minute:02d}</b>",
    ]
    if occasion:
        lines.append(f"\n✨ مناسبت: <b>{occasion}</b>")
    else:
        lines.append("\n✨ مناسبت رسمی خاصی ثبت نشده.")
    lines.append(f"\n📡 تقویم هجری شمسی")
    return "\n".join(lines)


def parse_reminder_time(text: str):
    """پارس ساده زمان یادآور. برمی‌گرداند (due_ts, clean_text) یا None."""
    from datetime import timedelta
    text = text.strip()
    now = datetime.now()
    t_norm = text.translate(str.maketrans("۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩", "01234567890123456789"))

    m = re.search(r"(?:یادآوری|یادآور|reminder)?\s*فردا\s*(?:ساعت\s*)?(\d{1,2})(?::(\d{2}))?\s+(.+)", t_norm, re.I)
    if m:
        hour = int(m.group(1))
        minute = int(m.group(2) or 0)
        body = m.group(3).strip()
        if 0 <= hour <= 23 and 0 <= minute <= 59 and body:
            due = (now + timedelta(days=1)).replace(hour=hour, minute=minute, second=0, microsecond=0)
            return due.timestamp(), body

    m = re.search(r"(?:یادآوری|یادآور|reminder)?\s*(?:ساعت\s*)?(\d{1,2})(?::(\d{2}))?\s+(.+)", t_norm, re.I)
    if m:
        hour = int(m.group(1))
        minute = int(m.group(2) or 0)
        body = m.group(3).strip()
        if 0 <= hour <= 23 and 0 <= minute <= 59 and body:
            due = now.replace(hour=hour, minute=minute, second=0, microsecond=0)
            if due <= now:
                due = due + timedelta(days=1)
            return due.timestamp(), body

    return None


async def reminder_loop(app):
    while True:
        try:
            now_ts = time.time()
            for uid, items in list(reminders.items()):
                remaining = []
                for r in items:
                    if r.get("due_ts", 0) <= now_ts:
                        try:
                            await bot_send_message_localized(
                                app.bot, uid,
                                f"🔔 <b>یادآور</b>\n\n{html.escape(r.get('text', ''))}",
                                parse_mode="HTML",
                            )
                        except Exception:
                            pass
                    else:
                        remaining.append(r)
                reminders[uid] = remaining
            save_reminders()
        except Exception as e:
            logger.warning(f"Reminder loop: {e}")
        await asyncio.sleep(30)




# =========================================================
# فونت و شکل‌دهی فارسی/کوردی برای استیکر (Termux-friendly)
# شکل‌دهی: اولویت با RAQM/HarfBuzz + fallback به arabic-reshaper
# کاربر می‌تواند از بین فونت‌های کوردی/عربی انتخاب کند
# =========================================================

# ۱۵ فونت مناسب فارسی/کوردی (دانلود خودکار در صورت نبود)
STICKER_FONTS = {
    "noto_naskh_b": {
        "label": "📖 Noto Naskh Bold",
        "file": "NotoNaskhArabic-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/googlefonts/noto-fonts@main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Bold.ttf",
            "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Bold.ttf",
        ],
    },
    "noto_naskh_r": {
        "label": "📖 Noto Naskh",
        "file": "NotoNaskhArabic-Regular.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/googlefonts/noto-fonts@main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
        ],
    },
    "noto_sans_b": {
        "label": "🔤 Noto Sans Bold",
        "file": "NotoSansArabic-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/googlefonts/noto-fonts@main/hinted/ttf/NotoSansArabic/NotoSansArabic-Bold.ttf",
        ],
    },
    "noto_kufi_b": {
        "label": "🕌 Noto Kufi Bold",
        "file": "NotoKufiArabic-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/googlefonts/noto-fonts@main/hinted/ttf/NotoKufiArabic/NotoKufiArabic-Bold.ttf",
        ],
    },
    "amiri_b": {
        "label": "✒️ Amiri Bold",
        "file": "Amiri-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/aliftype/amiri@main/fonts/ttf/Amiri-Bold.ttf",
            "https://github.com/aliftype/amiri/raw/main/fonts/ttf/Amiri-Bold.ttf",
        ],
    },
    "amiri_r": {
        "label": "✒️ Amiri",
        "file": "Amiri-Regular.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/aliftype/amiri@main/fonts/ttf/Amiri-Regular.ttf",
        ],
    },
    "scheherazade": {
        "label": "📜 Scheherazade",
        "file": "ScheherazadeNew-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/scheherazadenew/ScheherazadeNew-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/scheherazadenew/ScheherazadeNew-Bold.ttf",
        ],
    },
    "lateef": {
        "label": "📝 Lateef",
        "file": "Lateef-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/lateef/Lateef-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/lateef/Lateef-Bold.ttf",
        ],
    },
    "harmattan": {
        "label": "🪶 Harmattan",
        "file": "Harmattan-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/harmattan/Harmattan-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/harmattan/Harmattan-Bold.ttf",
        ],
    },
    "cairo_b": {
        "label": "🏛 Cairo Bold",
        "file": "Cairo-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/cairo/static/Cairo-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/cairo/static/Cairo-Bold.ttf",
        ],
    },
    "ibm_plex": {
        "label": "🟦 IBM Plex Arabic",
        "file": "IBMPlexSansArabic-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/ibmplexsansarabic/IBMPlexSansArabic-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/ibmplexsansarabic/IBMPlexSansArabic-Bold.ttf",
        ],
    },
    "vazirmatn": {
        "label": "🇮🇷 Vazirmatn Bold",
        "file": "Vazirmatn-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/rastikerdar/vazirmatn@master/fonts/ttf/Vazirmatn-Bold.ttf",
            "https://github.com/rastikerdar/vazirmatn/raw/master/fonts/ttf/Vazirmatn-Bold.ttf",
        ],
    },
    "reem_kufi": {
        "label": "🔷 Reem Kufi",
        "file": "ReemKufi-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/reemkufi/ReemKufi-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/reemkufi/ReemKufi-Bold.ttf",
        ],
    },
    "el_messiri": {
        "label": "✨ El Messiri",
        "file": "ElMessiri-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/elmessiri/ElMessiri-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/elmessiri/ElMessiri-Bold.ttf",
        ],
    },
    "markazi": {
        "label": "📰 Markazi Text",
        "file": "MarkaziText-Bold.ttf",
        "urls": [
            "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/markazitext/MarkaziText-Bold.ttf",
            "https://github.com/google/fonts/raw/main/ofl/markazitext/MarkaziText-Bold.ttf",
        ],
    },
}

DEFAULT_STICKER_FONT = "noto_naskh_b"


def _font_search_dirs():
    dirs = []
    try:
        dirs.append(os.path.dirname(os.path.abspath(__file__)))
    except Exception:
        pass
    dirs.extend([
        os.path.join(os.getcwd(), "fonts"),
        os.getcwd(),
        os.path.expanduser("~/downloads"),
        os.path.expanduser("~/fonts"),
        os.path.expanduser("~/insta_bot/fonts"),
        "/data/data/com.termux/files/home/downloads",
        "/data/data/com.termux/files/home/fonts",
        "/usr/share/fonts/truetype/noto",
        "/usr/share/fonts/truetype",
        "/system/fonts",
    ])
    return dirs


def _find_local_font(filename: str) -> str | None:
    for d in _font_search_dirs():
        try:
            fp = os.path.join(d, filename)
            if os.path.isfile(fp) and os.path.getsize(fp) > 15000:
                return fp
        except Exception:
            continue
    return None


def ensure_sticker_font(font_key: str | None = None) -> str | None:
    """پیدا کردن یا دانلود فونت انتخاب‌شده توسط کاربر."""
    key = font_key or DEFAULT_STICKER_FONT
    info = STICKER_FONTS.get(key) or STICKER_FONTS[DEFAULT_STICKER_FONT]
    filename = info["file"]

    local = _find_local_font(filename)
    if local:
        return local

    # مسیر ذخیره: کنار ربات / fonts
    try:
        base = os.path.dirname(os.path.abspath(__file__))
    except Exception:
        base = os.getcwd()
    fonts_dir = os.path.join(base, "fonts")
    try:
        os.makedirs(fonts_dir, exist_ok=True)
    except Exception:
        fonts_dir = base
    target = os.path.join(fonts_dir, filename)

    try:
        import urllib.request
        for url in info.get("urls", []):
            try:
                logger.info(f"Downloading font {key}: {url}")
                req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=60) as resp:
                    data = resp.read()
                if len(data) < 15000:
                    continue
                with open(target, "wb") as f:
                    f.write(data)
                if os.path.isfile(target) and os.path.getsize(target) > 15000:
                    logger.info(f"Font saved -> {target}")
                    return target
            except Exception as e:
                logger.warning(f"Font download fail ({key}): {e}")
    except Exception as e:
        logger.warning(f"ensure_sticker_font error: {e}")

    # آخرین تلاش: هر فونت عربی موجود
    for k, inf in STICKER_FONTS.items():
        loc = _find_local_font(inf["file"])
        if loc:
            return loc
    return None


# سازگاری با کد قدیمی
def _ensure_sticker_font() -> str | None:
    return ensure_sticker_font(DEFAULT_STICKER_FONT)


def _has_raqm() -> bool:
    try:
        from PIL import features
        return bool(features.check("raqm"))
    except Exception:
        return False


def _shape_rtl_text(text: str) -> str:
    if not text:
        return text
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        cfg = {"delete_harakat": False, "support_ligatures": True, "language": "ArabicV2"}
        try:
            reshaper = arabic_reshaper.ArabicReshaper(configuration=cfg)
            reshaped = reshaper.reshape(text)
        except Exception:
            reshaped = arabic_reshaper.reshape(text)
        try:
            return get_display(reshaped, base_dir="R")
        except TypeError:
            return get_display(reshaped)
    except ImportError:
        logger.warning("arabic-reshaper/python-bidi missing")
        return text
    except Exception as e:
        logger.warning(f"reshape error: {e}")
        return text


def _rtl_packages_ok() -> bool:
    if _has_raqm():
        return True
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        sample = "سلام"
        out = get_display(arabic_reshaper.reshape(sample))
        return bool(out) and out != sample
    except Exception:
        return False


def _get_arabic_font(size: int, font_key: str | None = None, use_raqm: bool = True):
    from PIL import ImageFont
    fp = ensure_sticker_font(font_key)
    if not fp:
        logger.error("No Arabic font found for stickers")
        return ImageFont.load_default(), False

    raqm_on = False
    font = None
    if use_raqm and _has_raqm():
        try:
            font = ImageFont.truetype(fp, size, layout_engine=ImageFont.Layout.RAQM)
            raqm_on = True
        except Exception as e:
            logger.warning(f"RAQM font load failed: {e}")
    if font is None:
        try:
            font = ImageFont.truetype(fp, size)
        except Exception as e:
            logger.error(f"Font load error {fp} @ {size}: {e}")
            try:
                font = ImageFont.truetype(fp, max(24, size // 2))
            except Exception:
                return ImageFont.load_default(), False
    return font, raqm_on


STICKER_BG_COLORS = {
    "bg_blue": ((40, 80, 180, 255), "🔵 آبی"),
    "bg_red": ((180, 40, 50, 255), "🔴 قرمز"),
    "bg_green": ((30, 130, 70, 255), "🟢 سبز"),
    "bg_purple": ((110, 50, 160, 255), "🟣 بنفش"),
    "bg_black": ((25, 25, 25, 255), "⚫ مشکی"),
    "bg_orange": ((220, 120, 30, 255), "🟠 نارنجی"),
    "bg_teal": ((20, 130, 130, 255), "🩵 فیروزه‌ای"),
    "bg_pink": ((200, 60, 120, 255), "🩷 صورتی"),
}
STICKER_FG_COLORS = {
    "fg_white": ((255, 255, 255, 255), "⬜ سفید"),
    "fg_black": ((20, 20, 20, 255), "⬛ مشکی"),
    "fg_yellow": ((255, 230, 50, 255), "🟡 زرد"),
    "fg_cyan": ((100, 240, 255, 255), "🩵 آبی روشن"),
}


def sticker_bg_keyboard():
    rows = []
    items = list(STICKER_BG_COLORS.items())
    for i in range(0, len(items), 2):
        row = []
        for j in range(2):
            if i + j < len(items):
                key, (_, label) = items[i + j]
                row.append(LButton(label, callback_data=f"stbg_{key}"))
        rows.append(row)
    rows.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)


def sticker_fg_keyboard():
    rows = []
    items = list(STICKER_FG_COLORS.items())
    for i in range(0, len(items), 2):
        row = []
        for j in range(2):
            if i + j < len(items):
                key, (_, label) = items[i + j]
                row.append(LButton(label, callback_data=f"stfg_{key}"))
        rows.append(row)
    rows.append([LButton("🔙 تغییر پس‌زمینه", callback_data="text_sticker")])
    rows.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)


def image_sticker_fg_keyboard():
    """انتخاب رنگ متن برای استیکرهای ساخته‌شده از عکس."""
    rows = []
    items = list(STICKER_FG_COLORS.items())
    for i in range(0, len(items), 2):
        row = []
        for j in range(2):
            if i + j < len(items):
                key, (_, label) = items[i + j]
                row.append(LButton(label, callback_data=f"istfg_{key}"))
        rows.append(row)
    rows.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)


def sticker_background_type_keyboard():
    """انتخاب نوع پس‌زمینه برای ساخت استیکر."""
    return InlineKeyboardMarkup([
        [
            LButton("🖼 با عکس", callback_data="sticker_with_photo"),
            LButton("🎨 بدون عکس", callback_data="sticker_no_photo"),
        ],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")],
    ])


def image_sticker_font_keyboard():
    """انتخاب فونت برای متن روی عکس استیکر."""
    items = list(STICKER_FONTS.items())
    rows = []
    for i in range(0, len(items), 2):
        row = []
        for key, info in items[i:i + 2]:
            row.append(LButton(
                info["label"],
                callback_data=f"istfont_{key}"
            ))
        rows.append(row)
    rows.append([LButton("❌ لغو", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)


def sticker_font_keyboard():
    rows = []
    items = list(STICKER_FONTS.items())
    for i in range(0, len(items), 2):
        row = []
        for j in range(2):
            if i + j < len(items):
                key, info = items[i + j]
                row.append(LButton(info["label"], callback_data=f"stfont_{key}"))
        rows.append(row)
    rows.append([LButton("🔙 تغییر رنگ متن", callback_data="sticker_pick_fg")])
    rows.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)


def _wrap_rtl_lines(text: str, max_chars: int = 16) -> list:
    lines = []
    for paragraph in text.split("\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        words = paragraph.split()
        if not words:
            continue
        current = []
        for w in words:
            trial = " ".join(current + [w])
            if len(trial) <= max_chars or not current:
                current.append(w)
                if len(w) > max_chars and len(current) == 1:
                    chunk = current.pop()
                    while len(chunk) > max_chars:
                        lines.append(chunk[:max_chars])
                        chunk = chunk[max_chars:]
                    if chunk:
                        current = [chunk]
            else:
                lines.append(" ".join(current))
                current = [w]
        if current:
            lines.append(" ".join(current))
    return lines or [text]


def _draw_rtl_line(draw, xy, line: str, font, fill, raqm_on: bool):
    x, y = xy
    if raqm_on:
        try:
            draw.text((x, y), line, font=font, fill=fill, direction="rtl", language="ckb")
            return
        except TypeError:
            try:
                draw.text((x, y), line, font=font, fill=fill, direction="rtl")
                return
            except TypeError:
                pass
        except Exception as e:
            logger.warning(f"RAQM draw fallback: {e}")
    draw.text((x, y), line, font=font, fill=fill)


def _measure_line(draw, line: str, font, raqm_on: bool):
    if raqm_on:
        try:
            bbox = draw.textbbox((0, 0), line, font=font, direction="rtl", language="ckb")
            return bbox[2] - bbox[0], bbox[3] - bbox[1]
        except TypeError:
            try:
                bbox = draw.textbbox((0, 0), line, font=font, direction="rtl")
                return bbox[2] - bbox[0], bbox[3] - bbox[1]
            except Exception:
                pass
        except Exception:
            pass
    try:
        bbox = draw.textbbox((0, 0), line, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]
    except Exception:
        return 100, 40


def create_image_text_sticker(
    image_path: str,
    text: str,
    font_key: str | None = None,
    size: int = 512,
    fg_rgba=(255, 255, 255, 255),
) -> BytesIO | None:
    """ساخت استیکر عکس با متن دقیقاً وسط تصویر."""
    try:
        from PIL import Image, ImageDraw

        text = (text or "").strip()
        if not text:
            return None
        if len(text) > 180:
            text = text[:180].rstrip() + "…"

        font_key = font_key or DEFAULT_STICKER_FONT
        fp = ensure_sticker_font(font_key)
        if not fp:
            return None

        img = Image.open(image_path).convert("RGBA")
        img.thumbnail((size, size), Image.Resampling.LANCZOS)

        # بوم مربعی استیکر، با خود عکس در مرکز
        canvas = Image.new("RGBA", (size, size), (0, 0, 0, 0))
        x0 = (size - img.width) // 2
        y0 = (size - img.height) // 2
        canvas.alpha_composite(img, (x0, y0))

        draw = ImageDraw.Draw(canvas)
        font_size = max(28, min(92, int(size * 0.105)))
        font, raqm_on = _get_arabic_font(font_size, font_key, True)

        # متن را با عرض واقعی فونت به خطوط تقسیم می‌کنیم
        max_w = int(size * 0.86)
        raw_lines = []
        for paragraph in text.splitlines() or [text]:
            words = paragraph.split()
            if not words:
                raw_lines.append("")
                continue
            line = words[0]
            for word in words[1:]:
                candidate = line + " " + word
                w, _ = _measure_line(draw, candidate, font, raqm_on)
                if w > max_w:
                    raw_lines.append(line)
                    line = word
                else:
                    line = candidate
            raw_lines.append(line)

        line_info = []
        for line in raw_lines:
            w, h = _measure_line(draw, line, font, raqm_on)
            line_info.append((line, w, h))

        spacing = max(6, int(font_size * 0.16))
        total_h = sum(h for _, _, h in line_info) + spacing * max(0, len(line_info)-1)

        # مرکز عمودی دقیق
        top_y = (size - total_h) / 2

        # سایه برای خوانایی، سپس رنگ انتخابی کاربر
        shadow = (0, 0, 0, 180)
        fill = tuple(fg_rgba) if fg_rgba else (255, 255, 255, 255)
        if len(fill) == 3:
            fill = (*fill, 255)

        y = top_y
        for line, w, h in line_info:
            # مرکز افقی دقیق با عرض اندازه‌گیری‌شده
            x = (size - w) / 2

            # برای RTL، _draw_rtl_line از همان x استفاده می‌کند.
            if raqm_on:
                _draw_rtl_line(draw, (x + 3, y + 3), line, font, shadow, raqm_on)
                _draw_rtl_line(draw, (x, y), line, font, fill, raqm_on)
            else:
                draw.text((x + 3, y + 3), line, font=font, fill=shadow)
                draw.text((x, y), line, font=font, fill=fill)

            y += h + spacing

        out = BytesIO()
        out.name = "sticker.webp"
        canvas.save(out, format="WEBP", lossless=True, method=6)
        out.seek(0)
        return out

    except Exception as e:
        logger.exception("create_image_text_sticker failed: %s", e)
        return None


def create_text_sticker(
    text: str,

    bg_rgba=(40, 80, 180, 255),
    fg_rgba=(255, 255, 255, 255),
    size: int = 512,
    font_key: str | None = None,
) -> BytesIO | None:
    """
    استیکر متنی فارسی/کوردی با فونت انتخابی کاربر.
    """
    try:
        from PIL import ImageDraw

        text = (text or "").strip()
        if not text:
            return None
        if len(text) > 120:
            text = text[:120].rstrip() + "…"

        font_key = font_key or DEFAULT_STICKER_FONT
        if not ensure_sticker_font(font_key):
            logger.error("create_text_sticker: font missing")
            return None

        raqm_on = _has_raqm()
        max_chars = 12 if len(text) > 30 else (14 if len(text) > 15 else 18)
        logical_lines = _wrap_rtl_lines(text, max_chars=max_chars)
        draw_lines = logical_lines if raqm_on else [_shape_rtl_text(line) for line in logical_lines]

        n_lines = len(draw_lines)
        if n_lines <= 1:
            font_size = 78
        elif n_lines == 2:
            font_size = 62
        elif n_lines <= 4:
            font_size = 48
        else:
            font_size = 38

        font, raqm_on = _get_arabic_font(font_size, font_key=font_key, use_raqm=True)
        img = Image.new("RGBA", (size, size), bg_rgba)
        draw = ImageDraw.Draw(img)

        def measure_all(lines, fnt, rq):
            widths, heights = [], []
            for line in lines:
                w, h = _measure_line(draw, line, fnt, rq)
                widths.append(max(1, w))
                heights.append(max(h, int(font_size * 0.85)))
            return widths, heights

        line_widths, line_heights = measure_all(draw_lines, font, raqm_on)
        gap = 14
        total_h = sum(line_heights) + max(0, n_lines - 1) * gap

        tries = 0
        while (total_h > size * 0.88 or max(line_widths) > size * 0.90) and font_size > 26 and tries < 6:
            font_size = max(26, int(font_size * 0.82))
            font, raqm_on = _get_arabic_font(font_size, font_key=font_key, use_raqm=True)
            line_widths, line_heights = measure_all(draw_lines, font, raqm_on)
            gap = max(8, int(font_size * 0.22))
            total_h = sum(line_heights) + max(0, n_lines - 1) * gap
            tries += 1

        y = (size - total_h) // 2
        shadow = (0, 0, 0, 160)
        for i, line in enumerate(draw_lines):
            w = line_widths[i]
            x = (size - w) // 2
            _draw_rtl_line(draw, (x + 2, y + 2), line, font, shadow, raqm_on)
            _draw_rtl_line(draw, (x, y), line, font, fg_rgba, raqm_on)
            y += line_heights[i] + gap

        bio = BytesIO()
        img.save(bio, format="WEBP", quality=95)
        bio.seek(0)
        return bio
    except Exception as e:
        logger.error(f"create_text_sticker error: {e}")
        return None


async def ocr_image(file_path: str) -> str:
    """استخراج متن از تصویر با مدل Vision فعال Groq."""
    try:
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        payload = {
            "model": GROQ_VISION_MODEL,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "تو یک موتور OCR دقیق هستی. تمام متن داخل تصویر را استخراج کن. "
                        "اگر فارسی یا کوردی است همان‌طور بنویس. فقط متن را برگردان، توضیح اضافه نده."
                    ),
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "متن داخل این تصویر را کامل استخراج کن."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                    ],
                },
            ],
            "temperature": 0.1,
            "max_tokens": 1500,
        }
        async with httpx.AsyncClient(timeout=90) as client:
            r = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json=payload,
            )
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"OCR error: {e}")
        return "❌ استخراج متن از تصویر ناموفق بود."


async def summarize_voice_text(text: str) -> str:
    """خلاصه متن استخراج‌شده از ویس."""
    if not text or len(text) < 20:
        return text or "—"
    try:
        async with httpx.AsyncClient(timeout=40) as client:
            resp = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": [
                        {
                            "role": "system",
                            "content": "متن را به فارسی خلاصه و مرتب کن. کوتاه و واضح باشد.",
                        },
                        {"role": "user", "content": text},
                    ],
                    "temperature": 0.3,
                    "max_tokens": 600,
                },
            )
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        return text



# =========================================================
# FREE CONFIG
# =========================================================

def extract_host_port(config: str):
    try:
        if config.startswith("vmess://"):
            import base64
            encoded = config[8:]
            padded = encoded + "=" * (-len(encoded) % 4)
            data = json.loads(base64.b64decode(padded).decode("utf-8", errors="ignore"))
            return data.get("add"), int(data.get("port", 0))
        elif config.startswith(("vless://", "trojan://", "ss://")):
            main_part = config.split("://", 1)[1]
            host_port = main_part.split("@")[1].split("?")[0].split("#")[0] if "@" in main_part else main_part.split("?")[0].split("#")[0]
            if ":" in host_port:
                host, port = host_port.rsplit(":", 1)
                return host.strip(), int(port)
    except Exception:
        pass
    return None, None


async def test_config_tcp(host: str, port: int, timeout: float = 2.5) -> bool:
    if not host or not port:
        return False
    try:
        _, writer = await asyncio.wait_for(asyncio.open_connection(host, port), timeout=timeout)
        writer.close()
        await writer.wait_closed()
        return True
    except Exception:
        return False


async def fetch_configs(protocol: str) -> list:
    urls = CONFIG_SOURCES.get(protocol, [])
    configs = []
    async with httpx.AsyncClient(timeout=12.0) as client:
        for url in urls:
            try:
                resp = await client.get(url)
                if resp.status_code == 200:
                    for line in resp.text.strip().splitlines():
                        line = line.strip()
                        if line.startswith(("vless://", "vmess://", "trojan://", "ss://")):
                            configs.append(line)
            except Exception as e:
                logger.error(f"Fetch error: {e}")
    return list(set(configs))


async def get_working_config(protocol: str, max_tries: int = 10) -> str | None:
    if protocol == "random":
        protocol = random.choice(["vless", "vmess", "trojan", "ss"])
    configs = await fetch_configs(protocol)
    if not configs:
        return None
    random.shuffle(configs)
    for config in configs[:max_tries]:
        host, port = extract_host_port(config)
        if host and port and await test_config_tcp(host, port):
            return config
    return None


# =========================================================
# TEXT BEAUTIFIER
# =========================================================

def convert_to_fancy_fonts(text: str) -> dict:
    """Create safe Unicode text styles. Persian text stays untouched."""
    maps = {
        "پررنگ": str.maketrans(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789",
            "𝐀𝐁𝐂𝐃𝐄𝐅𝐆𝐇𝐈𝐉𝐊𝐋𝐌𝐍𝐎𝐏𝐐𝐑𝐒𝐓𝐔𝐕𝐖𝐗𝐘𝐙"
            "𝐚𝐛𝐜𝐝𝐞𝐟𝐠𝐡𝐢𝐣𝐤𝐥𝐦𝐧𝐨𝐩𝐪𝐫𝐬𝐭𝐮𝐯𝐰𝐱𝐲𝐳"
            "𝟎𝟏𝟐𝟑𝟒𝟓𝟔𝟕𝟖𝟗"
        ),
        "ایتالیک": str.maketrans(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz",
            "𝘈𝘉𝘊𝘋𝘌𝘍𝘎𝘏𝘐𝘑𝘒𝘓𝘔𝘕𝘖𝘗𝘘𝘙𝘚𝘛𝘜𝘝𝘞𝘟𝘠𝘡"
            "𝘢𝘣𝘤𝘥𝘦𝘧𝘨𝘩𝘪𝘫𝘬𝘭𝘮𝘯𝘰𝘱𝘲𝘳𝘴𝘵𝘶𝘷𝘸𝘹𝘺𝘻"
        ),
        "پررنگ ایتالیک": str.maketrans(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz",
            "𝑨𝑩𝑪𝑫𝑬𝑭𝑮𝑯𝑰𝑱𝑲𝑳𝑴𝑵𝑶𝑷𝑸𝑹𝑺𝑻𝑼𝑽𝑾𝑿𝒀𝒁"
            "𝒂𝒃𝒄𝒅𝒆𝒇𝒈𝒉𝒊𝒋𝒌𝒍𝒎𝒏𝒐𝒑𝒒𝒓𝒔𝒕𝒖𝒗𝒘𝒙𝒚𝒛"
        ),
        "حروف دایره‌ای": str.maketrans(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789",
            "ⒶⒷⒸⒹⒺⒻⒼⒽⒾⒿⓀⓁⓂⓃⓄⓅⓆⓇⓈⓉⓊⓋⓌⓍⓎⓏ"
            "ⓐⓑⓒⓓⓔⓕⓖⓗⓘⓙⓚⓛⓜⓝⓞⓟⓠⓡⓢⓣⓤⓥⓦⓧⓨⓩ"
            "⓪①②③④⑤⑥⑦⑧⑨"
        ),
        "تمام‌عرض": str.maketrans(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789",
            "ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺ"
            "ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ"
            "０１２３４５６７８９"
        ),
    }
    result = {name: text.translate(table) for name, table in maps.items()}
    result["اصلی"] = text
    result["با فاصله"] = " ".join(list(text.replace(" ", "")))
    result["نقطه‌دار"] = "·".join(list(text.replace(" ", "")))
    return result

def create_word_document(text: str) -> BytesIO:
    doc = Document()
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =========================================================
# KEYBOARDS
# =========================================================

def get_reply_keyboard():
    return ReplyKeyboardMarkup([
        [LKeyboardButton("🏠 منوی اصلی"), LKeyboardButton("🔄 ری‌استارت")],
        [LKeyboardButton("ℹ️ راهنما"), LKeyboardButton("🌐 زبان / Language")]
    ], resize_keyboard=True, is_persistent=True)


def get_main_keyboard():
    return InlineKeyboardMarkup([
        # --- دانلود ---
        [LButton("📥 دانلود فیلم و ویدیو", callback_data="header_download")],
        [
            LButton("📸 اینستاگرام", callback_data="platform_instagram"),
            LButton("▶️ یوتیوب", callback_data="platform_youtube"),
            LButton("🎵 تیک‌تاک", callback_data="platform_tiktok"),
        ],
        [
            LButton("🐦 توییتر", callback_data="platform_twitter"),
            LButton("🔴 ردیت", callback_data="platform_reddit"),
            LButton("🔍 جستجوی یوتیوب", callback_data="youtube_search"),
        ],
        [LButton("🔗 تبدیل ویدیو به MP3", callback_data="video_to_mp3")],

        # --- موسیقی ---
        [LButton("🎵 موسیقی و آهنگ", callback_data="header_music")],
        [
            LButton("🎶 پیدا کردن آهنگ", callback_data="music_finder"),
            LButton("📝 متن آهنگ", callback_data="lyrics_menu"),
        ],
        [LButton("📜 زیرنویس یوتیوب", callback_data="subtitle_menu")],

        # --- قیمت‌ها ---
        [LButton("💰 قیمت دلار و سکه", callback_data="price_do_l4")],

        # --- امکانات هوشمند (یکجا) ---
        [LButton("🚀 امکانات هوشمند", callback_data="extended_tools")],

        # --- ابزارهای کاربردی ---
        [LButton("🛠 ابزارهای کاربردی", callback_data="header_tools")],
        [
            LButton("🖼 ساخت استیکر", callback_data="make_sticker"),
            LButton("📱 QR Code", callback_data="qr_code"),
            LButton("🗜 فشرده عکس", callback_data="compress_image"),
        ],
        [
            LButton("🗜 فشرده ویدیو", callback_data="compress_video"),
            LButton("📄 عکس به PDF", callback_data="photo_to_pdf"),
            LButton("📝 متن به ورد", callback_data="text_to_word"),
        ],
        [
            LButton("✍ زیباسازی متن", callback_data="fancy_text"),
            LButton("🔗 کوتاه لینک", callback_data="short_link"),
            LButton("🕔 ساعت", callback_data="show_clock"),
        ],
        [
            LButton("💱 تبدیل ارز", callback_data="currency_convert"),
            LButton("👤 پنل من", callback_data="user_panel"),
        ],
        [LButton("📁 فایل‌خوان هوشمند", callback_data="smart_file_reader")],

        # --- هوش مصنوعی ---
        [LButton("🤖 هوش مصنوعی", callback_data="header_ai")],
        [
            LButton("💬 چت AI", callback_data="ai_chat"),
            LButton("🌐 ترجمه", callback_data="translate_text"),
            LButton("📝 خلاصه‌سازی", callback_data="summarize_text"),
        ],
        [
            LButton("🎤 ویس به متن", callback_data="voice_to_text"),
            LButton("🇬🇧 معلم انگلیسی", callback_data="english_teacher"),
        ],
        [LButton("🎬 ساخت زیرنویس (کوردی / فارسی / انگلیسی)", callback_data="subtitle_maker")],

        # --- کانفیگ ---
        [LButton("🩷 کانفیگ رایگان", callback_data="free_config")],

        [
            LButton("ℹ️ راهنما", callback_data="help"),
            LButton("📩 پشتیبانی", url="https://t.me/sir_Aso"),
        ],
    ])


def get_subtitle_lang_keyboard():
    return InlineKeyboardMarkup([
        [LButton("☀️ کوردی", callback_data="sub_lang_ckb")],
        [LButton("🇮🇷 فارسی", callback_data="sub_lang_fa")],
        [LButton("🇬🇧 انگلیسی", callback_data="sub_lang_en")],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")],
    ])


def get_subtitle_result_keyboard():
    return InlineKeyboardMarkup([
        [LButton("📄 دریافت فایل SRT", callback_data="sub_send_srt")],
        [LButton("🎥 دریافت ویدیوی زیرنویس‌دار", callback_data="sub_send_video")],
        [LButton("📦 دریافت هر دو", callback_data="sub_send_both")],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")],
    ])


def get_config_keyboard():
    return InlineKeyboardMarkup([
        [LButton("VLESS", callback_data="cfg_vless")],
        [LButton("VMess", callback_data="cfg_vmess")],
        [LButton("Trojan", callback_data="cfg_trojan")],
        [LButton("Shadowsocks", callback_data="cfg_ss")],
        [LButton("🎲 تصادفی", callback_data="cfg_random")],
        [LButton("🔙 برگشت", callback_data="main_menu")],
    ])


def get_english_teacher_keyboard():
    return InlineKeyboardMarkup([
        [LButton("💬 مکالمه آزاد", callback_data="et_conversation")],
        [LButton("✍️ تصحیح جمله من", callback_data="et_correct")],
        [LButton("📚 یادگیری لغت جدید", callback_data="et_vocabulary")],
        [LButton("📖 تمرین گرامر", callback_data="et_grammar")],
        [LButton("🎯 تنظیم سطح من", callback_data="et_set_level")],
        [LButton("❌ خروج از معلم زبان", callback_data="et_exit")],
    ])


def get_level_keyboard():
    return InlineKeyboardMarkup([
        [LButton("🟢 مبتدی", callback_data="level_beginner")],
        [LButton("🟡 متوسط", callback_data="level_intermediate")],
        [LButton("🔴 پیشرفته", callback_data="level_advanced")],
        [LButton("🔙 برگشت", callback_data="english_teacher")],
    ])


def get_admin_keyboard():
    return InlineKeyboardMarkup([
        [LButton("📊 آمار ربات", callback_data="admin_stats")],
        [LButton("📢 پیام همگانی", callback_data="admin_broadcast")],
        [LButton("🚫 بلاک‌لیست", callback_data="admin_blacklist")],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")],
    ])


def get_youtube_quality_keyboard():
    return InlineKeyboardMarkup([
        [
            LButton("360p", callback_data="quality_360"),
            LButton("720p", callback_data="quality_720"),
            LButton("1080p", callback_data="quality_1080"),
        ],
        [LButton("🎵 فقط صدا (MP3)", callback_data="quality_audio")],
        [
            LButton("❌لغو", callback_data="cancel_download"),
            LButton("🏠 منوی اصلی", callback_data="main_menu"),
        ],
    ])


def get_after_download_keyboard(is_youtube=False):
    buttons = []
    if is_youtube:
        buttons.append([
            LButton("🔄 کیفیت دیگر", callback_data="change_quality"),
            LButton("🎵 صدا (MP3)", callback_data="quality_audio")
        ])
    buttons.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
    return InlineKeyboardMarkup(buttons)


def get_cancel_keyboard():
    return InlineKeyboardMarkup([[LButton("❌لغو", callback_data="cancel_download")]])


def get_back_keyboard():
    return InlineKeyboardMarkup([
        [LButton("❌لغو", callback_data="main_menu")],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")]
    ])


def get_ai_keyboard():
    return InlineKeyboardMarkup([[LButton("❌خروج از چت", callback_data="exit_ai")]])


# =========================================================
# AI + STT
# =========================================================

ASO_SYSTEM_PROMPT = (
    "تو آسو هستی؛ دستیار هوشمند، دوستانه و کمی شوخ‌طبع ربات AsoLand. "
    "اسمت آسو است و وقتی لازم است خودت را معرفی کن (مثلاً: سلام، من آسو هستم 😊). "
    "فقط و فقط به فارسی معیار و طبیعی (فارسی ایران) جواب بده. هرگز از واژه‌ها، جمله‌بندی یا املای کردی/سورانی استفاده نکن؛ حتی اگر کاربر کردی نوشت، پاسخ را فارسی بده. از حروف ویژه کردی مثل ڵ، ڕ، ێ، ۆ، ڤ استفاده نکن. مگر برای نام خاص، کد، دستور فنی یا نقل‌قول صریح، هیچ متن کردی یا زبان دیگری تولید نکن. جواب‌ها مفید، کوتاه تا متوسط و کمی شوخ باشند؛ "
    "از طنز ملایم و خودمونی استفاده کن ولی بی‌ادب یا توهین‌آمیز نباش. "
    "اگر کاربر احوال‌پرسی کرد، گرم و باحال جواب بده. "
    "در موضوعات جدی (مثل قیمت، محاسبه، درس) دقیق و جدی‌تر باش. "
    "هرگز خودت را مدل زبانی یا هوش مصنوعی گوگل/اوپن‌ای معرفی نکن؛ تو آسو هستی."
)

PERSIAN_ONLY_SYSTEM = """پاسخ زیر را فقط از نظر زبان و واژگان اصلاح کن.
آن را به فارسی معیار و طبیعی ایران تبدیل کن. هیچ واژه، عبارت یا املای کردی/سورانی باقی نگذار.
حروف ویژه کردی مانند ڵ، ڕ، ێ، ۆ، ڤ را حذف و معادل فارسی طبیعی‌شان را به کار ببر.
معنی، ساختار، اعداد، نام‌ها، کدها و اصطلاحات فنی را تغییر نده.
فقط متن اصلاح‌شده را برگردان و هیچ توضیحی اضافه نکن."""

def _has_kurdish_specific_chars(text: str) -> bool:
    return bool(re.search(r"[ڵڕێۆڤ]", text or ""))

async def _ensure_persian_output(text: str) -> str:
    if not text or not _has_kurdish_specific_chars(text):
        return text
    try:
        payload = {
            "model": GROQ_CHAT_MODEL,
            "messages": [
                {"role": "system", "content": PERSIAN_ONLY_SYSTEM},
                {"role": "user", "content": text},
            ],
            "temperature": 0.0,
            "max_tokens": max(1200, min(4096, len(text) * 3)),
        }
        async with httpx.AsyncClient(timeout=45.0) as client:
            resp = await client.post(
                GROQ_CHAT_URL,
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json=payload,
            )
            resp.raise_for_status()
            fixed = resp.json()["choices"][0]["message"]["content"].strip()
            return fixed or text
    except Exception as exc:
        logger.warning("Persian-only correction failed: %s", exc)
        return text

async def groq_chat(messages: list, *, temperature: float = 0.7, max_tokens: int = 1024, timeout: float = 60.0) -> str:
    """Reliable Groq chat wrapper used by AI Chat and English Teacher."""
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY is not configured")
    payload = {"model": GROQ_CHAT_MODEL, "messages": messages, "temperature": temperature, "max_completion_tokens": max_tokens}
    async with httpx.AsyncClient(timeout=timeout) as client:
        resp = await client.post(GROQ_CHAT_URL, headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}, json=payload)
    if resp.status_code >= 400:
        try:
            detail = resp.json().get("error", {})
            if isinstance(detail, dict): detail = detail.get("message") or detail
        except Exception:
            detail = resp.text
        raise RuntimeError(f"Groq HTTP {resp.status_code}: {str(detail)[:500]}")
    data = resp.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise RuntimeError(f"Invalid Groq response: {data!r}") from exc
    if not isinstance(content, str) or not content.strip():
        raise RuntimeError("Groq returned an empty response")
    content = content.strip()
    return await _ensure_persian_output(content)


def _valid_api_key(value: str) -> bool:
    return bool(value and not value.startswith("PASTE_YOUR_"))


def _split_telegram_text(text: str, limit: int = 3900):
    text = (text or "").strip()
    if not text:
        return ["❌ پاسخی دریافت نشد."]
    parts = []
    while len(text) > limit:
        cut = text.rfind("\n", 0, limit)
        if cut < limit // 2:
            cut = text.rfind(" ", 0, limit)
        if cut < limit // 2:
            cut = limit
        parts.append(text[:cut].rstrip())
        text = text[cut:].lstrip()
    parts.append(text)
    return parts


async def groq_chat(messages, *, temperature=0.7, max_tokens=1200, timeout=75.0):
    """مرکز واحد ارتباط با Groq برای جلوگیری از خطاهای تکراری."""
    if not _valid_api_key(GROQ_API_KEY):
        return None, "❌ کلید Groq تنظیم نشده است."

    try:
        async with httpx.AsyncClient(
            timeout=httpx.Timeout(timeout, connect=15.0)
        ) as client:
            response = await client.post(
                GROQ_CHAT_URL,
                headers={
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": messages,
                    "temperature": temperature,
                    "max_tokens": max_tokens,
                },
            )

        if response.status_code >= 400:
            try:
                error_data = response.json().get("error", {})
                detail = error_data.get("message") or error_data.get("code") or response.text
            except Exception:
                detail = response.text
            logger.error("Groq HTTP %s: %s", response.status_code, detail[:500])
            return None, f"❌ خطای Groq: {detail[:300]}"

        data = response.json()
        choices = data.get("choices") or []
        if not choices:
            logger.error("Groq returned no choices: %s", data)
            return None, "❌ هوش مصنوعی پاسخی برنگرداند."

        content = ((choices[0].get("message") or {}).get("content") or "").strip()
        if not content:
            return None, "❌ پاسخ هوش مصنوعی خالی بود."
        return await _ensure_persian_output(content), None

    except httpx.TimeoutException:
        logger.exception("Groq timeout")
        return None, "⏳ زمان پاسخ تمام شد؛ دوباره امتحان کن."
    except httpx.RequestError as exc:
        logger.exception("Groq request error: %s", exc)
        return None, "❌ اتصال به هوش مصنوعی برقرار نشد."
    except Exception as exc:
        logger.exception("Groq unexpected error: %s", exc)
        return None, "❌ خطای غیرمنتظره در هوش مصنوعی."


async def ask_ai(message: str, context: ContextTypes.DEFAULT_TYPE,
                 system_prompt: str = None) -> str:
    history = [
        x for x in context.user_data.get("ai_history", [])
        if isinstance(x, dict) and x.get("role") in {"user", "assistant"} and x.get("content")
    ][-10:]

    system_prompt = system_prompt or ASO_SYSTEM_PROMPT
    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(history)
    messages.append({"role": "user", "content": message})

    reply, error = await groq_chat(messages)
    if error:
        return error

    history.extend([
        {"role": "user", "content": message},
        {"role": "assistant", "content": reply},
    ])
    context.user_data["ai_history"] = history[-10:]
    return reply


async def summarize_text(text: str) -> str:
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": [
                        {"role": "system", "content": "You are a professional text summarizer. Detect language (Persian or English). Summarize clearly in the SAME language. Just return the summary."},
                        {"role": "user", "content": text}
                    ],
                    "temperature": 0.3, "max_tokens": 1024
                }
            )
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"Summarize error: {e}")
        return "❌خطا در خلاصه‌سازی متن."


def get_english_teacher_prompt(level: str = "intermediate",
                                mode: str = "conversation") -> str:
    levels = {
        "beginner": "Beginner (A1-A2)",
        "intermediate": "Intermediate (B1-B2)",
        "advanced": "Advanced (C1-C2)",
    }
    level_name = levels.get(level, levels["intermediate"])

    mode_instructions = {
        "conversation": (
            "Have a natural conversation in British English. "
            "Keep replies practical and not too long, then ask one follow-up question. "
            "If there is an important English mistake, briefly correct it after answering."
        ),
        "correct": (
            "Correct the student's sentence. Give: Correct sentence; a short Persian explanation; "
            "one natural British-English alternative; and pronunciation help when useful."
        ),
        "vocabulary": (
            "Teach exactly 3 useful words for this level. For each provide the word, "
            "British IPA pronunciation, Persian meaning, a natural example and a short usage note."
        ),
        "grammar": (
            "Teach one grammar point suitable for this level. Explain it simply in Persian, "
            "give the structure, 3 English examples with Persian meanings, and one short exercise."
        ),
    }

    return (
        "You are Miss Emma, a friendly professional British English teacher.\n"
        f"Student level: {level_name}.\n"
        "Use British English. Be accurate, natural and encouraging. "
        "Use Persian for explanations and English for examples. "
        "Do not invent grammar rules.\n\n"
        + mode_instructions.get(mode, mode_instructions["conversation"])
    )


async def translate_text(text: str) -> str:
    try:
        return await groq_chat([
            {"role":"system","content":"Translate to the other language (Persian↔English). Only return the translation."},
            {"role":"user","content":text},
        ], temperature=0.2, max_tokens=1200, timeout=45.0)
    except Exception as e:
        logger.exception("Translation error: %s", e)
        return "❌ خطا در ترجمه."


async def speech_to_text(file_path: str, language: str = "auto") -> str | None:
    try:
        if language == "ckb":
            async with httpx.AsyncClient(timeout=90.0) as client:
                with open(file_path, "rb") as f:
                    resp = await client.post(
                        "https://www.kurdishtts.com/api/stt-proxy",
                        headers={"x-api-key": KURDISH_STT_KEY},
                        data={"dialect": "sorani"},
                        files={"file": (os.path.basename(file_path), f)}
                    )
                if resp.status_code == 200:
                    return resp.json().get("text", "").strip() or None
        else:
            data = {"model": "whisper-large-v3", "response_format": "json", "temperature": 0.0}
            if language in ("fa", "en"):
                data["language"] = language
            async with httpx.AsyncClient(timeout=90.0) as client:
                with open(file_path, "rb") as f:
                    resp = await client.post(
                        "https://api.groq.com/openai/v1/audio/transcriptions",
                        headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
                        data=data,
                        files={"file": (os.path.basename(file_path), f, "audio/ogg")}
                    )
                if resp.status_code == 200:
                    return resp.json().get("text", "").strip() or None
    except Exception as e:
        logger.error(f"STT Error: {e}")
    return None


async def transcribe_with_segments(file_path: str) -> list:
    try:
        async with httpx.AsyncClient(timeout=180.0) as client:
            with open(file_path, "rb") as f:
                resp = await client.post(
                    "https://api.groq.com/openai/v1/audio/transcriptions",
                    headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
                    data={
                        "model": "whisper-large-v3",
                        "response_format": "verbose_json",
                        "temperature": 0.0
                    },
                    files={"file": (os.path.basename(file_path), f, "audio/ogg")}
                )
            if resp.status_code != 200:
                logger.error(f"Transcription failed: {resp.text}")
                return []
            data = resp.json()
            segments = []
            for seg in data.get("segments", []):
                text = (seg.get("text") or "").strip()
                if text:
                    segments.append({
                        "start": float(seg.get("start", 0)),
                        "end": float(seg.get("end", 0)),
                        "text": text
                    })
            return segments
    except Exception as e:
        logger.error(f"transcribe_with_segments error: {e}")
        return []


async def recognize_music_with_audd(file_path: str) -> dict | None:
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            with open(file_path, "rb") as f:
                resp = await client.post(
                    "https://api.audd.io/",
                    data={"api_token": AUDD_API_KEY, "return": "apple_music,spotify"},
                    files={"file": (os.path.basename(file_path), f)}
                )
            if resp.status_code == 200 and resp.json().get("status") == "success":
                song = resp.json().get("result")
                if song:
                    return {
                        "title": song.get("title") or "نامشخص",
                        "artist": song.get("artist") or "نامشخص",
                        "album": song.get("album") or "",
                        "release_date": song.get("release_date") or ""
                    }
    except Exception as e:
        logger.error(f"AudD Error: {e}")
    return None


async def search_youtube(query: str, max_results: int = 6) -> list:
    try:
        ydl_opts = {"quiet": True, "no_warnings": True, "extract_flat": True, "skip_download": True, "noplaylist": True}
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(f"ytsearch{max_results}:{query}", download=False)
        results = []
        for entry in (info.get("entries") or []):
            if entry and entry.get("id"):
                results.append({
                    "id": entry["id"],
                    "title": entry.get("title") or "بدون عنوان",
                    "url": f"https://www.youtube.com/watch?v={entry['id']}",
                    "duration": entry.get("duration"),
                    "uploader": entry.get("uploader") or entry.get("channel") or ""
                })
        return results
    except Exception as e:
        logger.error(f"YT Search Error: {e}")
        return []


async def get_lyrics(artist: str, title: str) -> str | None:
    try:
        artist = re.sub(r'\(.*?\)|\[.*?\]', '', (artist or "").strip()).strip()
        title = re.sub(r'\(.*?\)|\[.*?\]', '', (title or "").strip()).strip()
        urls = [f"https://api.lyrics.ovh/v1/{artist}/{title}", f"https://api.lyrics.ovh/v1/{title}"] if artist else [f"https://api.lyrics.ovh/v1/{title}"]
        async with httpx.AsyncClient(timeout=12.0) as client:
            for url in urls:
                try:
                    resp = await client.get(url)
                    if resp.status_code == 200:
                        lyrics = resp.json().get("lyrics")
                        if lyrics and len(lyrics.strip()) > 40:
                            return lyrics.strip()
                except Exception:
                    continue
    except Exception:
        pass
    return None


async def short_url(long_url: str) -> str | None:
    try:
        async with httpx.AsyncClient(timeout=12.0) as client:
            resp = await client.get(f"https://is.gd/create.php?format=simple&url={long_url}")
            if resp.status_code == 200 and resp.text.startswith("http"):
                return resp.text.strip()
    except Exception:
        pass
    return None


# =========================================================
# SUBTITLE MODULE
# =========================================================

SORANI_SUBTITLE_PROMPT = """
You are an expert native Central Kurdish (Sorani Kurdish) subtitle translator and editor.
Your job is to translate movie/TV dialogue into accurate, natural, spoken Sorani Kurdish.
Accuracy of meaning is more important than literal word-for-word translation.

STRICT RULES:
1. Output MUST be Central Kurdish (Sorani) in Kurdish Arabic script. NEVER output Persian or Kurmanji.
2. Translate the meaning of the SOURCE line exactly. Do not invent information, remove meaning, or add meaning.
3. Preserve tense, negation, questions, commands, pronouns, relationships, certainty, politeness, sarcasm, anger, fear, humour and emotional tone.
4. Use authentic modern spoken Sorani suitable for a film subtitle, not formal written Kurdish.
5. Prefer correct Sorani vocabulary and grammar over Persian-like wording. Do not replace a Sorani word with a Persian word merely because it sounds familiar.
6. Resolve pronouns and short expressions from the nearby dialogue context when the meaning is clear; otherwise keep the ambiguity instead of guessing.
7. Preserve names, places, brands, numbers and intentional slang when appropriate. Do not translate proper names unnecessarily.
8. Translate idioms by their actual meaning in context, not by their individual words.
9. Do not censor, soften, exaggerate, or morally interpret the dialogue.
10. Do not add explanations, notes, quotes, emojis, markdown, speaker labels or punctuation that is not needed.
11. Return EXACTLY one translated line for every numbered source line.
12. Keep the original numbering exactly and in the same order.
13. Never merge two numbered lines and never split one numbered line into multiple numbered lines.
14. Before answering, silently check every line against the source for missing negation, wrong subject/object, wrong tense, wrong person, and wrong emotional meaning.

Return ONLY the numbered Sorani translations.
"""

PERSIAN_SUBTITLE_PROMPT = """
You are a professional Persian subtitle translator.
Translate the following movie/TV dialogue into natural, fluent, modern spoken Persian (Farsi).

STRICT RULES:
1. Output MUST be Persian (Farsi) using Persian script.
2. Use natural spoken Persian suitable for movie subtitles (not overly formal).
3. Preserve meaning, emotion, attitude and humour.
4. Keep names, places and brands when appropriate.
5. Do not add explanations, quotes, emojis or markdown.
6. Return EXACTLY one translated line for each numbered input line.
7. Keep the original numbering exactly.

Return ONLY the numbered Persian translations.
"""

ENGLISH_SUBTITLE_PROMPT = """
You are a professional English subtitle translator.
Translate the following movie/TV dialogue into natural, fluent, modern spoken English.

STRICT RULES:
1. Output MUST be natural English suitable for movie subtitles.
2. Preserve meaning, emotion, attitude and humour.
3. Keep names, places and brands when appropriate.
4. Do not add explanations, quotes, emojis or markdown.
5. Return EXACTLY one translated line for each numbered input line.
6. Keep the original numbering exactly.

Return ONLY the numbered English translations.
"""


def normalize_text(text: str, lang: str = "fa") -> str:
    if not text:
        return ""
    text = re.sub(r"[*_`#]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+([،؛؟!.,:])", r"\1", text)
    if lang == "ckb":
        replacements = {
            "ي": "ی", "ى": "ی", "ك": "ک", "ۀ": "ە", "ة": "ە",
            "ؤ": "ۆ", "ـ": "", "\u200d": "", "\ufeff": "",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
    return text.strip()


def split_subtitle_line(text: str, max_chars: int = 42) -> list:
    text = text.strip()
    if len(text) <= max_chars:
        return [text]
    words = text.split()
    if len(words) <= 2:
        return [text]

    best_index = 1
    best_diff = float("inf")
    for i in range(1, len(words)):
        left = " ".join(words[:i])
        right = " ".join(words[i:])
        if len(left) <= max_chars and len(right) <= max_chars:
            diff = abs(len(left) - len(right))
            if diff < best_diff:
                best_diff = diff
                best_index = i

    line1 = " ".join(words[:best_index])
    line2 = " ".join(words[best_index:])

    if len(line1) > max_chars or len(line2) > max_chars:
        wrapped = textwrap.wrap(text, width=max_chars, break_long_words=False, break_on_hyphens=False)
        if len(wrapped) <= 2:
            return wrapped
        return [wrapped[0], " ".join(wrapped[1:])]
    return [line1, line2]


def create_srt(segments: list, output_path: str, lang: str = "fa"):
    with open(output_path, "w", encoding="utf-8-sig") as f:
        for number, seg in enumerate(segments, 1):
            lines = split_subtitle_line(seg["text"], max_chars=42)
            subtitle_text = "\n".join(lines)
            start = seconds_to_srt_time(seg["start"])
            end = seconds_to_srt_time(seg["end"])
            f.write(f"{number}\n")
            f.write(f"{start} --> {end}\n")
            f.write(subtitle_text + "\n\n")


def create_ass(segments: list, output_path: str, width: int = 1280, height: int = 720, lang: str = "fa"):
    if height <= 720:
        font_size = 34
    elif height <= 1080:
        font_size = 46
    else:
        font_size = 56

    if lang == "ckb":
        font_name = "Noto Naskh Arabic"
    elif lang == "fa":
        font_name = "Noto Naskh Arabic"
    else:
        font_name = "Arial"

    header = f"""[Script Info]
ScriptType: v4.00+
PlayResX: {width}
PlayResY: {height}
WrapStyle: 2
ScaledBorderAndShadow: yes
YCbCr Matrix: None

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,{font_name},{font_size},&H00FFFFFF,&H00FFFFFF,&H00000000,&H80000000,0,0,0,0,100,100,0,0,1,3.0,1.2,2,50,50,40,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""

    def ass_time(seconds: float) -> str:
        seconds = max(0.0, float(seconds))
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        cs = int(round((seconds - int(seconds)) * 100))
        if cs >= 100:
            s += 1
            cs = 0
        return f"{h}:{m:02d}:{s:02d}.{cs:02d}"

    with open(output_path, "w", encoding="utf-8-sig") as f:
        f.write(header)
        for seg in segments:
            lines = split_subtitle_line(seg["text"], max_chars=42)
            text = r"\N".join(lines)
            text = text.replace("{", r"\{").replace("}", r"\}")
            f.write(
                f"Dialogue: 0,"
                f"{ass_time(seg['start'])},"
                f"{ass_time(seg['end'])},"
                f"Default,,0,0,0,,"
                f"{text}\n"
            )


async def correct_transcript(segments: list) -> list:
    """Lightly clean ASR output without translating, paraphrasing or changing timing."""
    if not segments:
        return segments
    try:
        lines = [f"{i+1}. {seg['text']}" for i, seg in enumerate(segments)]
        text_block = "\n".join(lines)
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": [
                        {
                            "role": "system",
                            "content": (
                                "You are a highly conservative speech-to-text transcript editor. "
                                "Correct ONLY clear ASR recognition errors that are strongly supported by the surrounding dialogue. "
                                "Do NOT translate, paraphrase, shorten, expand, rewrite, or stylistically improve the dialogue. "
                                "Preserve the original language, wording, slang, names, numbers, emotion, tense, negation and meaning whenever possible. "
                                "If uncertain, keep the original text unchanged. "
                                "Keep exactly the same number of lines and numbering. "
                                "Return only corrected numbered lines."
                            )
                        },
                        {"role": "user", "content": text_block}
                    ],
                    "temperature": 0.0,
                    "max_tokens": 3072
                }
            )
            resp.raise_for_status()
            corrected = resp.json()["choices"][0]["message"]["content"].strip()

        parsed = {}
        for line in corrected.splitlines():
            line = line.strip()
            match = re.match(r"^(\d+)\.\s*(.+)$", line)
            if match:
                idx = int(match.group(1)) - 1
                if 0 <= idx < len(segments):
                    parsed[idx] = match.group(2).strip()

        if len(parsed) == len(segments) and set(parsed) == set(range(len(segments))):
            return [
                {
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": parsed[i] or seg["text"]
                }
                for i, seg in enumerate(segments)
            ]
    except Exception as e:
        logger.warning(f"Transcript correction failed: {e}")
    return segments


async def _translate_batch_once(batch: list, system_prompt: str, target_lang: str) -> dict:
    """Translate one subtitle batch and return {1-based-number: text}."""
    lines = [f"{j+1}. {seg['text']}" for j, seg in enumerate(batch)]
    text_block = "\n".join(lines)

    user_prompt = (
        "Translate the numbered SOURCE dialogue below. The lines are consecutive movie dialogue, "
        "so use their context to resolve short pronouns, references and idioms. "
        "Do not invent information. Do not merge or split lines. "
        "Return exactly one numbered translation for every input line and nothing else.\n\n"
        "SOURCE DIALOGUE:\n" + text_block
    )

    async with httpx.AsyncClient(timeout=90.0) as client:
        resp = await client.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={
                "model": GROQ_CHAT_MODEL,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 3072
            }
        )
        resp.raise_for_status()
        translated_block = resp.json()["choices"][0]["message"]["content"].strip()

    translated_lines = {}
    for line in translated_block.splitlines():
        line = line.strip()
        match = re.match(r"^(\d+)\.\s*(.+)$", line)
        if match:
            translated_lines[int(match.group(1))] = normalize_text(match.group(2).strip(), target_lang)
    return translated_lines


async def _review_sorani_batch(batch: list, translated_lines: dict) -> dict:
    """Second-pass Sorani QA: compare source and translation and fix only meaning errors."""
    if not translated_lines or len(translated_lines) != len(batch):
        return translated_lines

    source_block = "\n".join(f"{i+1}. {seg['text']}" for i, seg in enumerate(batch))
    translation_block = "\n".join(f"{i+1}. {translated_lines[i+1]}" for i in range(len(batch)))

    system_prompt = """
You are a senior native Central Kurdish (Sorani) subtitle proofreader.
Review the proposed Sorani translations against the original dialogue.

STRICT RULES:
1. Keep Central Kurdish (Sorani), Kurdish Arabic script only. Never Persian or Kurmanji.
2. Fix ONLY genuine translation errors: wrong meaning, missing/added meaning, wrong negation, wrong tense/person, wrong subject/object, wrong pronoun/reference, incorrect idiom, or unnatural wording that changes meaning.
3. Preserve correct translations exactly when they are already accurate.
4. Do not invent context or information that is not supported by the source dialogue.
5. Keep names, numbers, slang and emotional tone appropriate to the source.
6. Keep exactly one output line per numbered input line and keep numbering unchanged.
7. Return ONLY the corrected numbered Sorani translations.
""".strip()

    user_prompt = (
        "ORIGINAL SOURCE:\n" + source_block +
        "\n\nPROPOSED SORANI TRANSLATION:\n" + translation_block +
        "\n\nCheck every line carefully and return the final Sorani subtitles only."
    )

    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            resp = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.0,
                    "max_tokens": 3072
                }
            )
            resp.raise_for_status()
            reviewed = resp.json()["choices"][0]["message"]["content"].strip()

        parsed = {}
        for line in reviewed.splitlines():
            line = line.strip()
            match = re.match(r"^(\d+)\.\s*(.+)$", line)
            if match:
                idx = int(match.group(1))
                if 1 <= idx <= len(batch):
                    parsed[idx] = normalize_text(match.group(2).strip(), "ckb")
        if len(parsed) == len(batch):
            return parsed
    except Exception as e:
        logger.warning(f"Sorani subtitle QA failed: {e}")
    return translated_lines


async def batch_translate(segments: list, target_lang: str, batch_size: int = 6) -> list:
    """Translate subtitles in small context-aware batches; Sorani gets a second QA pass."""
    if target_lang == "ckb":
        system_prompt = SORANI_SUBTITLE_PROMPT
    elif target_lang == "fa":
        system_prompt = PERSIAN_SUBTITLE_PROMPT
    else:
        system_prompt = ENGLISH_SUBTITLE_PROMPT

    result = []
    for i in range(0, len(segments), batch_size):
        batch = segments[i:i + batch_size]
        translated_lines = {}

        try:
            translated_lines = await _translate_batch_once(batch, system_prompt, target_lang)

            # If the model omitted a line, retry the whole small batch once rather than
            # silently treating the source text as a translation.
            if len(translated_lines) != len(batch):
                await asyncio.sleep(0.2)
                translated_lines = await _translate_batch_once(batch, system_prompt, target_lang)

            if target_lang == "ckb" and len(translated_lines) == len(batch):
                translated_lines = await _review_sorani_batch(batch, translated_lines)

            for j, seg in enumerate(batch):
                translated = translated_lines.get(j + 1)
                if not translated:
                    # Safe fallback: keep the original line instead of inventing a translation.
                    translated = normalize_text(seg["text"], target_lang)
                result.append({
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": translated
                })

        except Exception as e:
            logger.error(f"Batch translate error: {e}")
            for seg in batch:
                result.append({
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": normalize_text(seg["text"], target_lang)
                })

        await asyncio.sleep(0.35)

    return result


async def process_subtitle(video_path: str, status_message, target_lang: str = "ckb") -> dict:
    async def _status(msg):
        try:
            await status_message.edit_text(msg)
        except Exception:
            pass
    temp_dir = tempfile.mkdtemp(prefix="sub_")
    audio_path = os.path.join(temp_dir, "audio.ogg")
    srt_path = os.path.join(temp_dir, "subtitle.srt")
    ass_path = os.path.join(temp_dir, "subtitle.ass")
    output_video = os.path.join(temp_dir, "output.mp4")
    trimmed_video = os.path.join(temp_dir, "trimmed.mp4")

    lang_names = {"ckb": "☀️ کوردی", "fa": "🇮🇷 فارسی", "en": "🇬🇧 انگلیسی"}

    try:
        info = get_media_info(video_path)
        duration = info.get("duration") or 0
        width = info.get("width") or 1280
        height = info.get("height") or 720
        work_video = video_path

        if duration > MAX_SUBTITLE_VIDEO_DURATION:
            await _status(f"⏳ویدیو طولانی است. فقط {MAX_SUBTITLE_VIDEO_DURATION} ثانیه اول پردازش می‌شود...")
            proc = await asyncio.create_subprocess_exec(
                "ffmpeg", "-y", "-i", video_path, "-t", str(MAX_SUBTITLE_VIDEO_DURATION),
                "-c", "copy", trimmed_video,
                stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.DEVNULL
            )
            await proc.wait()
            if os.path.exists(trimmed_video):
                work_video = trimmed_video

        await _status("⏳۱/۶ استخراج صدا...")
        proc = await asyncio.create_subprocess_exec(
            "ffmpeg", "-y", "-i", work_video, "-vn", "-acodec", "libopus", "-b:a", "64k", audio_path,
            stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.DEVNULL
        )
        await proc.wait()
        if not os.path.exists(audio_path):
            raise RuntimeError("استخراج صدا ناموفق بود")

        await _status("⏳ ۲/۶ تبدیل گفتار به متن...")
        segments = await transcribe_with_segments(audio_path)
        if not segments:
            full_text = await speech_to_text(audio_path, "auto")
            if not full_text:
                raise RuntimeError("متن استخراج نشد")
            dur = get_media_info(work_video).get("duration") or 60
            segments = [{"start": 0.0, "end": float(dur), "text": full_text}]

        await _status("⏳۳/۶ اصلاح متن...")
        segments = await correct_transcript(segments)

        await _status(f"⏳۴/۶ ترجمه به {lang_names.get(target_lang, target_lang)}...")
        translated = await batch_translate(segments, target_lang, batch_size=8)

        await _status("⏳ ۵/۶ ساخت فایل زیرنویس...")
        create_srt(translated, srt_path, lang=target_lang)
        create_ass(translated, ass_path, width=width, height=height, lang=target_lang)

        burned_video = None
        try:
            await _status("⏳۶/۶ چسباندن زیرنویس روی ویدیو...")

            cmd = [
                "ffmpeg", "-y", "-i", work_video,
                "-vf", f"ass={ass_path}",
                "-c:a", "copy",
                "-c:v", "libx264",
                "-preset", "veryfast",
                "-crf", "23",
                output_video
            ]
            proc = await asyncio.create_subprocess_exec(
                *cmd, stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.PIPE
            )
            try:
                _, stderr = await asyncio.wait_for(proc.communicate(), timeout=150)
                if proc.returncode == 0 and os.path.exists(output_video) and os.path.getsize(output_video) > 1000:
                    burned_video = output_video
            except asyncio.TimeoutError:
                proc.kill()

            if not burned_video:
                style = (
                    "FontSize=20,"
                    "PrimaryColour=&H00FFFFFF,"
                    "OutlineColour=&H00000000,"
                    "BackColour=&H80000000,"
                    "BorderStyle=3,"
                    "Outline=2,"
                    "Shadow=1,"
                    "MarginV=30,"
                    "Alignment=2"
                )
                cmd = [
                    "ffmpeg", "-y", "-i", work_video,
                    "-vf", f"subtitles={srt_path}:force_style='{style}'",
                    "-c:a", "copy",
                    "-c:v", "libx264",
                    "-preset", "veryfast",
                    "-crf", "23",
                    output_video
                ]
                proc = await asyncio.create_subprocess_exec(
                    *cmd, stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.PIPE
                )
                try:
                    _, stderr = await asyncio.wait_for(proc.communicate(), timeout=150)
                    if proc.returncode == 0 and os.path.exists(output_video) and os.path.getsize(output_video) > 1000:
                        burned_video = output_video
                except asyncio.TimeoutError:
                    proc.kill()

        except Exception as e:
            logger.warning(f"Burning subtitles failed: {e}")

        await _status(f"✅ زیرنویس {lang_names.get(target_lang)} آماده شد!")

        return {
            "srt_path": srt_path,
            "video_path": burned_video,
            "temp_dir": temp_dir,
            "original_video": video_path,
            "lang": target_lang
        }

    except Exception as e:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
        raise e


# =========================================================
# DOWNLOAD
# =========================================================

def _cookie_env_names(platform: str):
    """Return cookie env names in priority order for a platform."""
    key = {
        "YouTube": "YOUTUBE",
        "Instagram": "INSTAGRAM",
        "TikTok": "TIKTOK",
        "Facebook": "FACEBOOK",
        "Twitter": "TWITTER",
        "Reddit": "REDDIT",
        "Aparat": "APARAT",
    }.get(platform, "")
    names = []
    if key:
        names += [f"{key}_COOKIES_B64", f"{key}_COOKIES_TEXT"]
    names += ["YT_DLP_COOKIES_B64", "YT_DLP_COOKIES_TEXT"]
    return names


def _validate_netscape_cookies(raw: str, platform: str) -> str:
    """Validate Netscape cookie text; return a short warning string (empty if OK)."""
    lines = [ln for ln in raw.splitlines() if ln.strip() and not ln.strip().startswith("#")]
    if not lines:
        return "فایل کوکی خالی است (هیچ ردیف کوکی ندارد)."
    names = set()
    for ln in lines:
        parts = ln.split("\t")
        if len(parts) >= 6:
            names.add(parts[5].strip())
    if platform == "Instagram":
        missing = [k for k in ("sessionid",) if k not in names]
        if missing:
            return (
                "کوکی اینستاگرام ناقص است؛ این کلیدها پیدا نشد: "
                + ", ".join(missing)
                + ". از مرورگر لاگین‌شده دوباره export بگیر."
            )
        if "sessionid" in names and "ds_user_id" not in names:
            return "هشدار: ds_user_id در کوکی نیست؛ ممکن است session ناقص باشد."
    if platform == "YouTube":
        if not any(n.startswith("LOGIN_INFO") or n in {"SID", "HSID", "SSID", "APISID", "SAPISID", "__Secure-1PSID"} for n in names):
            return "کوکی یوتیوب شبیه حساب لاگین‌شده نیست (SID/LOGIN_INFO پیدا نشد)."
    return ""


def _decode_cookie_payload(payload: str, is_b64: bool) -> str:
    """Decode env cookie value; tolerate whitespace and accidental double-encoding."""
    payload = (payload or "").strip()
    # Render UI sometimes inserts newlines into long secrets
    compact = "".join(payload.split())
    if is_b64:
        try:
            raw = base64.b64decode(compact, validate=False).decode("utf-8", errors="replace")
        except Exception as exc:
            raise RuntimeError(
                "کوکی Base64 معتبر نیست. خروجی `base64 -w0 cookies.txt` را کامل در "
                "INSTAGRAM_COOKIES_B64 (یا YT_DLP_COOKIES_B64) بگذار."
            ) from exc
        # If user pasted plain text into *_B64 by mistake, accept Netscape text
        if "sessionid" not in raw and "Netscape" not in raw and payload.lstrip().startswith("#"):
            raw = payload
        return raw
    return payload


def _normalize_netscape_cookie_text(raw: str) -> str:
    """Ensure a usable Netscape cookie file (header + tab-separated rows)."""
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = raw.split("\n")
    out = ["# Netscape HTTP Cookie File", "# This file was generated by AsoLand"]
    for ln in lines:
        s = ln.strip()
        if not s or s.startswith("#"):
            continue
        # Some exporters use spaces; Netscape needs tabs between fields
        if "\t" not in s and " " in s:
            # best-effort: collapse multiple spaces only when it looks like 7 fields
            parts = s.split()
            if len(parts) >= 7:
                # domain flag path secure expiry name value(possibly with spaces joined)
                s = "\t".join(parts[:6] + [" ".join(parts[6:])])
        # HttpOnly prefix used by some browsers
        if s.startswith("#HttpOnly_"):
            s = s[len("#HttpOnly_"):]
        out.append(s)
    return "\n".join(out) + "\n"


def _prepare_cookie_file(platform: str):
    """Create a temporary Netscape cookie file from a Render secret if supplied.

    Returns (path, temporary_path). The caller must remove temporary_path.
    """
    configured_path = os.getenv("YT_DLP_COOKIES", "").strip()
    if configured_path and os.path.isfile(configured_path):
        return configured_path, None
    if os.path.isfile(COOKIES_FILE):
        return COOKIES_FILE, None

    payload = None
    is_b64 = False
    for name in _cookie_env_names(platform):
        value = os.getenv(name, "")
        if value and value.strip():
            payload = value
            is_b64 = name.endswith("_B64")
            break
    if not payload:
        return "", None

    raw = _decode_cookie_payload(payload, is_b64)
    raw = _normalize_netscape_cookie_text(raw)
    if "sessionid" not in raw and platform == "Instagram":
        # Also accept if original had sessionid before normalize edge-cases
        pass
    warn = _validate_netscape_cookies(raw, platform)
    if warn and ("ناقص" in warn or "خالی" in warn):
        raise RuntimeError(warn)

    fd, path = tempfile.mkstemp(prefix="asoland-cookies-", suffix=".txt")
    try:
        with os.fdopen(fd, "w", encoding="utf-8", newline="\n") as f:
            f.write(raw)
        try:
            os.chmod(path, 0o600)
        except OSError:
            pass
    except Exception:
        try:
            os.unlink(path)
        except OSError:
            pass
        raise
    return path, path


def _friendly_download_error(exc: Exception, platform: str) -> str:
    """Map common yt-dlp / network failures to clear Persian messages for the Mini App."""
    message = str(exc or "")
    lower = message.lower()

    if "sign in to confirm" in lower or "confirm you're not a bot" in lower or "confirm you are not a bot" in lower:
        return (
            "یوتیوب درخواست ورود (Sign-in) کرده است. "
            "روی سرور Render (IP دیتاسنتر) بدون کوکی معتبر کار نمی‌کند. "
            "در Environment Variables مقدار YOUTUBE_COOKIES_B64 یا YT_DLP_COOKIES_B64 را تنظیم کن "
            "(فایل کوکی Netscape از مرورگر خودت)."
        )
    # Instagram: redirected to login / anonymous rate-limit (even when cookie env is set,
    # the cookie may be expired, incomplete, or from a different account).
    if (
        "redirected to the login page" in lower
        or "rate-limit for accessing posts anonymously" in lower
        or ("instagram" in lower and "login page" in lower)
        or "accessing posts anonymously" in lower
    ):
        return (
            "اینستاگرام دسترسی ناشناس را بسته (login / rate-limit). "
            "yt-dlp گفته anonymously یعنی کوکی اعمال نشده یا منقضی است. "
            "۱) Logout/Login در اینستاگرام "
            "۲) Get cookies.txt LOCALLY → export "
            "۳) مطمئن شو sessionid و csrftoken داخل فایل است "
            "۴) base64 -w0 cookies.txt → INSTAGRAM_COOKIES_B64 "
            "۵) Redeploy. اگر باز هم همین بود → DOWNLOAD_PROXY اجباری است."
        )
    if "http error 403" in lower or "403: forbidden" in lower:
        return (
            f"دسترسی رد شد (HTTP 403) برای {platform or 'این سرویس'}. "
            "معمولاً به خاطر IP سرور یا نیاز به کوکی است. "
            "کوکی معتبر (YOUTUBE_COOKIES_B64 / INSTAGRAM_COOKIES_B64 / ...) یا DOWNLOAD_PROXY تنظیم کن."
        )
    if "429" in lower or "too many requests" in lower or "rate-limit" in lower or "rate limited" in lower:
        return (
            "سرویس مقصد درخواست‌های سرور را محدود کرده است (HTTP 429 / rate-limit). "
            "چند دقیقه صبر کن. اگر مکرر شد: کوکی را تازه کن یا DOWNLOAD_PROXY بگذار "
            "(IP رایگان Render اغلب بلاک می‌شود)."
        )
    if "private video" in lower or "this video is private" in lower or "sorry, this page isn't available" in lower:
        return "این ویدیو خصوصی است، حذف شده، یا برای این حساب قابل مشاهده نیست."
    if "video unavailable" in lower or "has been removed" in lower or "not available" in lower:
        return "ویدیو در دسترس نیست یا حذف شده است."
    if "age-restricted" in lower or "age restricted" in lower or "sign in to confirm your age" in lower:
        return (
            "ویدیو محدودیت سنی دارد. برای دانلود باید کوکی حساب لاگین‌شده در Environment Variable تنظیم شود."
        )
    if "login required" in lower or "please log in" in lower or ("cookies" in lower and "required" in lower):
        return (
            f"این لینک ({platform or 'پلتفرم'}) نیاز به ورود دارد. "
            "کوکی معتبر را در Environment Variable مربوطه (مثلاً INSTAGRAM_COOKIES_B64) قرار بده."
        )
    if "unsupported url" in lower or "no video formats" in lower or "requested format is not available" in lower:
        return "فرمت یا لینک پشتیبانی نمی‌شود. لینک مستقیم ویدیو را امتحان کن یا کیفیت پایین‌تر انتخاب کن."
    if "ffmpeg" in lower and ("not found" in lower or "error" in lower):
        return "خطا در ffmpeg هنگام ادغام ویدیو/صدا. روی سرور ffmpeg نصب باشد (در Dockerfile هست)."
    if "timed out" in lower or "timeout" in lower or "connection reset" in lower:
        return "اتصال قطع شد یا زمان تمام شد. لینک کوتاه‌تر/کیفیت پایین‌تر را امتحان کن یا دوباره تلاش کن."

    # Keep a readable but not too long original message
    cleaned = message.replace("ERROR: ", "").strip()
    if len(cleaned) > 280:
        cleaned = cleaned[:277] + "..."
    return cleaned or "خطای ناشناخته در دانلود"


def _pick_downloaded_file(output_dir, prefer_audio=False):
    """Return the newest media file under output_dir (recursive)."""
    exts_video = {".mp4", ".mkv", ".webm", ".mov", ".m4v"}
    exts_audio = {".mp3", ".m4a", ".aac", ".opus", ".ogg", ".wav"}
    exts = (exts_audio | exts_video) if prefer_audio else (exts_video | exts_audio)
    candidates = []
    root = Path(output_dir)
    if not root.exists():
        return None
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in exts:
            candidates.append(p)
    if not candidates:
        return None
    if prefer_audio:
        audio_only = [p for p in candidates if p.suffix.lower() in exts_audio]
        if audio_only:
            candidates = audio_only
    return str(max(candidates, key=lambda x: x.stat().st_mtime))


def _download_with_gallery_dl(url, output_dir, platform, cookie_path="", proxy="", quality="720"):
    """Fallback downloader using gallery-dl (strong for Instagram / Twitter / Reddit)."""
    try:
        import gallery_dl  # noqa: F401
    except ImportError as exc:
        raise RuntimeError("gallery-dl نصب نیست") from exc

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    dest = out / "gdl"
    dest.mkdir(exist_ok=True)

    cmd = [
        "gallery-dl",
        "--destination", str(dest),
        "--filename", "{id}.{extension}",
        "--no-mtime",
        "--quiet",
    ]
    if cookie_path and os.path.isfile(cookie_path):
        cmd += ["--cookies", cookie_path]
    if proxy:
        cmd += ["--proxy", proxy]

    cmd.append(url)

    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=180,
        env={**os.environ, "PYTHONUNBUFFERED": "1"},
    )
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "").strip() or f"gallery-dl exit {proc.returncode}"
        raise RuntimeError(err[:500])

    prefer_audio = quality == "audio"
    filename = _pick_downloaded_file(str(dest), prefer_audio=prefer_audio)
    if not filename:
        # sometimes gallery-dl nests deeper
        filename = _pick_downloaded_file(str(out), prefer_audio=prefer_audio)
    if not filename:
        raise RuntimeError("gallery-dl فایلی ذخیره نکرد")

    # Convert to mp3 if user asked for audio only
    if quality == "audio" and not filename.lower().endswith(".mp3"):
        mp3 = str(Path(filename).with_suffix(".mp3"))
        conv = subprocess.run(
            ["ffmpeg", "-y", "-i", filename, "-vn", "-acodec", "libmp3lame", "-q:a", "2", mp3],
            capture_output=True,
            timeout=120,
        )
        if conv.returncode == 0 and os.path.exists(mp3):
            try:
                os.unlink(filename)
            except OSError:
                pass
            filename = mp3

    title = Path(filename).stem
    return {
        "filename": filename,
        "title": clean_title(title),
        "is_audio": quality == "audio",
        "info": {"extractor": "gallery-dl", "id": Path(filename).stem},
        "platform": platform,
        "backend": "gallery-dl",
    }


def _download_with_ytdlp(url, output_dir, platform, quality="720", progress_callback=None,
                         cancel_check=None, download_subs=False, cookie_path=""):
    """Primary downloader using yt-dlp."""
    output = os.path.join(output_dir, "%(id)s.%(ext)s")
    last_update = [0]
    info = None
    filename = None

    def progress_hook(data):
        if cancel_check and cancel_check():
            raise Exception("دانلود توسط کاربر لغو شد")
        if data.get("status") != "downloading":
            return
        percent = data.get("_percent_str", "0%").strip()
        if time.time() - last_update[0] < 1.5:
            return
        last_update[0] = time.time()
        if progress_callback:
            progress_callback(percent)

    user_agent = os.getenv(
        "YT_DLP_USER_AGENT",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
    )
    ydl_opts = {
        "outtmpl": output,
        "noplaylist": True,
        "quiet": True,
        "no_warnings": True,
        "retries": 3,
        "extractor_retries": 3,
        "fragment_retries": 5,
        "socket_timeout": 60,
        "concurrent_fragment_downloads": 1,
        "http_chunk_size": 5 * 1024 * 1024,
        "http_headers": {
            "User-Agent": user_agent,
            "Accept-Language": "en-US,en;q=0.9",
        },
        "progress_hooks": [progress_hook],
    }

    if cookie_path:
        ydl_opts["cookiefile"] = cookie_path

    proxy = os.getenv("DOWNLOAD_PROXY", "").strip()
    if proxy:
        ydl_opts["proxy"] = proxy

    impersonate_target = os.getenv("YT_DLP_IMPERSONATE_TARGET", "").strip()
    if impersonate_target:
        try:
            import curl_cffi  # noqa: F401
            ydl_opts["impersonate"] = impersonate_target
        except Exception:
            pass

    if platform == "Instagram":
        ydl_opts["http_headers"].update({
            "Referer": "https://www.instagram.com/",
            "Origin": "https://www.instagram.com",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9,fa;q=0.8",
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "same-origin",
        })
        ydl_opts["retries"] = 5
        ydl_opts["extractor_retries"] = 5
        ydl_opts["sleep_interval_requests"] = 1
    elif platform == "TikTok":
        ydl_opts["http_headers"]["Referer"] = "https://www.tiktok.com/"
    elif platform == "YouTube":
        ydl_opts["http_headers"]["Referer"] = "https://www.youtube.com/"
        po_token = os.getenv("YT_DLP_YOUTUBE_PO_TOKEN", "").strip()
        clients_env = os.getenv("YT_DLP_YOUTUBE_PLAYER_CLIENT", "").strip()
        if clients_env:
            player_clients = [c.strip() for c in clients_env.split(",") if c.strip()]
        else:
            player_clients = ["android", "ios", "mweb", "web"]
        youtube_args = {"player_client": player_clients}
        if po_token:
            youtube_args["po_token"] = [f"{player_clients[0]}+{po_token}"]
        ydl_opts["extractor_args"] = {"youtube": youtube_args}

    if download_subs:
        ydl_opts.update({
            "writesubtitles": True,
            "writeautomaticsub": True,
            "subtitleslangs": ["fa", "en"],
            "skip_download": True,
        })
    elif quality == "audio":
        ydl_opts.update({
            "format": "bestaudio/best",
            "postprocessors": [{"key": "FFmpegExtractAudio", "preferredcodec": "mp3", "preferredquality": "192"}],
        })
    else:
        h = {"360": 360, "720": 720, "1080": 1080}.get(str(quality), 720)
        ydl_opts.update({
            "format": (
                f"bestvideo[height<={h}][ext=mp4]+bestaudio[ext=m4a]/"
                f"bestvideo[height<={h}]+bestaudio/"
                f"best[height<={h}][ext=mp4]/best[height<={h}]/best"
            ),
            "merge_output_format": "mp4",
        })

    last_error = None
    attempts = [dict(ydl_opts)]
    if quality not in ("audio",) and not download_subs:
        simple = dict(ydl_opts)
        simple["format"] = "best[ext=mp4]/best"
        simple.pop("merge_output_format", None)
        attempts.append(simple)

    for attempt_idx, opts in enumerate(attempts):
        for rate_try in range(3):
            try:
                with yt_dlp.YoutubeDL(opts) as ydl:
                    info = ydl.extract_info(url, download=True)
                    filename = ydl.prepare_filename(info)
                last_error = None
                break
            except Exception as exc:
                last_error = exc
                message = str(exc)
                lower = message.lower()
                if "429" in lower or "too many requests" in lower:
                    if rate_try >= 2:
                        raise RuntimeError(_friendly_download_error(exc, platform)) from exc
                    time.sleep(5 * (rate_try + 1))
                    continue
                if (
                    attempt_idx < len(attempts) - 1
                    and (
                        "requested format is not available" in lower
                        or "no video formats" in lower
                        or "format is not available" in lower
                    )
                ):
                    break
                raise RuntimeError(_friendly_download_error(exc, platform)) from exc
        if last_error is None:
            break
    if last_error is not None:
        raise RuntimeError(_friendly_download_error(last_error, platform)) from last_error

    if quality == "audio":
        mp3 = os.path.splitext(filename)[0] + ".mp3"
        if os.path.exists(mp3):
            filename = mp3
    elif not download_subs and not os.path.exists(filename):
        mp4 = os.path.splitext(filename)[0] + ".mp4"
        if os.path.exists(mp4):
            filename = mp4

    if not os.path.exists(filename):
        picked = _pick_downloaded_file(output_dir, prefer_audio=(quality == "audio"))
        if picked:
            filename = picked

    if not os.path.exists(filename):
        raise RuntimeError("فایل خروجی ساخته نشد")

    return {
        "filename": filename,
        "title": clean_title(info.get("title") if info else None),
        "is_audio": quality == "audio",
        "info": info,
        "platform": platform,
        "backend": "yt-dlp",
    }


# Platforms where gallery-dl is a useful second backend
_GALLERY_DL_PLATFORMS = {"Instagram", "Twitter / X", "Reddit", "TikTok", "Facebook"}


def download_media(url, output_dir, quality="best", progress_callback=None, cancel_check=None, download_subs=False):
    """Download media with multi-backend strategy.

    1) yt-dlp (primary) — best for YouTube, also works for others
    2) gallery-dl (fallback) — strong for Instagram / Twitter / Reddit when yt-dlp is rate-limited

    Authentication is supplied only through server-side secrets/files.
    """
    url = normalize_media_url(url) or url
    platform = get_platform(url)
    cookie_temp = None
    errors = []

    try:
        cookie_path, cookie_temp = _prepare_cookie_file(platform)
        proxy = os.getenv("DOWNLOAD_PROXY", "").strip()

        # Instagram on datacenter IPs almost always needs cookies. Fail early
        # with a clear message instead of an opaque yt-dlp anonymous error.
        if platform == "Instagram" and not cookie_path and not proxy:
            raise RuntimeError(
                "اینستاگرام روی سرور Render بدون کوکی کار نمی‌کند. "
                "INSTAGRAM_COOKIES_B64 را با cookies.txt تازه (شامل sessionid و csrftoken) "
                "تنظیم کن، یا DOWNLOAD_PROXY بگذار."
            )

        # --- Backend 1: yt-dlp ---
        try:
            result = _download_with_ytdlp(
                url, output_dir, platform, quality=quality,
                progress_callback=progress_callback, cancel_check=cancel_check,
                download_subs=download_subs, cookie_path=cookie_path or "",
            )
            return result
        except Exception as exc:
            errors.append(f"yt-dlp: {exc}")
            # Subtitles path has no gallery-dl equivalent worth keeping
            if download_subs:
                raise

        # --- Backend 2: gallery-dl (Instagram, Twitter, Reddit, TikTok, Facebook) ---
        if platform in _GALLERY_DL_PLATFORMS:
            try:
                # Small pause after rate-limit before second backend
                time.sleep(2)
                result = _download_with_gallery_dl(
                    url, output_dir, platform,
                    cookie_path=cookie_path or "",
                    proxy=proxy,
                    quality=quality,
                )
                return result
            except Exception as exc:
                errors.append(f"gallery-dl: {exc}")

        # Both failed — show a clear combined message
        def _strip_prefix(s: str) -> str:
            for prefix in ("yt-dlp: ", "gallery-dl: "):
                if s.startswith(prefix):
                    return s[len(prefix):]
            return s
        parts = [_strip_prefix(e) for e in errors]
        if len(parts) >= 2:
            msg = (
                "هر دو موتور دانلود شکست خوردند.\n"
                f"• yt-dlp: {parts[0][:220]}\n"
                f"• gallery-dl: {parts[1][:220]}\n"
                "راه‌حل: کوکی تازه (sessionid+csrftoken) یا DOWNLOAD_PROXY با IP معمولی."
            )
        elif parts:
            msg = parts[0]
        else:
            msg = "دانلود ناموفق"
        raise RuntimeError(msg[:700])

    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(_friendly_download_error(exc, platform)) from exc
    finally:
        if cookie_temp:
            try:
                os.unlink(cookie_temp)
            except OSError:
                pass


async def upload_file(chat_id, filename, caption, context, is_audio=False):
    with open(filename, "rb") as f:
        if is_audio:
            await context.bot.send_audio(chat_id=chat_id, audio=f, caption=tr(caption, get_user_lang(chat_id)), title=os.path.basename(filename), read_timeout=120, write_timeout=120)
        elif os.path.getsize(filename) <= MAX_FILE_SIZE:
            await context.bot.send_video(chat_id=chat_id, video=f, caption=tr(caption, get_user_lang(chat_id)), supports_streaming=True, read_timeout=120, write_timeout=120)
        else:
            await context.bot.send_document(chat_id=chat_id, document=f, caption=tr(caption, get_user_lang(chat_id)) + "\n\n" + tr("📦 فایل حجیم", get_user_lang(chat_id)), read_timeout=180, write_timeout=180)


# =========================================================
# HANDLERS
# =========================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    saved_lang = get_user_lang(update.effective_user.id)
    context.user_data.clear()
    context.user_data["language"] = saved_lang
    _CURRENT_UI_LANG.set(saved_lang)
    register_user(update.effective_user.id)
    name = update.effective_user.first_name or "کاربر"
    text = (
        f"سلام <b>{html.escape(name)}</b> 👋\n\n"
        f"{get_jalali_datetime()}\n\n"
        "🌟 <b>AsoLand</b>\n"
        "ربات همه‌کاره و هوشمند\n\n"
        "از منوی زیر استفاده کن."
    )
    await reply_text_localized(update.message, context, text, parse_mode="HTML", reply_markup=get_reply_keyboard())
    await reply_text_localized(update.message, context, "منوی اصلی:", reply_markup=get_main_keyboard())


async def admin_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    _CURRENT_UI_LANG.set(get_user_lang(update.effective_user.id, context))
    if not is_admin(update.effective_user.id):
        await reply_text_localized(update.message, context, "⛔️ دسترسی ندارید.")
        return
    register_user(update.effective_user.id)
    await reply_text_localized(update.message, context, 
        "🛡 <b>پنل مدیریت</b>\n\nآمار، پیام همگانی و بلاک‌لیست از دکمه‌های زیر:",
        parse_mode="HTML",
        reply_markup=get_admin_keyboard()
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    lang = get_user_lang(user_id, context)
    _CURRENT_UI_LANG.set(lang)
    help_texts = {
        "fa": (
            "🌟 <b>راهنمای AsoLand</b>\n\n"
            "📥 دانلود فیلم و ویدیو از اینستاگرام، یوتیوب، تیک‌تاک و ...\n"
            "🎵 موسیقی، پیدا کردن آهنگ و متن ترانه\n"
            "💰 قیمت دلار و سکه + تبدیل ارز\n"
            "🚀 امکانات هوشمند: آب‌وهوا، فال روزانه، محاسبه‌گر، نمودار، هشدار قیمت، اخبار و ابزار دانشجویی\n"
            "🛠 ابزارهای کاربردی: استیکر، QR، فشرده‌سازی، PDF، ورد و زیباسازی متن\n"
            "🤖 هوش مصنوعی: چت، ترجمه، خلاصه‌سازی، ویس به متن و معلم انگلیسی\n"
            "🎬 ساخت زیرنویس (کوردی / فارسی / انگلیسی)\n"
            "🩷 کانفیگ رایگان\n\n"
            f"📊 محدودیت روزانه — دانلود: {DAILY_LIMIT} | کانفیگ: {DAILY_CONFIG_LIMIT}\n\n"
            "📩 پشتیبانی: @sir_Aso"
        ),
        "ckb": (
            "🌟 <b>ڕێنمایی AsoLand</b>\n\n"
            "📥 دابەزاندنی فیلم و ڤیدیۆ لە ئینستاگرام، یوتیوب، تیک‌تۆک و هتد\n"
            "🎵 میوزیک، دۆزینەوەی گۆرانی و دەقی گۆرانی\n"
            "💰 نرخی دۆلار و دراو + گۆڕینی دراو\n"
            "🚀 تایبەتمەندییە زیرەکەکان: کەش و هەوا، بەختی ڕۆژانە، حیسابکەر، هێڵکاری، ئاگادارکردنەوەی نرخ، هەواڵ و ئامرازە خوێندکارییەکان\n"
            "🛠 ئامرازە بەسوودەکان: ستیکەر، QR، پەستاندن، PDF، Word و جوانکردنی دەق\n"
            "🤖 زیرەکی دەستکرد: گفتوگۆ، وەرگێڕان، کورتەکردنەوە، دەنگ بۆ دەق و مامۆستای ئینگلیزی\n"
            "🎬 دروستکردنی ژێرنووس (کوردی / فارسی / ئینگلیزی)\n"
            "🩷 کۆنفیگی بەخۆڕایی\n\n"
            f"📊 سنووری ڕۆژانە — داگرتن: {DAILY_LIMIT} | کۆنفیگ: {DAILY_CONFIG_LIMIT}\n\n"
            "📩 پشتگیری: @sir_Aso"
        ),
        "en": (
            "🌟 <b>AsoLand Help</b>\n\n"
            "📥 Download videos from Instagram, YouTube, TikTok and more\n"
            "🎵 Music, song search and lyrics\n"
            "💰 Dollar and coin prices + currency conversion\n"
            "🚀 Smart features: weather, daily fortune, calculator, charts, price alerts, news and student tools\n"
            "🛠 Useful tools: stickers, QR, compression, PDF, Word and fancy text\n"
            "🤖 AI: chat, translation, summarization, voice-to-text and English teacher\n"
            "🎬 Create subtitles (Kurdish / Persian / English)\n"
            "🩷 Free config\n\n"
            f"📊 Daily limits — downloads: {DAILY_LIMIT} | configs: {DAILY_CONFIG_LIMIT}\n\n"
            "📩 Support: @sir_Aso"
        ),
    }[lang]
    markup = InlineKeyboardMarkup([[LButton("🏠 منوی اصلی", callback_data="main_menu", _lang=lang)]])
    if update.callback_query:
        await edit_query_localized(update.callback_query, context, help_texts, parse_mode="HTML", reply_markup=markup)
    else:
        await reply_text_localized(update.message, context, help_texts, parse_mode="HTML", reply_markup=get_reply_keyboard())


# =========================================================
# SMART FILE READER — PDF / DOCX / TXT
# =========================================================

SMART_FILE_EXTENSIONS = {".pdf", ".docx", ".txt"}


def smart_file_keyboard():
    return InlineKeyboardMarkup([
        [
            LButton("📌 خلاصه", callback_data="file_action_summary"),
            LButton("🔍 جستجو", callback_data="file_action_search"),
        ],
        [
            LButton("❓ سؤال از فایل", callback_data="file_action_question"),
            LButton("🌐 ترجمه", callback_data="file_action_translate"),
        ],
        [
            LButton("📝 استخراج نکات", callback_data="file_action_notes"),
            LButton("🎓 ساخت آزمون", callback_data="file_action_quiz"),
        ],
        [LButton("🧠 فلش‌کارت", callback_data="file_action_flashcards")],
        [LButton("🗑 حذف فایل", callback_data="file_action_clear")],
        [LButton("🏠 منوی اصلی", callback_data="main_menu")],
    ])


def _smart_normalize(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_smart_file_text(path: str, extension: str) -> str:
    extension = extension.lower()
    if extension == ".txt":
        raw = Path(path).read_bytes()
        for enc in ("utf-8-sig", "utf-8", "cp1256", "cp1252", "latin-1"):
            try:
                return _smart_normalize(raw.decode(enc))
            except UnicodeDecodeError:
                continue
        return _smart_normalize(raw.decode("utf-8", errors="replace"))

    if extension == ".docx":
        doc = Document(path)
        chunks = []
        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if value:
                chunks.append(value)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    chunks.append(" | ".join(cells))
        return _smart_normalize("\n".join(chunks))

    if extension == ".pdf":
        if PdfReader is None:
            raise RuntimeError("برای خواندن PDF کتابخانه pypdf نصب نیست.\n\npip install pypdf")
        reader = PdfReader(path)
        chunks = []
        for page_no, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                page_text = f"[خطا در خواندن صفحه {page_no}: {exc}]"
            if page_text.strip():
                chunks.append(f"[صفحه {page_no}]\n{page_text}")
        return _smart_normalize("\n\n".join(chunks))

    raise ValueError("فرمت فایل پشتیبانی نمی‌شود. فقط PDF، DOCX و TXT مجاز است.")


def smart_file_excerpt(text: str, limit: int = SMART_FILE_MAX_CHARS) -> str:
    text = text or ""
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[ادامه فایل برای جلوگیری از طول بیش از حد درخواست حذف شد.]"


def smart_split_message(text: str, limit: int = 3900):
    text = text or ""
    return [text[i:i + limit] for i in range(0, len(text), limit)] or [""]


async def smart_ai(prompt: str, system: str, max_tokens: int = 1800) -> str:
    if not GROQ_API_KEY:
        return "❌ کلید هوش مصنوعی تنظیم نشده است."
    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            response = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": GROQ_CHAT_MODEL,
                    "messages": [
                        {"role": "system", "content": system},
                        {"role": "user", "content": prompt},
                    ],
                    "temperature": 0.2,
                    "max_completion_tokens": max_tokens,
                },
            )
            response.raise_for_status()
            data = response.json()
            return data["choices"][0]["message"]["content"].strip()
    except Exception as exc:
        logger.exception("Smart file AI error: %s", exc)
        return "❌ پردازش هوشمند فایل با خطا مواجه شد. دوباره تلاش کن."


async def smart_file_ai_action(text: str, action: str, user_input: str = "") -> str:
    content = smart_file_excerpt(text)
    base = (
        "تو دستیار فایل AsoLand هستی. فقط بر اساس محتوای فایل پاسخ بده. "
        "اگر اطلاعات موردنظر در فایل نیست، صریحاً بگو در فایل پیدا نشد. "
        "خروجی باید کاملاً فارسی معیار ایران باشد و از واژه‌ها یا ساختار کردی استفاده نکند. "
        "از Markdown سنگین استفاده نکن."
    )
    if action == "summary":
        prompt = "فایل زیر را منظم و دقیق خلاصه کن. عنوان‌های اصلی، ایده‌های مهم و نتیجه‌گیری را حفظ کن.\n\n" + content
        system = base + "\nوظیفه: خلاصه‌سازی دقیق بدون ساختن اطلاعات جدید."
    elif action == "notes":
        prompt = "از فایل زیر مهم‌ترین نکات را به صورت فهرست شماره‌دار استخراج کن. نکات کلیدی، تعریف‌ها، اعداد و نتیجه‌ها را جا نینداز.\n\n" + content
        system = base + "\nوظیفه: استخراج نکات کلیدی."
    elif action == "question":
        prompt = "به سؤال زیر فقط با استفاده از فایل پاسخ بده. اگر لازم بود به بخش یا صفحه مربوط اشاره کن.\n\nسؤال: " + user_input + "\n\nفایل:\n" + content
        system = base + "\nوظیفه: پرسش و پاسخ مستند از متن فایل."
    elif action == "translate":
        prompt = "متن فایل را به زبان مقصد زیر ترجمه کن. ساختار و معنی را حفظ کن و چیزی اضافه نکن.\n\nزبان مقصد: " + user_input + "\n\nمتن:\n" + content
        system = "تو مترجم دقیق هستی. ترجمه روان و وفادار به متن بده."
    elif action == "quiz":
        prompt = "از محتوای فایل دقیقاً ۲۰ سؤال امتحانی بساز. برای هر سؤال ۴ گزینه و در پایان همان سؤال پاسخ صحیح را بنویس. سؤال‌ها باید فقط از فایل باشند و تکراری نباشند.\n\n" + content
        system = base + "\nوظیفه: ساخت آزمون ۲۰ سؤالی چهارگزینه‌ای."
    elif action == "flashcards":
        prompt = "از فایل ۲۰ فلش‌کارت آموزشی بساز. هر مورد شامل «سؤال/مفهوم» و «پاسخ کوتاه» باشد. فقط از اطلاعات فایل استفاده کن.\n\n" + content
        system = base + "\nوظیفه: ساخت فلش‌کارت آموزشی."
    else:
        return "❌ عملیات فایل نامعتبر است."
    return await smart_ai(prompt, system, max_tokens=3500 if action in ("quiz", "flashcards") else 2200)


async def handle_smart_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    _CURRENT_UI_LANG.set(get_user_lang(update.effective_user.id, context))
    user_id = update.effective_user.id
    if user_id in blacklist:
        return
    document = update.message.document
    if not document:
        return

    filename = document.file_name or "file"
    extension = Path(filename).suffix.lower()
    if extension not in SMART_FILE_EXTENSIONS:
        await reply_text_localized(update.message, context, 
            "❌ این فایل پشتیبانی نمی‌شود.\n\n📄 فرمت‌های مجاز: PDF، DOCX، TXT",
            reply_markup=get_main_keyboard(),
        )
        return
    if document.file_size and document.file_size > SMART_FILE_MAX_SIZE:
        await reply_text_localized(update.message, context, "❌ حجم فایل بیشتر از ۲۵ مگابایت است.", reply_markup=get_main_keyboard())
        return

    status = await reply_text_localized(update.message, context, "📥 فایل دریافت شد.\n⏳ در حال استخراج متن...")
    temp_dir = tempfile.mkdtemp(prefix="smart_file_")
    local_path = os.path.join(temp_dir, filename)
    try:
        tg_file = await context.bot.get_file(document.file_id)
        await tg_file.download_to_drive(local_path)
        extracted = await asyncio.to_thread(extract_smart_file_text, local_path, extension)
        if not extracted:
            await edit_text_localized(status, context, "❌ از این فایل متن قابل خواندن پیدا نشد.", reply_markup=get_main_keyboard())
            return
        context.user_data["smart_file_text"] = extracted
        context.user_data["smart_file_name"] = filename
        context.user_data["smart_file_temp_dir"] = temp_dir
        context.user_data["mode"] = "smart_file_menu"
        size_text = f"{len(extracted):,} کاراکتر"
        await edit_text_localized(status, context, 
            f"📄 <b>فایل دریافت شد</b>\n\n"
            f"📎 {html.escape(filename)}\n"
            f"📝 متن استخراج‌شده: {size_text}\n\n"
            f"چه کاری انجام بدم؟",
            parse_mode="HTML",
            reply_markup=smart_file_keyboard(),
        )
    except Exception as exc:
        logger.exception("Smart document error: %s", exc)
        await edit_text_localized(status, context, f"❌ خطا در خواندن فایل:\n{html.escape(str(exc))}", parse_mode="HTML", reply_markup=get_main_keyboard())
        shutil.rmtree(temp_dir, ignore_errors=True)


async def smart_file_cleanup(context):
    temp_dir = context.user_data.pop("smart_file_temp_dir", None)
    context.user_data.pop("smart_file_text", None)
    context.user_data.pop("smart_file_name", None)
    context.user_data.pop("smart_file_pending_action", None)
    if temp_dir:
        shutil.rmtree(temp_dir, ignore_errors=True)


async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    _CURRENT_UI_LANG.set(get_user_lang(update.effective_user.id, context))

    if data == "language_menu":
        await edit_query_localized(query, context, 
            tr("لطفاً زبان را انتخاب کنید:", _CURRENT_UI_LANG.get()),
            reply_markup=language_keyboard()
        )
        return

    if data.startswith("set_lang_"):
        lang = data.replace("set_lang_", "", 1)
        if lang not in SUPPORTED_LANGUAGES:
            lang = "fa"

        # تغییر زبان = ری‌استارت کاملِ وضعیت همان کاربر
        # پردازش‌های ربات متوقف نمی‌شوند، فقط session کاربر از صفر ساخته می‌شود.
        set_user_lang(update.effective_user.id, lang, context)
        context.user_data.clear()
        context.user_data["language"] = lang
        _CURRENT_UI_LANG.set(lang)
        register_user(update.effective_user.id)

        name = update.effective_user.first_name or ("کاربر" if lang == "fa" else "User")
        welcome = {
            "fa": (f"سلام <b>{html.escape(name)}</b> 👋\n\n{get_jalali_datetime('fa')}\n\n🌟 <b>AsoLand</b>\nربات همه‌کاره و هوشمند\n\nاز منوی زیر استفاده کن."),
            "ckb": (f"سڵاو <b>{html.escape(name)}</b> 👋\n\n{get_jalali_datetime('ckb')}\n\n🌟 <b>AsoLand</b>\nبۆتی زیرەکی هەمووکارە\n\nلە پێڕستی خوارەوە هەڵبژێرە."),
            "en": (f"Hello <b>{html.escape(name)}</b> 👋\n\n{get_jalali_datetime('en')}\n\n🌟 <b>AsoLand</b>\nAll-in-one smart bot\n\nChoose an option from the menu below."),
        }[lang]

        # پیام فعلی به پیام شروعِ تازه تبدیل می‌شود و Reply Keyboard هم بلافاصله عوض می‌شود.
        await query.edit_message_text(
            welcome, parse_mode="HTML", reply_markup=get_main_keyboard()
        )
        await context.bot.send_message(
            chat_id=update.effective_chat.id,
            text={
                "fa": "✅ زبان با موفقیت تغییر کرد.",
                "ckb": "✅ زمان بە سەرکەوتوویی گۆڕدرا.",
                "en": "✅ Language changed successfully.",
            }[lang],
            reply_markup=get_reply_keyboard()
        )
        return
    user_id = update.effective_user.id
    register_user(user_id)

    if data == "smart_file_reader":
        context.user_data["mode"] = "smart_file_waiting"
        await edit_query_localized(query, context, 
            "📁 <b>فایل‌خوان هوشمند AsoLand</b>\n\n"
            "فایل PDF، Word (DOCX) یا TXT را همین‌جا ارسال کن.\n\n"
            "بعد از دریافت فایل می‌توانی خلاصه، جستجو، سؤال، ترجمه، نکات، آزمون یا فلش‌کارت بگیری.",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([[LButton("🏠 منوی اصلی", callback_data="main_menu")]]),
        )
        return

    if data.startswith("file_action_"):
        file_text = context.user_data.get("smart_file_text")
        if not file_text:
            await edit_query_localized(query, context, "❌ اول یک فایل PDF، DOCX یا TXT ارسال کن.", reply_markup=get_main_keyboard())
            return
        action = data.replace("file_action_", "", 1)
        if action == "clear":
            await smart_file_cleanup(context)
            context.user_data["mode"] = None
            await edit_query_localized(query, context, "🗑 فایل از حافظه این جلسه حذف شد.", reply_markup=get_main_keyboard())
            return
        if action == "search":
            context.user_data["mode"] = "smart_file_search"
            context.user_data["smart_file_pending_action"] = "search"
            await edit_query_localized(query, context, "🔍 عبارت موردنظر برای جستجو در فایل را بفرست.", reply_markup=get_back_keyboard())
            return
        if action == "question":
            context.user_data["mode"] = "smart_file_question"
            context.user_data["smart_file_pending_action"] = "question"
            await edit_query_localized(query, context, "❓ سؤالت را درباره محتوای فایل بنویس.", reply_markup=get_back_keyboard())
            return
        if action == "translate":
            context.user_data["mode"] = "smart_file_translate"
            context.user_data["smart_file_pending_action"] = "translate"
            await edit_query_localized(query, context, "🌐 زبان مقصد را بنویس. مثال: انگلیسی، عربی، فارسی، کردی", reply_markup=get_back_keyboard())
            return
        await edit_query_localized(query, context, "⏳ در حال پردازش فایل...")
        result = await smart_file_ai_action(file_text, action)
        for index, part in enumerate(smart_split_message(result)):
            if index == 0:
                await edit_query_localized(query, context, part, reply_markup=smart_file_keyboard() if len(result) <= 3900 else None)
            else:
                await bot_send_message_localized(context.bot, query.message.chat_id, part)
        return

    if data.startswith("header_"):
        return

    if data == "extended_tools":
        await edit_query_localized(query, context, "🚀 <b>امکانات هوشمند AsoLand</b>\n\nیکی از سرویس‌ها را انتخاب کن:", parse_mode="HTML", reply_markup=extended_tools_keyboard())
        return

    if data == "chart_menu":
        await edit_query_localized(query, context, "📊 <b>نمودار قیمت</b>\n\nرمزارز را انتخاب کن:", parse_mode="HTML", reply_markup=InlineKeyboardMarkup([
            [LButton("₿ بیت‌کوین ۷ روزه", callback_data="chart_btc_7"), LButton("🔷 اتریوم ۷ روزه", callback_data="chart_eth_7")],
            [LButton("🌤 نمودار دمای شهر", callback_data="chart_weather")],
            [LButton("🏠 منوی اصلی", callback_data="main_menu")]
        ]))
        return

    if data.startswith("chart_") and data not in ("chart_menu", "chart_weather"):
        parts=data.split("_"); coin="bitcoin" if parts[1]=="btc" else "ethereum"; days=int(parts[2])
        path=await get_crypto_history_chart(coin,days)
        if not path: await edit_query_localized(query, context, "❌ساخت نمودار ناموفق بود.", reply_markup=extended_tools_keyboard()); return
        with open(path,"rb") as f: await context.bot.send_photo(query.message.chat_id,f,caption=tr(f"📊 نمودار {coin} - {days} روز", get_user_lang(query.message.chat_id)))
        try: os.remove(path)
        except: pass
        await edit_query_localized(query, context, "✅ نمودار ارسال شد.",reply_markup=extended_tools_keyboard()); return

    if data == "chart_weather":
        context.user_data["mode"]="weather_chart"
        await edit_query_localized(query, context, "🌤 نام شهر را بفرست؛ نمودار دمای ۷ روزه می‌سازم.",reply_markup=get_back_keyboard()); return

    if data == "alert_menu":
        context.user_data["mode"]="price_alert"
        await edit_query_localized(query, context, "🔔 هشدار قیمت\n\nفرمت:\nدلار above 190000\nدلار below 180000\nطلا below 20000000\nسکه above 190000000",reply_markup=get_back_keyboard()); return

    if data == "news_menu":
        await edit_query_localized(query, context, "📰 دسته خبر را انتخاب کن:",reply_markup=InlineKeyboardMarkup([
            [LButton("🌠عمومی",callback_data="news_general"),LButton("💻 فناوری",callback_data="news_tech")],
            [LButton("💵 اقتصاد",callback_data="news_economy"),LButton("₿ کریپتو",callback_data="news_crypto")],
            [LButton("🎓 آموزش",callback_data="news_student")],[LButton("🏠 منوی اصلی",callback_data="main_menu")]
        ])); return

    if data.startswith("news_") and data != "news_menu":
        category=data.replace("news_",""); text=await news_text(category)
        await edit_query_localized(query, context, text,parse_mode="HTML",disable_web_page_preview=True,reply_markup=InlineKeyboardMarkup([[LButton("🔄 بروزرسانی",callback_data=data)],[LButton("🏠 منوی اصلی",callback_data="main_menu")]])); return

    if data == "student_menu":
        context.user_data["mode"]="student_text"
        await edit_query_localized(query, context, "🎓 <b>ابزارهای دانشجویی</b>\n\nسؤال ریاضی، فیزیک، شیمی یا مسئله درسی را بفرست.\n\nبرای حل از روی عکس، گزینه «حل سؤال از عکس» را بزن.",parse_mode="HTML",reply_markup=InlineKeyboardMarkup([[LButton("📸 حل از عکس",callback_data="solve_image")],[LButton("🏠 منوی اصلی",callback_data="main_menu")]])); return

    if data == "solve_image":
        context.user_data["mode"]="solve_image"
        await edit_query_localized(query, context, "📸 عکس سؤال را بفرست. متن سؤال را تشخیص می‌دهم و مرحله‌به‌مرحله حل می‌کنم.",reply_markup=get_back_keyboard()); return

    # ---------- قیمت‌ها ----------
    # ---------- قیمت‌ها (DO_L4) ----------
    if data == "price_do_l4":
        try:
            await edit_query_localized(query, context, "⏳ در حال دریافت قیمت دلار و سکه...")
        except Exception:
            pass
        try:
            text = await get_do_l4_prices(get_user_lang(query.message.chat_id, context))
            if len(text) > 4000:
                text = text[:3900] + "\n\n…"
            await edit_query_localized(query, context, 
                text,
                parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([
                    [LButton("🔄 بروزرسانی", callback_data="price_do_l4")],
                    [LButton("🏠 منوی اصلی", callback_data="main_menu")]
                ]),
                disable_web_page_preview=True,
            )
        except Exception as e:
            logger.error(f"price_do_l4 error: {e}")
            try:
                await edit_query_localized(query, context, 
                    f"❌ خطا در دریافت قیمت.\n<code>{html.escape(str(e)[:200])}</code>",
                    parse_mode="HTML",
                    reply_markup=InlineKeyboardMarkup([
                        [LButton("🔄 تلاش دوباره", callback_data="price_do_l4")],
                        [LButton("🏠 منوی اصلی", callback_data="main_menu")]
                    ]),
                )
            except Exception:
                pass
        return

    # ---------- امکانات جدید ----------
    if data == "weather_menu":
        context.user_data["mode"] = "weather"
        await edit_query_localized(query, context, 
            "🌤 <b>آب‌وهوا</b>\n\nنام شهر را بفرست؛ مثال: تهران، آمستردام، اربیل",
            parse_mode="HTML", reply_markup=get_back_keyboard()
        )
        return

    if data == "daily_fortune":
        name = update.effective_user.first_name or "دوست من"
        text = await get_daily_fortune(name)
        await edit_query_localized(query, context, 
            text, parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🔄 فال دوباره", callback_data="daily_fortune")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ])
        )
        return

    if data == "advanced_calculator":
        context.user_data["mode"] = "advanced_calculator"
        await edit_query_localized(query, context, 
            "🧮 <b>محاسبه‌گر پیشرفته</b>\n\n"
            "عبارت ریاضی، تابع یا معادله را بفرست.\n"
            "مثال‌ها:\n"
            "• (25+7)*3\n"
            "• sqrt(144)+sin(pi/2)\n"
            "• x^2-5*x+6=0\n"
            "• 2**10",
            parse_mode="HTML", reply_markup=get_back_keyboard()
        )
        return

    
    # ---------- امکانات جدید ----------
    if data == "currency_convert":
        context.user_data["mode"] = "currency_convert"
        await edit_query_localized(query, context, 
            "💱 <b>تبدیل ارز</b>\n\n"
            "مقدار را بفرست؛ مثال:\n"
            "• ۱۰۰ دلار\n"
            "• ۱۰۰ دلار به تومان\n"
            "• ۵۰۰۰۰۰۰ تومان به دلار",
            parse_mode="HTML",
            reply_markup=get_back_keyboard(),
        )
        return

    if data == "user_panel":
        name = update.effective_user.first_name or "کاربر"
        text = get_user_panel(user_id, name)
        await edit_query_localized(query, context, 
            text,
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🔄 بروزرسانی", callback_data="user_panel")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")],
            ]),
        )
        return

    if data == "today_calendar":
        text = get_today_calendar()
        await edit_query_localized(query, context, 
            text,
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🔄 بروزرسانی", callback_data="today_calendar")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")],
            ]),
        )
        return

    if data == "reminder_menu":
        context.user_data["mode"] = "reminder"
        await edit_query_localized(query, context, 
            "🔔 <b>یادآور</b>\n\n"
            "مثال:\n"
            "• یادآوری ساعت ۱۰ جلسه\n"
            "• یادآوری ۱۰:۳۰ خرید نان\n"
            "• یادآوری فردا ساعت ۹ کلاس",
            parse_mode="HTML",
            reply_markup=get_back_keyboard(),
        )
        return

    if data == "text_sticker":
        context.user_data["mode"] = None
        context.user_data.pop("sticker_bg", None)
        context.user_data.pop("sticker_fg", None)
        context.user_data.pop("sticker_font", None)
        try:
            await edit_query_localized(query, context, 
                "🎨 <b>استیکر متنی</b>\n\nاول رنگ <b>پس‌زمینه</b> را انتخاب کن:",
                parse_mode="HTML",
                reply_markup=sticker_bg_keyboard(),
            )
        except Exception as e:
            logger.error(f"text_sticker open error: {e}")
            try:
                await bot_send_message_localized(context.bot, query.message.chat_id,
                    "🎨 <b>استیکر متنی</b>\n\nاول رنگ <b>پس‌زمینه</b> را انتخاب کن:",
                    parse_mode="HTML",
                    reply_markup=sticker_bg_keyboard(),
                )
            except Exception as e2:
                logger.error(f"text_sticker send error: {e2}")
                await query.answer(tr("خطا در باز کردن استیکر متنی", get_user_lang(query.message.chat_id, context)), show_alert=True)
        return

    if data.startswith("stbg_"):
        key = data  # stbg_bg_blue
        # map callback to color key
        color_key = data.replace("stbg_", "", 1)  # bg_blue
        if color_key not in STICKER_BG_COLORS:
            await query.answer(tr("رنگ نامعتبر", get_user_lang(query.message.chat_id, context)), show_alert=True)
            return
        context.user_data["sticker_bg"] = color_key
        label = STICKER_BG_COLORS[color_key][1]
        await edit_query_localized(query, context, 
            f"🎨 پس‌زمینه: <b>{label}</b>\n\nحالا رنگ <b>متن</b> را انتخاب کن:",
            parse_mode="HTML",
            reply_markup=sticker_fg_keyboard(),
        )
        return

    if data.startswith("stfg_"):
        color_key = data.replace("stfg_", "", 1)  # fg_white
        if color_key not in STICKER_FG_COLORS:
            await query.answer(tr("رنگ نامعتبر", get_user_lang(query.message.chat_id, context)), show_alert=True)
            return
        if not context.user_data.get("sticker_bg"):
            await edit_query_localized(query, context, 
                "ابتدا رنگ پس‌زمینه را انتخاب کن:",
                reply_markup=sticker_bg_keyboard(),
            )
            return
        context.user_data["sticker_fg"] = color_key
        bg_label = STICKER_BG_COLORS[context.user_data["sticker_bg"]][1]
        fg_label = STICKER_FG_COLORS[color_key][1]
        await edit_query_localized(query, context, 
            f"🎨 پس‌زمینه: <b>{bg_label}</b>\n"
            f"✏️ متن: <b>{fg_label}</b>\n\n"
            "حالا <b>فونت</b> را انتخاب کن (۱۵ فونت کوردی/عربی):",
            parse_mode="HTML",
            reply_markup=sticker_font_keyboard(),
        )
        return

    if data == "sticker_pick_fg":
        if not context.user_data.get("sticker_bg"):
            await edit_query_localized(query, context, 
                "ابتدا رنگ پس‌زمینه را انتخاب کن:",
                reply_markup=sticker_bg_keyboard(),
            )
            return
        bg_label = STICKER_BG_COLORS[context.user_data["sticker_bg"]][1]
        await edit_query_localized(query, context, 
            f"🎨 پس‌زمینه: <b>{bg_label}</b>\n\nحالا رنگ <b>متن</b> را انتخاب کن:",
            parse_mode="HTML",
            reply_markup=sticker_fg_keyboard(),
        )
        return

    if data.startswith("istfg_"):
        color_key = data.replace("istfg_", "", 1)
        if color_key not in STICKER_FG_COLORS:
            await query.answer(tr("رنگ نامعتبر", get_user_lang(query.message.chat_id, context)), show_alert=True)
            return
        photo_path = context.user_data.get("sticker_photo_path")
        if not photo_path or not os.path.exists(photo_path):
            context.user_data["mode"] = "image_sticker_wait_photo"
            await edit_query_localized(query, context, 
                "📸 عکس پیدا نشد. دوباره عکس را بفرست.",
                reply_markup=get_back_keyboard(),
            )
            return
        context.user_data["sticker_fg"] = color_key
        fg_label = STICKER_FG_COLORS[color_key][1]
        context.user_data["mode"] = "image_sticker_wait_font"
        await edit_query_localized(query, context, 
            f"✏️ <b>رنگ متن:</b> {html.escape(fg_label)}\n\n"
            "🔤 حالا <b>فونت</b> موردنظرت را انتخاب کن:",
            parse_mode="HTML",
            reply_markup=image_sticker_font_keyboard(),
        )
        return

    if data == "image_sticker_change_color":
        if not context.user_data.get("sticker_photo_path"):
            context.user_data["mode"] = "image_sticker_wait_photo"
            await edit_query_localized(query, context, "📸 ابتدا عکس را بفرست.", reply_markup=get_back_keyboard())
            return
        context.user_data["mode"] = "image_sticker_wait_fg"
        await edit_query_localized(query, context, 
            "🎨 رنگ جدید متن را انتخاب کن:",
            reply_markup=image_sticker_fg_keyboard(),
        )
        return

    if data.startswith("istfont_"):
        font_key = data.replace("istfont_", "", 1)
        if font_key not in STICKER_FONTS:
            await query.answer(tr("فونت نامعتبر", get_user_lang(query.message.chat_id, context)), show_alert=True)
            return

        photo_path = context.user_data.get("sticker_photo_path")
        if not photo_path or not os.path.exists(photo_path):
            context.user_data["mode"] = "image_sticker_wait_photo"
            await edit_query_localized(query, context, 
                "📸 عکس پیدا نشد. دوباره عکس را بفرست.",
                reply_markup=get_back_keyboard(),
            )
            return

        if not context.user_data.get("sticker_fg"):
            context.user_data["mode"] = "image_sticker_wait_fg"
            await edit_query_localized(query, context, 
                "🎨 ابتدا رنگ متن را انتخاب کن:",
                reply_markup=image_sticker_fg_keyboard(),
            )
            return
        context.user_data["sticker_font"] = font_key
        context.user_data["mode"] = "image_sticker_wait_text"
        font_label = STICKER_FONTS[font_key]["label"]
        fg_label = STICKER_FG_COLORS[context.user_data["sticker_fg"]][1]
        await edit_query_localized(query, context, 
            f"✏️ <b>رنگ متن:</b> {html.escape(fg_label)}\n"
            f"🔤 <b>فونت:</b> {html.escape(font_label)}\n\n"
            "✍️ حالا <b>متن موردنظرت</b> را بفرست تا روی عکس نوشته شود.",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🔤 تغییر فونت", callback_data="image_sticker_change_font")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")],
            ]),
        )
        return

    if data == "image_sticker_change_font":
        if not context.user_data.get("sticker_photo_path"):
            context.user_data["mode"] = "image_sticker_wait_photo"
            await edit_query_localized(query, context, 
                "📸 ابتدا عکس را بفرست.",
                reply_markup=get_back_keyboard(),
            )
            return
        context.user_data["mode"] = "image_sticker_wait_font"
        await edit_query_localized(query, context, 
            "🔤 فونت موردنظرت را انتخاب کن:",
            reply_markup=image_sticker_font_keyboard(),
        )
        return

    if data.startswith("stfont_"):
        font_key = data.replace("stfont_", "", 1)
        if font_key not in STICKER_FONTS:
            await query.answer(tr("فونت نامعتبر", get_user_lang(query.message.chat_id, context)), show_alert=True)
            return
        if not context.user_data.get("sticker_bg") or not context.user_data.get("sticker_fg"):
            await edit_query_localized(query, context, 
                "ابتدا رنگ‌ها را انتخاب کن:",
                reply_markup=sticker_bg_keyboard(),
            )
            return
        context.user_data["sticker_font"] = font_key
        context.user_data["mode"] = "text_sticker"
        bg_label = STICKER_BG_COLORS[context.user_data["sticker_bg"]][1]
        fg_label = STICKER_FG_COLORS[context.user_data["sticker_fg"]][1]
        font_label = STICKER_FONTS[font_key]["label"]
        await edit_query_localized(query, context, 
            f"🎨 پس‌زمینه: <b>{bg_label}</b>\n"
            f"✏️ متن: <b>{fg_label}</b>\n"
            f"🔤 فونت: <b>{html.escape(font_label)}</b>\n\n"
            "حالا <b>متن فارسی یا کوردی</b> را بفرست.\n"
            "مثال: سلام دوستان / سڵاو چۆنی\n\n"
            "ℹ️ اگر فونت روی دستگاه نباشد، یک‌بار دانلود می‌شود.",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🔙 تغییر فونت", callback_data="sticker_pick_font")],
                [LButton("🎨 تغییر رنگ‌ها", callback_data="text_sticker")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")],
            ]),
        )
        return

    if data == "sticker_pick_font":
        if not context.user_data.get("sticker_bg") or not context.user_data.get("sticker_fg"):
            await edit_query_localized(query, context, 
                "ابتدا رنگ‌ها را انتخاب کن:",
                reply_markup=sticker_bg_keyboard(),
            )
            return
        context.user_data["mode"] = None
        bg_label = STICKER_BG_COLORS[context.user_data["sticker_bg"]][1]
        fg_label = STICKER_FG_COLORS[context.user_data["sticker_fg"]][1]
        await edit_query_localized(query, context, 
            f"🎨 پس‌زمینه: <b>{bg_label}</b>\n"
            f"✏️ متن: <b>{fg_label}</b>\n\n"
            "فونت را انتخاب کن:",
            parse_mode="HTML",
            reply_markup=sticker_font_keyboard(),
        )
        return


    if data == "ocr_image":
        context.user_data["mode"] = "ocr_image"
        await edit_query_localized(query, context, 
            "🔤 <b>OCR — استخراج متن از عکس</b>\n\nعکس دارای متن را بفرست.",
            parse_mode="HTML",
            reply_markup=get_back_keyboard(),
        )
        return

    # ---------- زیرنویس ----------
    if data == "subtitle_maker":
        await edit_query_localized(query, context, 
            "🎬 <b>ساخت زیرنویس</b>\n\n"
            "زبان زیرنویس مورد نظرت رو انتخاب کن:\n\n"
            "☀️ کوردی\n"
            "🇮🇷 فارسی\n"
            "🇬🇧 انگلیسی",
            parse_mode="HTML",
            reply_markup=get_subtitle_lang_keyboard()
        )
        return

    if data.startswith("sub_lang_"):
        lang = data.replace("sub_lang_", "")
        context.user_data["mode"] = "subtitle_maker"
        context.user_data["subtitle_lang"] = lang
        lang_names = {"ckb": "☀️ کوردی", "fa": "🇮🇷 فارسی", "en": "🇬🇧 انگلیسی"}
        await edit_query_localized(query, context, 
            f"✅ زبان زیرنویس: <b>{lang_names.get(lang)}</b>\n\n"
            "حالا ویدیوی کوتاه (حداکثر ۹۰ ثانیه) رو بفرست.\n"
            "ربات متن را استخراج و به زبان انتخاب‌شده ترجمه می‌کند.",
            parse_mode="HTML",
            reply_markup=get_back_keyboard()
        )
        return

    if data in ["sub_send_srt", "sub_send_video", "sub_send_both"]:
        result = context.user_data.get("subtitle_result")
        if not result:
            await edit_query_localized(query, context, "❌اطلاعات منقضی شده. دوباره ویدیو بفرست.", reply_markup=get_main_keyboard())
            return

        chat_id = query.message.chat_id
        lang_names = {"ckb": "☀️ کوردی", "fa": "🇮🇷 فارسی", "en": "🇬🇧 انگلیسی"}
        lang_name = lang_names.get(result.get("lang", ""), "زیرنویس")

        try:
            if data in ["sub_send_srt", "sub_send_both"]:
                with open(result["srt_path"], "rb") as f:
                    await context.bot.send_document(chat_id, f, caption=tr(f"📄 فایل زیرنویس {lang_name} (.srt)", get_user_lang(chat_id)))

            if data in ["sub_send_video", "sub_send_both"]:
                if result.get("video_path") and os.path.exists(result["video_path"]):
                    with open(result["video_path"], "rb") as f:
                        await context.bot.send_video(chat_id, f, caption=tr(f"🎥 ویدیو با زیرنویس {lang_name}", get_user_lang(chat_id)), supports_streaming=True)
                else:
                    await bot_send_message_localized(context.bot, chat_id, "⚠️ ویدیوی زیرنویس‌دار در دسترس نیست.\nاز فایل SRT استفاده کنید.")

            await edit_query_localized(query, context, "✅ ارسال شد.", reply_markup=get_main_keyboard())
        except Exception as e:
            logger.error(e)
            await edit_query_localized(query, context, f"❌خطا در ارسال: {e}", reply_markup=get_main_keyboard())
        return

    # ---------- کانفیگ ----------
    if data == "free_config":
        if not check_config_limit(user_id):
            await edit_query_localized(query, context, f"❌محدودیت روزانه تموم شده ({DAILY_CONFIG_LIMIT} تا).", reply_markup=get_main_keyboard())
            return
        await edit_query_localized(query, context, "🩷 <b>کانفیگ رایگان</b>\n\nنوع کانفیگ رو انتخاب کن:", parse_mode="HTML", reply_markup=get_config_keyboard())
        return

    if data.startswith("cfg_"):
        if not check_config_limit(user_id):
            await edit_query_localized(query, context, "❌محدودیت روزانه تموم شده.", reply_markup=get_main_keyboard())
            return
        protocol = data.replace("cfg_", "")
        await edit_query_localized(query, context, "⏳در حال پیدا کردن کانفیگ سالم...")
        config = await get_working_config(protocol)
        if not config:
            await edit_query_localized(query, context, "❌کانفیگ سالم پیدا نشد.", reply_markup=get_config_keyboard())
            return
        increase_config_count(user_id)
        remaining = DAILY_CONFIG_LIMIT - user_config_usage[user_id]["count"]
        await edit_query_localized(query, context, 
            f"✅ <b>کانفیگ آماده است</b>\n\n<code>{html.escape(config)}</code>\n\n📊 باقی‌مانده: {remaining}",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🩷 کانفیگ جدید", callback_data="free_config")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ])
        )
        return

    if data == "summarize_text":
        context.user_data["mode"] = "summarize"
        await edit_query_localized(query, context, "📠<b>خلاصه‌سازی متن</b>\n\nمتن یا مقاله رو بفرست.", parse_mode="HTML", reply_markup=get_back_keyboard())
        return

    if data == "text_to_word":
        context.user_data["mode"] = "text_to_word"
        await edit_query_localized(query, context, "📠متن رو بفرست تا به فایل ورد تبدیل کنم.", reply_markup=get_back_keyboard())
        return

    if data == "show_clock":
        await edit_query_localized(query, context, f"🕠<b>ساعت و تاریخ شمسی</b>\n\n{get_jalali_datetime()}", parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🔄 بروزرسانی", callback_data="show_clock")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ]))
        return

    if data == "fancy_text":
        context.user_data["mode"] = "fancy_text"
        await edit_query_localized(query, context, "✅ <b>زیباسازی متن</b>\n\nمتن مورد نظرت رو بفرست.", parse_mode="HTML", reply_markup=get_back_keyboard())
        return

    if data == "english_teacher":
        context.user_data["mode"] = "english_teacher"
        context.user_data["et_mode"] = "menu"
        level = context.user_data.get("english_level", "intermediate")
        await edit_query_localized(query, context, f"<b>English Teacher - Miss Emma</b>\n\nسطح فعلی: <b>{level}</b>", parse_mode="HTML", reply_markup=get_english_teacher_keyboard())
        return

    if data == "et_exit":
        context.user_data.clear()
        await edit_query_localized(query, context, "👋 خداحافظ!", reply_markup=get_main_keyboard())
        return

    if data == "et_set_level":
        await edit_query_localized(query, context, "🎯 سطح خودت رو انتخاب کن:", reply_markup=get_level_keyboard())
        return

    if data.startswith("level_"):
        context.user_data["english_level"] = data.replace("level_", "")
        await edit_query_localized(query, context, "✅ سطح تنظیم شد.", reply_markup=get_english_teacher_keyboard())
        return

    if data in ["et_conversation", "et_correct", "et_vocabulary", "et_grammar"]:
        mode_map = {"et_conversation": "conversation", "et_correct": "correct", "et_vocabulary": "vocabulary", "et_grammar": "grammar"}
        context.user_data["mode"] = "english_teacher"
        context.user_data["et_mode"] = mode_map[data]
        context.user_data["ai_history"] = []
        msgs = {"conversation": "💬 مکالمه آزاد", "correct": "✅ جمله رو بفرست", "vocabulary": "📚 لغت جدید", "grammar": "📖 گرامر"}
        await edit_query_localized(query, context, msgs[mode_map[data]], reply_markup=InlineKeyboardMarkup([[LButton("🔙 منوی معلم", callback_data="english_teacher")]]))
        return

    if data == "admin_stats" and is_admin(user_id):
        today = str(date.today())
        today_dl = sum(1 for u in user_downloads.values() if u.get("date") == today and u.get("count", 0) > 0)
        today_cfg = sum(1 for u in user_config_usage.values() if u.get("date") == today and u.get("count", 0) > 0)
        await edit_query_localized(query, context, 
            "📊 <b>آمار ربات</b>\n\n"
            f"👥 کل کاربران: <b>{len(all_users)}</b>\n"
            f"📥 کل دانلودها: <b>{total_downloads}</b>\n"
            f"🟢 دانلود امروز (کاربر): <b>{today_dl}</b>\n"
            f"🩷 کانفیگ امروز (کاربر): <b>{today_cfg}</b>\n"
            f"🚫 بلاک‌لیست: <b>{len(blacklist)}</b>",
            parse_mode="HTML",
            reply_markup=get_admin_keyboard()
        )
        return

    if data == "admin_broadcast" and is_admin(user_id):
        context.user_data["mode"] = "broadcast"
        await edit_query_localized(query, context, "📢 پیام همگانی رو بفرست:", reply_markup=InlineKeyboardMarkup([[LButton("❌انصراف", callback_data="admin_cancel_broadcast")]]))
        return

    if data == "admin_cancel_broadcast":
        context.user_data.pop("mode", None)
        await edit_query_localized(query, context, "❌لغو شد.", reply_markup=get_admin_keyboard())
        return

    if data == "admin_blacklist" and is_admin(user_id):
        text = "🚫 لیست بلاک خالی است." if not blacklist else "🚫 بلاک‌شده‌ها:\n" + "\n".join([f"• <code>{uid}</code>" for uid in list(blacklist)[:40]])
        await edit_query_localized(query, context, text, parse_mode="HTML", reply_markup=get_admin_keyboard())
        return

    if data == "voice_to_text":
        context.user_data["mode"] = "voice_to_text"
        await edit_query_localized(query, context, "🎤 <b>تبدیل ویس به متن</b>\n\nزبان رو انتخاب کن:", parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("☀️ کوردی", callback_data="stt_lang_ckb")],
                [LButton("فارسی", callback_data="stt_lang_fa")],
                [LButton("انگلیسی", callback_data="stt_lang_en")],
                [LButton("تشخیص خودکار", callback_data="stt_lang_auto")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ]))
        return

    if data.startswith("stt_lang_"):
        lang = data.replace("stt_lang_", "")
        context.user_data["stt_language"] = lang
        context.user_data["mode"] = "voice_to_text"
        names = {"ckb": "کوردی", "fa": "فارسی", "en": "انگلیسی", "auto": "تشخیص خودکار"}
        await edit_query_localized(query, context, f"✅ زبان: <b>{names.get(lang)}</b>\n\nحالا ویس رو بفرست 🎤", parse_mode="HTML", reply_markup=get_back_keyboard())
        return

    if data == "main_menu":
        context.user_data.clear()
        await edit_query_localized(query, context, f"🌟 <b>AsoLand</b>\n\n{get_jalali_datetime()}\n\nاز منوی زیر استفاده کن.", parse_mode="HTML", reply_markup=get_main_keyboard())
        return

    if data == "help":
        await help_command(update, context)
        return

    if data == "ai_chat":
        context.user_data["mode"] = "ai"
        context.user_data["ai_history"] = []
        await edit_query_localized(query, context, 
            "👋 سلام! من <b>آسو</b> هستم 😎\n\n"
            "دستیار باحال AsoLand. هر چی بخوای بپرس؛ "
            "از قیمت دلار تا حل مسئله ریاضی، از شوخی تا ترجمه.\n\n"
            "بزن بریم 💬",
            parse_mode="HTML",
            reply_markup=get_ai_keyboard()
        )
        return

    if data == "exit_ai":
        context.user_data.clear()
        await edit_query_localized(query, context, "✅ خارج شدی.", reply_markup=get_main_keyboard())
        return

    if data == "translate_text":
        context.user_data["mode"] = "translate"
        await edit_query_localized(query, context, "🌠متن رو بفرست تا ترجمه کنم.", reply_markup=get_back_keyboard())
        return

    if data == "make_sticker":
        context.user_data.pop("sticker_font", None)
        context.user_data.pop("sticker_fg", None)
        context.user_data.pop("sticker_photo_path", None)
        context.user_data.pop("sticker_photo_dir", None)
        context.user_data.pop("sticker_bg", None)
        context.user_data.pop("sticker_fg", None)
        context.user_data["mode"] = "sticker_choose_background"
        await edit_query_localized(query, context, 
            "🖼 <b>ساخت استیکر</b>\n\n"
            "اول مشخص کن استیکر را با چه پس‌زمینه‌ای می‌خواهی:\n\n"
            "🖼 <b>با عکس:</b> عکس خودت را می‌فرستی و متن روی عکس قرار می‌گیرد.\n"
            "🎨 <b>بدون عکس:</b> یک رنگ برای پس‌زمینه انتخاب می‌کنی و متن روی آن قرار می‌گیرد.",
            parse_mode="HTML",
            reply_markup=sticker_background_type_keyboard(),
        )
        return

    if data == "sticker_with_photo":
        context.user_data["mode"] = "image_sticker_wait_photo"
        context.user_data.pop("sticker_font", None)
        context.user_data.pop("sticker_fg", None)
        context.user_data.pop("sticker_photo_path", None)
        context.user_data.pop("sticker_photo_dir", None)
        context.user_data.pop("sticker_bg", None)
        context.user_data.pop("sticker_fg", None)
        await edit_query_localized(query, context, 
            "🖼 <b>ساخت استیکر با عکس</b>\n\n"
            "عکس موردنظرت را بفرست.\n"
            "بعد فونت را انتخاب می‌کنی و متن دلخواهت روی عکس نوشته می‌شود.",
            parse_mode="HTML",
            reply_markup=get_back_keyboard(),
        )
        return

    if data == "sticker_no_photo":
        # اگر قبلاً عکس انتخاب شده بود، مسیر موقت آن را پاک کن.
        old_dir = context.user_data.get("sticker_photo_dir")
        if old_dir and os.path.isdir(old_dir):
            try:
                shutil.rmtree(old_dir, ignore_errors=True)
            except Exception:
                pass

        context.user_data.pop("sticker_photo_path", None)
        context.user_data.pop("sticker_photo_dir", None)
        context.user_data.pop("sticker_font", None)
        context.user_data.pop("sticker_fg", None)
        context.user_data["mode"] = "text_sticker"

        await edit_query_localized(query, context, 
            "🎨 <b>استیکر بدون عکس</b>\n\n"
            "رنگ پس‌زمینه را انتخاب کن:",
            parse_mode="HTML",
            reply_markup=sticker_bg_keyboard(),
        )
        return

    if data == "qr_code":
        context.user_data["mode"] = "qr_code"
        await edit_query_localized(query, context, "📱 متن یا لینک بفرست تا QR بسازم.", reply_markup=get_back_keyboard())
        return

    if data == "compress_image":
        context.user_data["mode"] = "compress_image"
        await edit_query_localized(query, context, "🗜 عکس بفرست تا فشرده کنم.", reply_markup=get_back_keyboard())
        return

    if data == "compress_video":
        context.user_data["mode"] = "compress_video"
        await edit_query_localized(query, context, "🗜 ویدیو بفرست تا فشرده کنم.", reply_markup=get_back_keyboard())
        return

    if data == "photo_to_pdf":
        context.user_data["mode"] = "photo_to_pdf"
        context.user_data["pdf_photos"] = []
        await edit_query_localized(query, context, "📄 عکس‌ها رو یکی‌یکی بفرست.\nوقتی تموم شد بنویس: <b>تمام</b>", parse_mode="HTML", reply_markup=get_back_keyboard())
        return

    if data == "short_link":
        context.user_data["mode"] = "short_link"
        await edit_query_localized(query, context, "🔗 لینک بلند رو بفرست.", reply_markup=get_back_keyboard())
        return

    if data == "music_finder":
        context.user_data["mode"] = "music_finder"
        await edit_query_localized(query, context, "🎶 اسم آهنگ، ویس یا لینک بفرست.", reply_markup=get_back_keyboard())
        return

    if data == "lyrics_menu":
        context.user_data["mode"] = "lyrics_search"
        await edit_query_localized(query, context, "📠اسم آهنگ و خواننده رو بفرست\nمثال: Shape of You - Ed Sheeran", reply_markup=get_back_keyboard())
        return

    if data == "video_to_mp3":
        context.user_data["mode"] = "music_finder"
        await edit_query_localized(query, context, "🔗 لینک ویدیو رو بفرست تا به MP3 تبدیل کنم.", reply_markup=get_back_keyboard())
        return

    if data == "subtitle_menu":
        context.user_data["mode"] = "subtitle"
        await edit_query_localized(query, context, "📜 لینک یوتیوب رو بفرست.", reply_markup=get_back_keyboard())
        return

    if data == "youtube_search":
        context.user_data["mode"] = "youtube_search"
        await edit_query_localized(query, context, "🔠اسم ویدیو یا آهنگ رو بنویس.", reply_markup=get_back_keyboard())
        return

    if data == "music_download_mp3":
        url = context.user_data.get("music_selected_url") or context.user_data.get("pending_url")
        if not url:
            query_text = context.user_data.get("music_download_query")
            if query_text:
                results = await search_youtube(query_text, 1)
                if results:
                    url = results[0]["url"]
        if not url:
            await edit_query_localized(query, context, "❌لینک منقضی شده.", reply_markup=get_main_keyboard())
            return
        await edit_query_localized(query, context, "⏳در حال آماده‌سازی MP3...")
        await process_download(update, context, url, quality="audio", status_message=query.message)
        return

    if data == "get_lyrics":
        artist = context.user_data.get("music_artist", "")
        title = context.user_data.get("music_title", "")
        if not title:
            await query.answer(tr("اطلاعاتی نیست", get_user_lang(query.message.chat_id, context)), show_alert=True)
            return
        await edit_query_localized(query, context, "📠در حال دریافت متن...")
        lyrics = await get_lyrics(artist, title)
        if not lyrics:
            await edit_query_localized(query, context, "😕 متن پیدا نشد.", reply_markup=get_back_keyboard())
            return
        header = f"📠<b>{html.escape(artist)} - {html.escape(title)}</b>\n\n"
        if len(lyrics) > 3800:
            parts = [lyrics[i:i+3800] for i in range(0, len(lyrics), 3800)]
            await edit_query_localized(query, context, header + f"<code>{html.escape(parts[0])}</code>", parse_mode="HTML")
            for p in parts[1:]:
                await bot_send_message_localized(context.bot, query.message.chat_id, f"<code>{html.escape(p)}</code>", parse_mode="HTML")
        else:
            await edit_query_localized(query, context, header + f"<code>{html.escape(lyrics)}</code>", parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([[LButton("🏠 منوی اصلی", callback_data="main_menu")]]))
        return

    if data.startswith("music_select_") or data.startswith("yt_select_"):
        try:
            idx = int(data.split("_")[-1])
            results = context.user_data.get("search_results", [])
            if idx >= len(results):
                await edit_query_localized(query, context, "❌نامعتبر")
                return
            item = results[idx]
            title = item["title"]
            url = item["url"]
            if " - " in title:
                artist, song_title = title.split(" - ", 1)
            else:
                artist, song_title = item.get("uploader") or "Unknown", title
            context.user_data.update({
                "music_download_query": f"{artist} - {song_title}",
                "music_artist": artist,
                "music_title": song_title,
                "music_selected_url": url,
                "pending_url": url
            })
            await edit_query_localized(query, context, 
                f"✅ <b>{html.escape(song_title)}</b>\n👤 {html.escape(artist)}",
                parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([
                    [LButton("⬇️ دانلود ویدیو", callback_data="quality_720")],
                    [LButton("🎵 دانلود MP3", callback_data="music_download_mp3")],
                    [LButton("📠متن آهنگ", callback_data="get_lyrics")],
                    [LButton("🏠 منوی اصلی", callback_data="main_menu")]
                ])
            )
        except Exception as e:
            logger.error(e)
            await edit_query_localized(query, context, "❌خطا")
        return

    if data == "cancel_download":
        cancel_flags[user_id] = True
        await edit_query_localized(query, context, "❌لغو شد.", reply_markup=get_main_keyboard())
        return

    if data.startswith("platform_"):
        msgs = {
            "instagram": "📸 لینک اینستاگرام رو بفرست",
            "youtube": "▶️ لینک یوتیوب رو بفرست",
            "tiktok": "🎵 لینک تیک‌تاک رو بفرست",
            "twitter": "🦠لینک توییتر رو بفرست",
            "reddit": "🔴 لینک ردیت رو بفرست"
        }
        await edit_query_localized(query, context, msgs.get(data.replace("platform_", ""), "لینک بفرست"), reply_markup=get_back_keyboard())
        return

    if data.startswith("quality_"):
        quality = data.replace("quality_", "")
        url = context.user_data.get("pending_url") or context.user_data.get("music_selected_url")
        if not url:
            await edit_query_localized(query, context, "❌لینک منقضی شده.", reply_markup=get_main_keyboard())
            return
        await edit_query_localized(query, context, "⏳در حال آماده‌سازی...")
        await process_download(update, context, url, quality, status_message=query.message)
        return

    if data == "change_quality":
        url = context.user_data.get("last_url")
        if not url:
            await edit_query_localized(query, context, "❌لینک قبلی نیست.")
            return
        context.user_data["pending_url"] = url
        await edit_query_localized(query, context, "🎞 کیفیت رو انتخاب کن:", reply_markup=get_youtube_quality_keyboard())
        return


async def process_download(update, context, url, quality="best", status_message=None):
    user_id = update.effective_user.id
    chat_id = update.effective_chat.id
    if user_id in blacklist:
        await bot_send_message_localized(context.bot, chat_id, "⛔️ بلاک هستید.")
        return
    if not check_daily_limit(user_id):
        msg = f"❌محدودیت روزانه تموم شده ({DAILY_LIMIT})"
        if status_message:
            await edit_text_localized(status_message, context, msg)
        else:
            await bot_send_message_localized(context.bot, chat_id, msg)
        return

    cancel_flags[user_id] = False
    platform = get_platform(url)
    is_youtube = platform == "YouTube"

    if status_message is None:
        status = await bot_send_message_localized(context.bot, chat_id, f"🔎 {platform}\n⬇️ دانلود...\n📊 ۰٪", reply_markup=get_cancel_keyboard())
    else:
        status = status_message
        await edit_text_localized(status, context, f"⬇️ دانلود از {platform}\n📊 ۰٪", reply_markup=get_cancel_keyboard())

    temp_dir = tempfile.mkdtemp(prefix="video_")
    loop = asyncio.get_running_loop()
    last_percent = [None]

    def progress_callback(percent):
        if cancel_flags.get(user_id) or percent == last_percent[0]:
            return
        last_percent[0] = percent
        async def edit():
            try:
                await edit_text_localized(status, context, f"⬇️ دانلود از {platform}\n📊 <b>{percent}</b>", parse_mode="HTML", reply_markup=get_cancel_keyboard())
            except Exception:
                pass
        asyncio.run_coroutine_threadsafe(edit(), loop)

    try:
        async with DOWNLOAD_SEMAPHORE:
            result = await asyncio.to_thread(download_media, url, temp_dir, quality, progress_callback, lambda: cancel_flags.get(user_id, False))
        if cancel_flags.get(user_id):
            await edit_text_localized(status, context, "❌لغو شد.", reply_markup=get_main_keyboard())
            return
        filename = result["filename"]
        if not os.path.exists(filename):
            raise RuntimeError("فایل پیدا نشد")
        if os.path.getsize(filename) > MAX_DOCUMENT_SIZE:
            await edit_text_localized(status, context, "❌حجم بیش از حد مجاز")
            return

        media_info = await asyncio.to_thread(get_media_info, filename)
        duration_text = format_duration(media_info["duration"])
        quality_text = f"{media_info['width']}×{media_info['height']}" if media_info.get("width") else "نامشخص"
        title = result["title"]
        size_text = format_size(os.path.getsize(filename))
        is_audio = result.get("is_audio", False)

        await edit_text_localized(status, context, f"✅ دانلود کامل\n\n📠{html.escape(title)}\n💾 {size_text} | Ⱡ{duration_text}\n🎞 {quality_text}\n\n📤 در حال ارسال...", parse_mode="HTML")
        caption = f"🎬 {title}\n💾 {size_text}\nⱠ{duration_text}\n🌠{platform}"
        async with UPLOAD_SEMAPHORE:
            await upload_file(chat_id, filename, caption, context, is_audio=is_audio)
        increase_download_count(user_id)
        context.user_data["last_url"] = url
        await edit_text_localized(status, context, "✅ با موفقیت ارسال شد.", reply_markup=get_after_download_keyboard(is_youtube))
    except Exception as e:
        logger.error(f"Download error: {e}")
        err = "لغو شد" if "لغو" in str(e) else str(e)[:250]
        try:
            await edit_text_localized(status, context, f"❌خطا\n<code>{html.escape(err)}</code>", parse_mode="HTML", reply_markup=get_main_keyboard())
        except Exception:
            pass
    finally:
        cancel_flags.pop(user_id, None)
        try:
            for f in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, f))
            os.rmdir(temp_dir)
        except Exception:
            pass


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    _CURRENT_UI_LANG.set(get_user_lang(update.effective_user.id, context))
    user_id = update.effective_user.id
    if user_id in blacklist:
        return
    mode = context.user_data.get("mode")
    if mode not in [
        "make_sticker",
        "image_sticker_wait_photo",
        "compress_image",
        "photo_to_pdf",
        "solve_image",
        "ocr_image",
    ]:
        await reply_text_localized(update.message, context, "📸 اول از منو ابزار مورد نظر رو انتخاب کن.", reply_markup=get_main_keyboard())
        return
    photo = update.message.photo[-1]
    status = await reply_text_localized(update.message, context, "⏳پردازش...")
    temp_dir = tempfile.mkdtemp(prefix="img_")
    input_file = os.path.join(temp_dir, "input.jpg")
    try:
        await (await context.bot.get_file(photo.file_id)).download_to_drive(input_file)
        if mode == "solve_image":
            result = await solve_image_question(input_file, context)
            await edit_text_localized(status, context, result, reply_markup=InlineKeyboardMarkup([[LButton("📸 سؤال جدید", callback_data="solve_image")],[LButton("🏠 منوی اصلی", callback_data="main_menu")]]))
            context.user_data["mode"] = None
        elif mode == "ocr_image":
            result = await ocr_image(input_file)
            await edit_text_localized(status, context, 
                f"🔤 <b>متن استخراج‌شده:</b>\n\n<code>{html.escape(result[:3500])}</code>",
                parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([
                    [LButton("🔤 عکس جدید", callback_data="ocr_image")],
                    [LButton("🏠 منوی اصلی", callback_data="main_menu")],
                ]),
            )
            context.user_data["mode"] = None
        elif mode in ["make_sticker", "image_sticker_wait_photo"]:
            sticker_dir = tempfile.mkdtemp(prefix="sticker_")
            sticker_photo = os.path.join(sticker_dir, "photo.jpg")
            shutil.copy2(input_file, sticker_photo)

            context.user_data["sticker_photo_path"] = sticker_photo
            context.user_data["sticker_photo_dir"] = sticker_dir
            context.user_data.pop("sticker_fg", None)
            context.user_data["mode"] = "image_sticker_wait_fg"

            await edit_text_localized(status, context, 
                "✅ عکس دریافت شد.\n\n"
                "🎨 حالا <b>رنگ متن</b> را انتخاب کن:",
                parse_mode="HTML",
                reply_markup=image_sticker_fg_keyboard(),
            )
        elif mode == "compress_image":
            img = Image.open(input_file).convert("RGB")
            out = os.path.join(temp_dir, "compressed.jpg")
            img.save(out, "JPEG", quality=60, optimize=True)
            with open(out, "rb") as f:
                await context.bot.send_photo(update.effective_chat.id, f, caption=tr(f"قبل: {format_size(os.path.getsize(input_file))}\nبعد: {format_size(os.path.getsize(out))}", get_user_lang(update.effective_user.id)))
            await edit_text_localized(status, context, "✅ فشرده شد.", reply_markup=get_main_keyboard())
        elif mode == "photo_to_pdf":
            photos = context.user_data.get("pdf_photos", [])
            photos.append(input_file)
            context.user_data["pdf_photos"] = photos
            await edit_text_localized(status, context, f"✅ عکس {len(photos)} اضافه شد.\nعکس بعدی یا بنویس: تمام", reply_markup=get_back_keyboard())
            return
        context.user_data["mode"] = None
    except Exception as e:
        await edit_text_localized(status, context, f"❌خطا: {e}", reply_markup=get_main_keyboard())
    finally:
        if mode != "photo_to_pdf":
            try:
                for f in os.listdir(temp_dir):
                    os.remove(os.path.join(temp_dir, f))
                os.rmdir(temp_dir)
            except Exception:
                pass


async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE):
    _CURRENT_UI_LANG.set(get_user_lang(update.effective_user.id, context))
    mode = context.user_data.get("mode")

    # This handler is registered before handle_music_media; delegate music-finder videos.
    if mode == "music_finder":
        await handle_music_media(update, context)
        return

    if mode == "subtitle_maker":
        video = update.message.video or update.message.document
        if not video:
            return

        target_lang = context.user_data.get("subtitle_lang", "ckb")
        duration = getattr(video, "duration", None) or 0
        if duration > MAX_SUBTITLE_VIDEO_DURATION + 5:
            await reply_text_localized(update.message, context, 
                f"⚠️ ویدیو طولانی است. فقط {MAX_SUBTITLE_VIDEO_DURATION} ثانیه اول پردازش می‌شود.",
                reply_markup=get_back_keyboard()
            )

        lang_names = {"ckb": "☀️ کوردی", "fa": "🇮🇷 فارسی", "en": "🇬🇧 انگلیسی"}
        status = await reply_text_localized(update.message, context, f"🎬 شروع ساخت زیرنویس {lang_names.get(target_lang)}...")
        temp_dir = tempfile.mkdtemp(prefix="sub_in_")
        input_video = os.path.join(temp_dir, "input.mp4")

        try:
            await (await context.bot.get_file(video.file_id)).download_to_drive(input_video)
            result = await process_subtitle(input_video, status, target_lang=target_lang)

            context.user_data["subtitle_result"] = result
            context.user_data["mode"] = None

            await edit_text_localized(status, context, 
                f"🎬 <b>زیرنویس {lang_names.get(target_lang)} آماده شد</b>\n\nیکی از گزینه‌ها رو انتخاب کن:",
                parse_mode="HTML",
                reply_markup=get_subtitle_result_keyboard()
            )

        except Exception as e:
            logger.error(f"Subtitle error: {e}")
            await edit_text_localized(status, context, 
                f"❌خطا در ساخت زیرنویس:\n<code>{html.escape(str(e)[:350])}</code>",
                parse_mode="HTML",
                reply_markup=get_main_keyboard()
            )
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:
                pass
        return

    if mode != "compress_video":
        return

    video = update.message.video or update.message.document
    if not video:
        return
    status = await reply_text_localized(update.message, context, "🗜 در حال فشرده‌سازی...")
    temp_dir = tempfile.mkdtemp(prefix="vid_")
    input_file = os.path.join(temp_dir, "input.mp4")
    output_file = os.path.join(temp_dir, "compressed.mp4")
    try:
        await (await context.bot.get_file(video.file_id)).download_to_drive(input_file)
        process = await asyncio.create_subprocess_exec(
            "ffmpeg", "-y", "-i", input_file, "-vcodec", "libx264", "-crf", "28", "-preset", "fast",
            "-acodec", "aac", "-b:a", "128k", output_file,
            stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.DEVNULL
        )
        await process.wait()
        if not os.path.exists(output_file):
            raise RuntimeError("ناموفق")
        with open(output_file, "rb") as f:
            await context.bot.send_video(update.effective_chat.id, f, caption=tr(f"قبل: {format_size(os.path.getsize(input_file))}\nبعد: {format_size(os.path.getsize(output_file))}", get_user_lang(update.effective_user.id)), supports_streaming=True)
        await edit_text_localized(status, context, "✅ فشرده شد.", reply_markup=get_main_keyboard())
        context.user_data["mode"] = None
    except Exception as e:
        await edit_text_localized(status, context, f"❌خطا: {e}", reply_markup=get_main_keyboard())
    finally:
        try:
            for f in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, f))
            os.rmdir(temp_dir)
        except Exception:
            pass


async def handle_music_media(update: Update, context: ContextTypes.DEFAULT_TYPE):
    _CURRENT_UI_LANG.set(get_user_lang(update.effective_user.id, context))
    if context.user_data.get("mode") != "music_finder":
        return
    message = update.message
    file_id = file_name = None
    if message.voice:
        file_id, file_name = message.voice.file_id, "voice.ogg"
    elif message.audio:
        file_id, file_name = message.audio.file_id, message.audio.file_name or "audio.mp3"
    elif message.video:
        file_id, file_name = message.video.file_id, "video.mp4"
    elif message.video_note:
        file_id, file_name = message.video_note.file_id, "note.mp4"
    else:
        return
    status = await reply_text_localized(message, context, "🎧 در حال شناسایی آهنگ...")
    temp_dir = tempfile.mkdtemp(prefix="music_")
    input_file = os.path.join(temp_dir, file_name)
    try:
        await (await context.bot.get_file(file_id)).download_to_drive(input_file)
        song = await recognize_music_with_audd(input_file)
        if not song:
            await edit_text_localized(status, context, "😕 شناسایی نشد.", reply_markup=get_back_keyboard())
            return
        context.user_data.update({
            "music_download_query": f"{song['artist']} - {song['title']}",
            "music_artist": song["artist"],
            "music_title": song["title"]
        })
        await edit_text_localized(status, context, 
            f"✅ <b>{html.escape(song['title'])}</b>\n👤 {html.escape(song['artist'])}",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("⬇️ دانلود MP3", callback_data="music_download_mp3")],
                [LButton("📠متن آهنگ", callback_data="get_lyrics")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ])
        )
    except Exception as e:
        logger.error(e)
        await edit_text_localized(status, context, "❌خطا در شناسایی آهنگ.", reply_markup=get_back_keyboard())
    finally:
        try:
            for f in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, f))
            os.rmdir(temp_dir)
        except Exception:
            pass


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # This handler is registered before handle_music_media; delegate music-finder voice messages.
    if context.user_data.get("mode") == "music_finder":
        await handle_music_media(update, context)
        return
    if context.user_data.get("mode") != "voice_to_text":
        return
    language = context.user_data.get("stt_language", "auto")
    status = await reply_text_localized(update.message, context, "🎧 در حال تبدیل ویس به متن...")
    temp_dir = tempfile.mkdtemp(prefix="stt_")
    voice_path = os.path.join(temp_dir, "voice.ogg")
    try:
        await (await context.bot.get_file(update.message.voice.file_id)).download_to_drive(voice_path)
        text = await speech_to_text(voice_path, language)
        if not text:
            await edit_text_localized(status, context, "❌تبدیل نشد.", reply_markup=InlineKeyboardMarkup([
                [LButton("🎤 دوباره", callback_data="voice_to_text")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ]))
            return
        await edit_text_localized(status, context, f"✅ <b>متن:</b>\n\n<code>{html.escape(text)}</code>", parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("🎤 ویس جدید", callback_data="voice_to_text")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ]))
    except Exception as e:
        logger.error(e)
        await edit_text_localized(status, context, "❌خطا")
    finally:
        try:
            for f in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, f))
            os.rmdir(temp_dir)
        except Exception:
            pass


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    register_user(user_id)
    _CURRENT_UI_LANG.set(get_user_lang(user_id, context))
    text0 = (update.message.text or "").strip()
    if text0 in {"🌐 زبان / Language", "🌐 Language", "🌐 زمان", "🌐 زمان / Language"}:
        await reply_text_localized(update.message, context, tr("لطفاً زبان را انتخاب کنید:", _CURRENT_UI_LANG.get()), reply_markup=language_keyboard())
        return
    if user_id in blacklist:
        await reply_text_localized(update.message, context, "⛔️ بلاک هستید.")
        return
    text = (update.message.text or "").strip()
    mode = context.user_data.get("mode")

    # ========== فایل‌خوان هوشمند ==========
    if mode == "smart_file_search":
        file_text = context.user_data.get("smart_file_text", "")
        if not file_text:
            context.user_data["mode"] = None
            await reply_text_localized(update.message, context, "❌ فایل فعالی وجود ندارد.", reply_markup=get_main_keyboard())
            return
        needle_text = text.casefold()
        lines = file_text.splitlines()
        matches = [line for line in lines if needle_text in line.casefold()]
        context.user_data["mode"] = "smart_file_menu"
        if matches:
            result = "🔍 <b>نتیجه جستجو</b>\n\n" + "\n".join(f"• {html.escape(x[:700])}" for x in matches[:20])
            if len(matches) > 20:
                result += f"\n\n… و {len(matches)-20} نتیجه دیگر."
        else:
            result = "🔍 چیزی با این عبارت در متن استخراج‌شده پیدا نشد."
        await reply_text_localized(update.message, context, result, parse_mode="HTML", reply_markup=smart_file_keyboard())
        return

    if mode in ("smart_file_question", "smart_file_translate"):
        file_text = context.user_data.get("smart_file_text", "")
        action = "question" if mode == "smart_file_question" else "translate"
        status = await reply_text_localized(update.message, context, "⏳ در حال پردازش فایل...")
        result = await smart_file_ai_action(file_text, action, text)
        context.user_data["mode"] = "smart_file_menu"
        parts = smart_split_message(result)
        await edit_text_localized(status, context, parts[0])
        for part in parts[1:]:
            await reply_text_localized(update.message, context, part)
        await reply_text_localized(update.message, context, "📁 عملیات دیگری روی همین فایل؟", reply_markup=smart_file_keyboard())
        return

    # ========== دکمه‌های پایین، مستقل از زبان ==========
    # Telegram sends the localized label, so never compare only with Persian text.
    bottom_actions = {
        "home": {"🏠 منوی اصلی", "🏠 Main Menu", "🏠 سەرەکی"},
        "restart": {"🔄 ری‌استارت", "🔄 Restart", "🔄 دووبارە دەستپێکردن"},
        "help": {"ℹ️ راهنما", "ℹ️ Help", "ℹ️ ڕێنمایی"},
        "language": {"🌐 زبان / Language", "🌐 Language", "🌐 زمان"},
    }
    if text in bottom_actions["language"]:
        await reply_text_localized(update.message, context, "لطفاً زبان را انتخاب کنید:", reply_markup=language_keyboard())
        return
    if text in bottom_actions["restart"]:
        current_lang = get_user_lang(user_id, context)
        context.user_data.clear()
        context.user_data["language"] = current_lang
        _CURRENT_UI_LANG.set(current_lang)
        await reply_text_localized(update.message, context, f"🔄 ری‌استارت شد\n\n{get_jalali_datetime(current_lang)}\n\n🌟 AsoLand", parse_mode="HTML", reply_markup=get_reply_keyboard())
        await reply_text_localized(update.message, context, "منوی اصلی:", reply_markup=get_main_keyboard())
        return
    if text in bottom_actions["home"]:
        current_lang = get_user_lang(user_id, context)
        context.user_data.clear()
        context.user_data["language"] = current_lang
        _CURRENT_UI_LANG.set(current_lang)
        await reply_text_localized(update.message, context, f"🌟 <b>AsoLand</b>\n\n{get_jalali_datetime(current_lang)}", parse_mode="HTML", reply_markup=get_reply_keyboard())
        await reply_text_localized(update.message, context, "منوی اصلی:", reply_markup=get_main_keyboard())
        return
    if text in bottom_actions["help"]:
        await help_command(update, context)
        return
    # ===================================

    
    if mode == "currency_convert":
        status = await reply_text_localized(update.message, context, "💱 در حال تبدیل...")
        result = await convert_currency(text)
        await edit_text_localized(status, context, result, parse_mode="HTML", reply_markup=InlineKeyboardMarkup([
            [LButton("💱 تبدیل جدید", callback_data="currency_convert")],
            [LButton("🏠 منوی اصلی", callback_data="main_menu")],
        ]))
        return

    if mode == "reminder":
        parsed = parse_reminder_time(text)
        if not parsed:
            await reply_text_localized(update.message, context, 
                "❌ فرمت یادآور نامعتبر.\nمثال: یادآوری ساعت ۱۰ جلسه",
                reply_markup=get_back_keyboard(),
            )
            return
        due_ts, body = parsed
        reminders[user_id].append({"text": body, "due_ts": due_ts, "created": time.time()})
        save_reminders()
        from datetime import datetime as _dt
        due_str = _dt.fromtimestamp(due_ts).strftime("%Y-%m-%d %H:%M")
        context.user_data["mode"] = None
        await reply_text_localized(update.message, context, 
            f"✅ یادآور ثبت شد.\n📌 {html.escape(body)}\n🕐 {due_str}",
            parse_mode="HTML",
            reply_markup=extended_tools_keyboard(),
        )
        return

    if mode == "image_sticker_wait_text":
        photo_path = context.user_data.get("sticker_photo_path")
        font_key = context.user_data.get("sticker_font", DEFAULT_STICKER_FONT)
        fg_key = context.user_data.get("sticker_fg", "fg_white")
        fg_rgba = STICKER_FG_COLORS.get(fg_key, STICKER_FG_COLORS["fg_white"])[0]

        if not photo_path or not os.path.exists(photo_path):
            context.user_data["mode"] = "image_sticker_wait_photo"
            await reply_text_localized(update.message, context, 
                "📸 عکس پیدا نشد. دوباره عکس را بفرست.",
                reply_markup=get_back_keyboard(),
            )
            return

        status = await reply_text_localized(update.message, context, "🎨 در حال ساخت استیکر...")
        try:
            bio = await asyncio.to_thread(
                create_image_text_sticker,
                photo_path,
                text,
                font_key,
                512,
                fg_rgba,
            )
            if not bio:
                await edit_text_localized(status, context, 
                    "❌ ساخت استیکر ناموفق بود.",
                    reply_markup=get_main_keyboard(),
                )
                return

            await context.bot.send_sticker(update.effective_chat.id, bio)
            await edit_text_localized(status, context, 
                "✅ <b>استیکر آماده شد.</b>\n\n"
                "متن جدید بفرست تا با همین عکس و فونت، استیکر دیگری ساخته شود.",
                parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([
                    [LButton("🎨 تغییر رنگ متن", callback_data="image_sticker_change_color")],
                    [LButton("🔤 تغییر فونت", callback_data="image_sticker_change_font")],
                    [LButton("📸 عکس جدید", callback_data="sticker_with_photo")],
                    [LButton("🎨 ساخت بدون عکس", callback_data="sticker_no_photo")],
                    [LButton("🏠 منوی اصلی", callback_data="main_menu")],
                ]),
            )
        except Exception as e:
            logger.error(f"image_sticker_text handler error: {e}")
            await edit_text_localized(status, context, 
                f"❌ خطا در ساخت استیکر:\n<code>{html.escape(str(e)[:250])}</code>",
                parse_mode="HTML",
                reply_markup=get_main_keyboard(),
            )
        return

    if mode == "text_sticker":
        # این حالت همیشه استیکر با پس‌زمینه رنگی می‌سازد، نه با عکس قبلی.
        bg_key = context.user_data.get("sticker_bg", "bg_blue")
        fg_key = context.user_data.get("sticker_fg", "fg_white")
        font_key = context.user_data.get("sticker_font", DEFAULT_STICKER_FONT)
        bg = STICKER_BG_COLORS.get(bg_key, STICKER_BG_COLORS["bg_blue"])[0]
        fg = STICKER_FG_COLORS.get(fg_key, STICKER_FG_COLORS["fg_white"])[0]
        if not _rtl_packages_ok():
            await reply_text_localized(update.message, context, 
                "⚠️ برای چسبیدن حروف فارسی/کوردی یکی از این‌ها لازم است:\n\n"
                "۱) Pillow با RAQM (بهترین برای کوردی):\n"
                "<code>pkg install libraqm harfbuzz fribidi</code>\n"
                "<code>pip install -U pillow</code>\n\n"
                "۲) یا حداقل:\n"
                "<code>pip install arabic-reshaper python-bidi==0.4.2</code>\n\n"
                "فونت Noto Naskh Arabic هم خودکار دانلود می‌شود.\n"
                "بعد ربات را کامل ببند و دوباره روشن کن.",
                parse_mode="HTML",
                reply_markup=get_back_keyboard(),
            )
            return
        status = await reply_text_localized(update.message, context, "🎨 در حال ساخت استیکر...")
        try:
            bio = await asyncio.to_thread(create_text_sticker, text, bg, fg, 512, font_key)
            if not bio:
                await edit_text_localized(status, context, "❌ ساخت استیکر ناموفق بود.", reply_markup=get_main_keyboard())
                return
            bio.name = "sticker.webp"
            await context.bot.send_sticker(update.effective_chat.id, bio)
            await edit_text_localized(status, context, 
                "✅ استیکر آماده شد.\nمی‌تونی متن جدید بفرستی یا تنظیمات را عوض کنی.",
                reply_markup=InlineKeyboardMarkup([
                    [LButton("🔤 تغییر فونت", callback_data="sticker_pick_font")],
                    [LButton("🎨 تغییر پس‌زمینه", callback_data="text_sticker")],
                    [LButton("🖼 ساخت با عکس", callback_data="sticker_with_photo")],
                    [LButton("🏠 منوی اصلی", callback_data="main_menu")],
                ]),
            )
        except Exception as e:
            logger.error(f"text_sticker handler error: {e}")
            await edit_text_localized(status, context, 
                f"❌ خطا در ساخت استیکر:\n<code>{html.escape(str(e)[:200])}</code>",
                parse_mode="HTML",
                reply_markup=get_main_keyboard(),
            )
        return

    if mode == "weather":
        status = await reply_text_localized(update.message, context, "🌤 در حال دریافت وضعیت آب‌وهوا...")
        result = await get_weather(text)
        await edit_text_localized(status, context, result, parse_mode="HTML", reply_markup=InlineKeyboardMarkup([
            [LButton("🌤 شهر دیگر", callback_data="weather_menu")],
            [LButton("🏠 منوی اصلی", callback_data="main_menu")]
        ]))
        return

    if mode == "advanced_calculator":
        result = calculate_advanced(text)
        await reply_text_localized(update.message, context, result, parse_mode="HTML", reply_markup=InlineKeyboardMarkup([
            [LButton("🧮 محاسبه جدید", callback_data="advanced_calculator")],
            [LButton("🏠 منوی اصلی", callback_data="main_menu")]
        ]))
        return

    if mode == "weather_chart":
        status=await reply_text_localized(update.message, context, "📊 در حال ساخت نمودار دما...")
        path=await get_weather_chart(text)
        if path:
            with open(path,"rb") as f: await context.bot.send_photo(update.effective_chat.id,f,caption=tr(f"🌤 نمودار ۷ روزه {text}", get_user_lang(update.effective_user.id)))
            try: os.remove(path)
            except: pass
            await edit_text_localized(status, context, "✅ نمودار آماده شد.",reply_markup=extended_tools_keyboard())
        else: await edit_text_localized(status, context, "❌نمودار آب‌وهوا ساخته نشد.",reply_markup=extended_tools_keyboard())
        context.user_data["mode"]=None; return

    if mode == "price_alert":
        m=re.match(r"^\s*([^\s]+)\s+(above|below|بالا|پایین)\s+([0-9۰-۹٠-٩,٬.]+)\s*$",text,re.I)
        if not m:
            await reply_text_localized(update.message, context, "❌ فرمت نادرست. مثال: BTC above 100000",reply_markup=get_back_keyboard()); return
        symbol=m.group(1); direction="above" if m.group(2).lower() in ("above","بالا") else "below"
        target=float(m.group(3).translate(str.maketrans("۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩","01234567890123456789")).replace(",","").replace("٬",""))
        if await get_alert_price(symbol) is None:
            await reply_text_localized(update.message, context, "❌این نماد پشتیبانی نمی‌شود. دلار، طلا یا سکه را امتحان کن.",reply_markup=get_back_keyboard()); return
        await set_price_alert(user_id,symbol,target,direction)
        context.user_data["mode"]=None
        await reply_text_localized(update.message, context, f"✅ هشدار ثبت شد: {symbol} {direction} {target:,.2f}",reply_markup=extended_tools_keyboard()); return

    if mode == "student_text":
        await update.message.chat.send_action(ChatAction.TYPING)
        reply=await solve_student_problem_with_context(text,context)
        await reply_text_localized(update.message, context, reply,reply_markup=InlineKeyboardMarkup([[LButton("🎓 سؤال جدید",callback_data="student_menu")],[LButton("🏠 منوی اصلی",callback_data="main_menu")]]))
        return

    if mode == "smart_ai":
        await update.message.chat.send_action(ChatAction.TYPING)
        reply=await ask_ai(text,context)
        await reply_text_localized(update.message, context, reply,reply_markup=get_main_keyboard()); context.user_data["mode"]=None; return

    if mode == "summarize":
        if not text or len(text) < 30:
            await reply_text_localized(update.message, context, "❌متن خیلی کوتاهه.", reply_markup=get_back_keyboard())
            return
        await update.message.chat.send_action(ChatAction.TYPING)
        status = await reply_text_localized(update.message, context, "📠در حال خلاصه‌سازی...")
        summary = await summarize_text(text)
        await edit_text_localized(status, context, f"📠<b>خلاصه متن:</b>\n\n{html.escape(summary)}", parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [LButton("📠خلاصه جدید", callback_data="summarize_text")],
                [LButton("🏠 منوی اصلی", callback_data="main_menu")]
            ]))
        return

    if mode == "text_to_word":
        if not text:
            await reply_text_localized(update.message, context, "❌متن بفرست.", reply_markup=get_back_keyboard())
            return
        status = await reply_text_localized(update.message, context, "📠در حال ساخت فایل ورد...")
        try:
            buffer = create_word_document(text)
            await context.bot.send_document(chat_id=update.effective_chat.id, document=buffer, filename="AsoLand.docx", caption=tr("✅ فایل ورد آماده شد.", get_user_lang(update.effective_user.id)))
            await edit_text_localized(status, context, "✅ فایل ورد ارسال شد.", reply_markup=get_main_keyboard())
        except Exception as e:
            logger.error(f"Word error: {e}")
            await edit_text_localized(status, context, "❌خطا در ساخت فایل ورد.", reply_markup=get_back_keyboard())
        context.user_data["mode"] = None
        return

    if mode == "fancy_text":
        if not text:
            await reply_text_localized(update.message, context, "❌متن بفرست.", reply_markup=get_back_keyboard())
            return
        fonts = convert_to_fancy_fonts(text)
        result = "✅ <b>فونت‌های زیبا:</b>\n\n"
        for k, v in fonts.items():
            result += f"<b>{k}:</b>\n<code>{html.escape(v)}</code>\n\n"
        if len(result) > 4000:
            for i in range(0, len(result), 4000):
                await reply_text_localized(update.message, context, result[i:i+4000], parse_mode="HTML")
        else:
            await reply_text_localized(update.message, context, result, parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([
                    [LButton("✅ متن جدید", callback_data="fancy_text")],
                    [LButton("🏠 منوی اصلی", callback_data="main_menu")]
                ]))
        return

    if mode == "english_teacher":
        et_mode = context.user_data.get("et_mode", "conversation")
        if et_mode == "menu":
            await reply_text_localized(update.message, context, "از دکمه‌ها استفاده کن:", reply_markup=get_english_teacher_keyboard())
            return
        await update.message.chat.send_action(ChatAction.TYPING)
        reply = await ask_ai(text, context, get_english_teacher_prompt(context.user_data.get("english_level", "intermediate"), et_mode))
        teacher_markup = InlineKeyboardMarkup([[LButton("🔙 منوی معلم", callback_data="english_teacher")]])
        parts = smart_split_message(reply, 3900)
        for index, part in enumerate(parts):
            await reply_text_localized(update.message, context, part, reply_markup=teacher_markup if index == len(parts)-1 else None)
        return

    if mode == "broadcast" and is_admin(user_id):
        context.user_data.pop("mode", None)
        targets = sorted(all_users - blacklist)
        if not targets:
            await reply_text_localized(update.message, context, "❌ هیچ کاربری برای ارسال وجود ندارد.", reply_markup=get_admin_keyboard())
            return
        status = await reply_text_localized(update.message, context, f"📢 در حال ارسال به {len(targets)} کاربر...")
        success = fail = 0
        for uid in targets:
            try:
                await bot_send_message_localized(context.bot, uid, text, parse_mode="HTML")
                success += 1
                await asyncio.sleep(0.05)
            except Exception:
                fail += 1
        await edit_text_localized(status, context, 
            f"✅ پیام همگانی تمام شد\nموفق: <b>{success}</b>\nناموفق: <b>{fail}</b>\nکل هدف: <b>{len(targets)}</b>",
            parse_mode="HTML",
            reply_markup=get_admin_keyboard()
        )
        return

    if mode == "photo_to_pdf" and text.lower() in ["تمام", "تموم", "done"]:
        photos = context.user_data.get("pdf_photos", [])
        if not photos:
            await reply_text_localized(update.message, context, "❌عکسی نیست.", reply_markup=get_main_keyboard())
            context.user_data.clear()
            return
        status = await reply_text_localized(update.message, context, "📄 ساخت PDF...")
        try:
            images = [Image.open(p).convert("RGB") for p in photos if os.path.exists(p)]
            if not images:
                raise RuntimeError("عکسی نیست")
            pdf_path = os.path.join(tempfile.gettempdir(), f"pdf_{user_id}.pdf")
            images[0].save(pdf_path, save_all=True, append_images=images[1:] if len(images) > 1 else [])
            with open(pdf_path, "rb") as f:
                await context.bot.send_document(update.effective_chat.id, f, caption=tr(f"📄 PDF با {len(images)} صفحه", get_user_lang(update.effective_user.id)))
            await edit_text_localized(status, context, "✅ آماده شد.", reply_markup=get_main_keyboard())
        except Exception as e:
            await edit_text_localized(status, context, f"❌خطا: {e}", reply_markup=get_main_keyboard())
        finally:
            context.user_data.clear()
            for p in photos:
                try: os.remove(p)
                except: pass
            try: os.remove(pdf_path)
            except: pass
        return

    if mode == "lyrics_search":
        status = await reply_text_localized(update.message, context, "📠جستجو...")
        artist, title = ("", text)
        if " - " in text:
            artist, title = text.split(" - ", 1)
        lyrics = await get_lyrics(artist.strip(), title.strip())
        if not lyrics:
            await edit_text_localized(status, context, "😕 پیدا نشد.", reply_markup=get_back_keyboard())
            return
        header = f"📠<b>{html.escape(text)}</b>\n\n"
        if len(lyrics) > 3800:
            parts = [lyrics[i:i+3800] for i in range(0, len(lyrics), 3800)]
            await edit_text_localized(status, context, header + f"<code>{html.escape(parts[0])}</code>", parse_mode="HTML")
            for p in parts[1:]:
                await reply_text_localized(update.message, context, f"<code>{html.escape(p)}</code>", parse_mode="HTML")
        else:
            await edit_text_localized(status, context, header + f"<code>{html.escape(lyrics)}</code>", parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup([[LButton("🏠 منوی اصلی", callback_data="main_menu")]]))
        return

    if mode == "qr_code":
        try:
            qr = qrcode.QRCode(version=1, box_size=10, border=4)
            qr.add_data(text)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            bio = BytesIO()
            img.save(bio, "PNG")
            bio.seek(0)
            await update.message.reply_photo(bio, caption=f"<code>{html.escape(text[:200])}</code>", parse_mode="HTML")
            context.user_data["mode"] = None
        except Exception as e:
            await reply_text_localized(update.message, context, f"❌خطا: {e}")
        return

    if mode == "short_link":
        if not text.startswith("http"):
            await reply_text_localized(update.message, context, "❌لینک معتبر بفرست.", reply_markup=get_back_keyboard())
            return
        status = await reply_text_localized(update.message, context, "🔗 در حال کوتاه کردن...")
        short = await short_url(text)
        if short:
            await edit_text_localized(status, context, f"✅ <code>{short}</code>", parse_mode="HTML", reply_markup=get_main_keyboard())
        else:
            await edit_text_localized(status, context, "❌خطا", reply_markup=get_back_keyboard())
        context.user_data["mode"] = None
        return

    if mode == "translate":
        await update.message.chat.send_action(ChatAction.TYPING)
        result = await translate_text(text)
        await reply_text_localized(update.message, context, f"🌠<b>ترجمه:</b>\n\n{html.escape(result)}", parse_mode="HTML", reply_markup=get_back_keyboard())
        return

    if mode == "ai":
        await update.message.chat.send_action(ChatAction.TYPING)
        reply = await ask_ai(text, context)
        parts = smart_split_message(reply, 3900)
        # AI chat is always Persian, regardless of the selected UI language.
        # Do not run the AI response through the Kurdish/English UI translator.
        for index, part in enumerate(parts):
            await update.message.reply_text(
                part,
                reply_markup=get_ai_keyboard() if index == len(parts)-1 else None,
            )
        return

    if mode == "youtube_search":
        status = await reply_text_localized(update.message, context, "🔠جستجو...")
        results = await search_youtube(text, 6)
        if not results:
            await edit_text_localized(status, context, "😕 چیزی پیدا نشد.", reply_markup=get_back_keyboard())
            return
        context.user_data["search_results"] = results
        lines = ["🎵 نتایج:\n"]
        buttons = []
        for i, item in enumerate(results):
            lines.append(f"{i+1}. {html.escape(item['title'][:50])}")
            buttons.append([LButton(f"{i+1}. {item['title'][:40]}", callback_data=f"yt_select_{i}")])
        buttons.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
        await edit_text_localized(status, context, "\n".join(lines), parse_mode="HTML", reply_markup=InlineKeyboardMarkup(buttons))
        return

    if mode == "music_finder":
        url = extract_url(text)
        if url:
            await process_download(update, context, url, quality="audio")
            return
        status = await reply_text_localized(update.message, context, "🔠جستجو...")
        results = await search_youtube(text, 5)
        if not results:
            await edit_text_localized(status, context, "😕 پیدا نشد.", reply_markup=get_back_keyboard())
            return
        context.user_data["search_results"] = results
        lines = ["🎵 نتایج:\n"]
        buttons = []
        for i, item in enumerate(results):
            lines.append(f"{i+1}. {html.escape(item['title'][:50])}")
            buttons.append([LButton(f"{i+1}. {item['title'][:38]}", callback_data=f"music_select_{i}")])
        buttons.append([LButton("🏠 منوی اصلی", callback_data="main_menu")])
        await edit_text_localized(status, context, "\n".join(lines), parse_mode="HTML", reply_markup=InlineKeyboardMarkup(buttons))
        return

    if mode == "subtitle":
        url = extract_url(text)
        if not url or get_platform(url) != "YouTube":
            await reply_text_localized(update.message, context, "❌فقط لینک یوتیوب.", reply_markup=get_back_keyboard())
            return
        status = await reply_text_localized(update.message, context, "📜 دریافت زیرنویس...")
        temp_dir = tempfile.mkdtemp(prefix="sub_")
        try:
            await asyncio.to_thread(download_media, url, temp_dir, download_subs=True)
            subs = [f for f in os.listdir(temp_dir) if f.endswith((".vtt", ".srt"))]
            if not subs:
                await edit_text_localized(status, context, "😕 زیرنویس پیدا نشد.", reply_markup=get_main_keyboard())
                return
            for sf in subs[:2]:
                with open(os.path.join(temp_dir, sf), "rb") as f:
                    await context.bot.send_document(update.effective_chat.id, f, caption=tr(f"📜 {sf}", get_user_lang(update.effective_user.id)))
            await edit_text_localized(status, context, "✅ ارسال شد.", reply_markup=get_main_keyboard())
        except Exception as e:
            await edit_text_localized(status, context, f"❌{str(e)[:200]}", reply_markup=get_main_keyboard())
        finally:
            try:
                for f in os.listdir(temp_dir):
                    os.remove(os.path.join(temp_dir, f))
                os.rmdir(temp_dir)
            except Exception:
                pass
        context.user_data["mode"] = None
        return

    if not mode:
        intent,value=await smart_route(text,context)
        if intent == "weather_direct":
            status=await reply_text_localized(update.message, context, "🌤 در حال دریافت آب‌وهوا...")
            result=await get_weather(value)
            await edit_text_localized(status, context, result,parse_mode="HTML",reply_markup=extended_tools_keyboard()); return
        if intent == "price_do_l4":
            status = await reply_text_localized(update.message, context, "⏳ در حال دریافت قیمت دلار و سکه...")
            result = await get_do_l4_prices()
            await edit_text_localized(status, context, result, parse_mode="HTML", reply_markup=extended_tools_keyboard())
            return
        if intent == "currency_convert":
            context.user_data["mode"] = "currency_convert"
            await reply_text_localized(update.message, context, "💱 مقدار را بفرست؛ مثال: ۱۰۰ دلار", reply_markup=get_back_keyboard())
            return
        if intent == "today_calendar":
            await reply_text_localized(update.message, context, get_today_calendar(), parse_mode="HTML", reply_markup=extended_tools_keyboard())
            return
        if intent == "reminder_menu":
            context.user_data["mode"] = "reminder"
            await reply_text_localized(update.message, context, "🔔 مثال: یادآوری ساعت ۱۰ جلسه", reply_markup=get_back_keyboard())
            return
        if intent == "user_panel":
            name = update.effective_user.first_name or "کاربر"
            await reply_text_localized(update.message, context, get_user_panel(user_id, name), parse_mode="HTML", reply_markup=extended_tools_keyboard())
            return
        if intent in ("weather_menu","daily_fortune","news_menu","chart_menu","alert_menu","student_menu"):
            if intent=="daily_fortune":
                result=await get_daily_fortune(update.effective_user.first_name or "دوست من")
                await reply_text_localized(update.message, context, result,parse_mode="HTML",reply_markup=extended_tools_keyboard()); return
            if intent=="weather_menu":
                context.user_data["mode"]="weather"; await reply_text_localized(update.message, context, "🌤 نام شهر را بفرست.",reply_markup=get_back_keyboard()); return
            if intent=="student_menu":
                context.user_data["mode"]="student_text"; await reply_text_localized(update.message, context, "🎓 سؤال درسی را بفرست.",reply_markup=get_back_keyboard()); return
            if intent=="news_menu":
                await reply_text_localized(update.message, context, "📰 دسته خبر را از منو انتخاب کن.",reply_markup=extended_tools_keyboard()); return
            if intent=="chart_menu":
                await reply_text_localized(update.message, context, "📊 از منوی امکانات هوشمند نمودار را انتخاب کن.",reply_markup=extended_tools_keyboard()); return
            if intent=="alert_menu":
                context.user_data["mode"]="price_alert"; await reply_text_localized(update.message, context, "🔔 مثال:\nدلار above 190000\nطلا below 20000000",reply_markup=get_back_keyboard()); return

    url = extract_url(text)
    if not url:
        await reply_text_localized(update.message, context, "❌لینک معتبر پیدا نشد.\nاز دکمه‌ها استفاده کن.", reply_markup=get_main_keyboard())
        return
    if get_platform(url) == "YouTube":
        context.user_data["pending_url"] = url
        await reply_text_localized(update.message, context, "▶️ یوتیوب شناسایی شد\nکیفیت رو انتخاب کن:", reply_markup=get_youtube_quality_keyboard())
        return
    await process_download(update, context, url, quality="best")


async def post_init(app):
    # بعد از استارت اپلیکیشن، حلقه‌های پس‌زمینه را اجرا کن
    asyncio.create_task(price_alert_loop(app))
    asyncio.create_task(reminder_loop(app))
    # دانلود فونت استیکر فارسی (اگر نباشد)
    try:
        await asyncio.to_thread(_ensure_sticker_font)
    except Exception as e:
        logger.warning(f"Sticker font init: {e}")


def main():
    if not _valid_api_key(TOKEN):
        raise RuntimeError("TOKEN تنظیم نشده است؛ توکن BotFather را در بخش API KEYS / TOKENS قرار بده.")
    load_data()
    load_languages()
    app = (
        Application.builder()
        .token(TOKEN)
        .connect_timeout(30.0)
        .read_timeout(30.0)
        .write_timeout(30.0)
        .pool_timeout(30.0)
        .get_updates_connect_timeout(30.0)
        .get_updates_read_timeout(50.0)
        .get_updates_pool_timeout(30.0)
        .connection_pool_size(8)
        .post_init(post_init)
        .build()
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("admin", admin_command))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CallbackQueryHandler(button_handler))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    # Music Finder باید قبل از Handlerهای عمومی صدا/ویدیو ثبت شود.
    app.add_handler(MessageHandler(filters.VOICE | filters.AUDIO | filters.VIDEO_NOTE, handle_music_media))
    app.add_handler(MessageHandler(filters.VIDEO | filters.Document.VIDEO, handle_video))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    # Smart File Reader must be registered before the generic document handlers.
    app.add_handler(MessageHandler(filters.Document.ALL, handle_smart_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    logger.info("AsoLand Bot started")
    print("=" * 50)
    print("  AsoLand Bot is running...")
    print("=" * 50)
    app.run_polling(drop_pending_updates=True)


if __name__ == "__main__":
    main()
